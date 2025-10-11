import io
import streamlit as st
import pandas as pd
import numpy as np 
import re
import xlsxwriter
import random
from openpyxl import load_workbook # Ensure load_workbook is imported
import io, zipfile
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from datetime import timedelta
from xlsxwriter import Workbook as Workbook
from collections import defaultdict
from datetime import datetime, timedelta
from collections import Counter
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED
import re
from collections import defaultdict
from copy import copy
from openpyxl import load_workbook, Workbook
from openpyxl.utils import get_column_letter
from openpyxl.cell.cell import MergedCell
from openpyxl.utils import column_index_from_string
from openpyxl.styles import Font, PatternFill, Color

st.set_page_config(page_title="PSUCOM PEDIATRIC CLERKSHIP SCHEDULE CREATOR", layout="wide")
st.title("PSUCOM PEDIATRIC CLERKSHIP SCHEDULE CREATOR")

# ─── Sidebar mode selector ─────────────────────────────────────────────────────
mode = st.sidebar.radio("What do you want to do?",("Instructions", "Format OPD + Summary", "Create Student Schedule","OPD Check","Create Individual Schedules","OPD MD PA Conflict Detector","Shift Availability Tracker"))
# ─── Sidebar mode selector ─────────────────────────────────────────────────────

if mode == "OPD Check":
    DAYS = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
    
    def detect_am_pm_blocks(df):
        """
        Scan down column A and group runs whose cell text begins with 'AM' or 'PM',
        ignoring everything else.
        Returns a list of (label, start_row, end_row) tuples.
        """
        runs = []
        current_label = None
        run_start = None
        prev_row = None
    
        # Excel row 6 is df index 5
        for idx in range(5, len(df)):
            raw = df.iat[idx, 0]
            if isinstance(raw, str):
                ru = raw.upper()
                if ru.startswith("AM"):
                    label = "AM"
                elif ru.startswith("PM"):
                    label = "PM"
                else:
                    continue
            else:
                continue
    
            row = idx + 1  # convert back to 1‑based Excel row
            if current_label is None:
                # start new run
                current_label = label
                run_start = row
            elif label != current_label or row != prev_row + 1:
                # close out previous run
                runs.append((current_label, run_start, prev_row))
                current_label = label
                run_start = row
    
            prev_row = row
    
        # finish last run
        if current_label is not None:
            runs.append((current_label, run_start, prev_row))
    
        return runs
    
    st.title("OPD PRECEPTOR CHECK")
    
    baseline_file = st.file_uploader(
        "1) Upload Latest Updated OPD", 
        type=["xlsx"], key="baseline"
    )
    assigned_file = st.file_uploader(
        "2) Upload OPD with Student Assignments", 
        type=["xlsx"], key="assigned"
    )
    
    if baseline_file and assigned_file:
        SHEETS = [
            'HOPE_DRIVE','ETOWN','NYES','COMPLEX','WARD A',
            'PSHCH_NURSERY','HAMPDEN_NURSERY','SJR_HOSP','AAC',
            'AHOLOUKPE','ADOLMED'
        ]
    
        # Read all relevant sheets at once
        base_sheets = pd.read_excel(baseline_file, sheet_name=SHEETS, header=None)
        assn_sheets = pd.read_excel(assigned_file, sheet_name=SHEETS, header=None)

        # 1) Remove student suffix from every baseline sheet
        for sheet_name, df in base_sheets.items():
            for col in df.columns[1:]:
                df[col] = (
                    df[col]
                      .where(df[col].notna(), np.nan)      # keep NaN
                      .astype(str)
                      .str.partition('~')[0]               # take text before "~"
                      .replace({'nan': np.nan})            # restore NaN
                )
            base_sheets[sheet_name] = df
    
        results = {}
        
        for sheet in SHEETS:
            df_base = base_sheets[sheet]
            df_assn = assn_sheets[sheet]
        
            runs = detect_am_pm_blocks(df_base)
            week_pairs = [(runs[i], runs[i+1]) for i in range(0, len(runs), 2)]
        
            base_map = {}    # (w,period,day,pre) → (cell, student)
            assn_map = {}
        
            for w_idx, ((am_lbl, am_s, am_e), (pm_lbl, pm_s, pm_e)) in enumerate(week_pairs, start=1):
                for period, (lbl, start, end) in [('AM',(am_lbl,am_s,am_e)), ('PM',(pm_lbl,pm_s,pm_e))]:
                    for col, day in enumerate(DAYS, start=1):
                        for row in range(start, end+1):
                            cell = f"{chr(ord('A')+col)}{row}"
        
                            # baseline
                            vb = df_base.iat[row-1, col]
                            if pd.notna(vb) and isinstance(vb, str):
                                parts = str(vb).split('~', 1)
                                pre = parts[0].strip()
                                stu = parts[1].strip() if len(parts)==2 else None
                                base_map.setdefault((w_idx,period,day,pre), (cell, stu))
        
                            # assigned
                            va = df_assn.iat[row-1, col]
                            if pd.notna(va) and isinstance(va, str):
                                parts = str(va).split('~', 1)
                                pre = parts[0].strip()
                                stu = parts[1].strip() if len(parts)==2 else None
                                assn_map.setdefault((w_idx,period,day,pre), (cell, stu))
        
            dropped = []
            added   = []
        
            # drops: in base not in assigned
            for key, (cell, stu) in base_map.items():
                if key not in assn_map:
                    w,p,d,pre = key
                    dropped.append((w,p,d,pre,cell,stu))
        
            # adds: in assigned not in base
            for key, (cell, stu) in assn_map.items():
                if key not in base_map:
                    w,p,d,pre = key
                    added.append((w,p,d,pre,cell,stu))
        
            # sort exactly as before...
            dropped_sorted = sorted(
                dropped,
                key=lambda x: (x[0], {'AM':0,'PM':1}[x[1]], DAYS.index(x[2]), x[4])
            )
            added_sorted = sorted(
                added,
                key=lambda x: (x[0], {'AM':0,'PM':1}[x[1]], DAYS.index(x[2]), x[4])
            )
        
            results[sheet] = {"dropped": dropped_sorted, "added": added_sorted}



        
        #
    
    doc = Document()
    doc.add_heading('Change Report', level=1)
    
    for sheet, change in (locals().get('results') or {}).items():
        doc.add_heading(sheet, level=2)
    
        # build week→day map (collect both AM & PM under each day)
        week_map = defaultdict(lambda: defaultdict(lambda: {'dropped': [], 'added': []}))
        for w,p,d,pre,cell,stu in change['dropped']:
            week_map[w][d]['dropped'].append((p, pre, cell, stu))
        for w,p,d,pre,cell,stu in change['added']:
            week_map[w][d]['added'].append((p, pre, cell, stu))
    
        # emit in week order
        for w in sorted(week_map):
            doc.add_heading(f'Week {w}', level=3)
    
            for day in DAYS:
                slot = week_map[w].get(day)
                if not slot or (not slot['dropped'] and not slot['added']):
                    continue
    
                doc.add_heading(day, level=4)
    
                # DROPS
                for p, pre, cell, stu in slot['dropped']:
                    line = f"- Dropped: {pre} — was at {cell}"
                    if stu:
                        line += f"  (Student impacted: {stu})"
                    doc.add_paragraph(line, style='List Bullet')
    
                # ADDS
                for p, pre, cell, stu in slot['added']:
                    line = f"- Added: {pre} — now at {cell}"
                    if stu:
                        line += f"  (Student impacted: {stu})"
                    doc.add_paragraph(line, style='List Bullet')
    
            doc.add_paragraph()  # blank line between weeks




    # Save to in-memory buffer
    word_file = io.BytesIO()
    doc.save(word_file)
    word_file.seek(0)
    
    # Download button
    st.download_button(
        label="📄 Download Word Report",
        data=word_file,
        file_name="change_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


elif mode == "Instructions":
    d = st.text_input('Start date (m/d/yyyy)')
    if d:
        try:
            s = datetime.strptime(d, '%m/%d/%Y')
            e = s + timedelta(days=34)
            st.write(f"{s:%B %d, %Y} → {e:%B %d, %Y}")
    
            st.write('Please go to https://login.qgenda.com/')
    
            # Display on‐screen instructions
            st.markdown(f"""Download four files and create reports based on **{s:%B %d, %Y}** → **{e:%B %d, %Y}**.""")
            st.write("Download instructions here:")


        except ValueError:
            st.error('Invalid format – use m/d/yyyy (e.g. 7/6/2021)')

        # --- Generate a Word document with the same instructions ---
        doc = Document()
        doc.add_heading('Qgenda Report Instructions', level=1)
        doc.styles['Normal'].font.size = Pt(8)
        
        # Bold the date range
        p = doc.add_paragraph()
        p.add_run('Date range: ')
        run_start = p.add_run(f'{s:%B %d, %Y}')
        run_start.bold = True
        p.add_run(' → ')
        run_end = p.add_run(f'{e:%B %d, %Y}')
        run_end.bold = True
        
        
        doc.add_paragraph('1. Go to https://login.qgenda.com/')
        
        # Helper to add each report block
        def add_report(title, steps):
            doc.add_heading(title, level=2)
            for step in steps:
                # If the step contains dates, bold them
                if 'Enter Start Date:' in step:
                    p = doc.add_paragraph(style='List Bullet')
                    prefix, dates = step.split(':', 1)
                    p.add_run(prefix + ': ')
                    # split the two dates on " and End Date:"
                    start_part, end_part = dates.strip().split(' and End Date:')
                    r1 = p.add_run(start_part.strip())
                    r1.bold = True
                    p.add_run(' and End Date: ')
                    r2 = p.add_run(end_part.strip())
                    r2.bold = True
                else:
                    doc.add_paragraph(step, style='List Bullet')
        
        add_report(
            'Report 1 – Penn State Health Hershey Medical Center - Academic General Pediatrics',
            [
                'Click Penn State Health Hershey Medical Center - Academic General Pediatrics → Schedule → Reports',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Short Name',
                'Click Run Report'
            ]
        )
        add_report(
            "Report 2 – Penn State Health Children's Hospital – Hospitalists",
            [
                'Click Penn State Health Children\'s Hospital → Schedule → Reports',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )
        add_report(
            'Report 3 – Department of Pediatrics (Admin - Adolescent Med)',
            [
                'Click Department of Pediatrics → Schedule → Reports',
                'Select Admin - Adolescent Med in top-right corner',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )
        add_report(
            'Report 4 – Department of Pediatrics (Complex Care)',
            [
                'Click Department of Pediatrics → Schedule → Reports',
                'Select Complex Care in top-right corner',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )


        # Save to bytes and offer download
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(label="📄 Download Instructions (Word)",data=buf.getvalue(),file_name="Qgenda_Report_Instructions.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


elif mode == "Format OPD + Summary":
    # ─── Inputs ────────────────────────────────────────────────────────────────────
    # Required keywords to look for in the content
    #required_keywords = ["academic general pediatrics", "hospitalists", "complex care", "adol med"]
    required_keywords = ["academic general pediatrics", "hospitalists", "complex care"]
    found_keywords = set()
    
    schedule_files = st.file_uploader("1) Upload one or more QGenda calendar Excel(s)",type=["xlsx", "xls"],accept_multiple_files=True)
    
    if schedule_files:
        for file in schedule_files:
            try:
                # Read the first sheet
                df = pd.read_excel(file, sheet_name=0, header=None)
    
                # Flatten all string values to a list of lowercase strings
                cell_values = df.astype(str).apply(lambda x: x.str.lower()).values.flatten().tolist()
    
                # Check if any keyword is found in cell values
                for keyword in required_keywords:
                    if any(keyword in val for val in cell_values):
                        found_keywords.add(keyword)
    
            except Exception as e:
                st.error(f"Error reading {file.name}: {e}")
    
        # Identify missing calendars
        missing_keywords = [k for k in required_keywords if k not in found_keywords]
    
        if missing_keywords:
            st.warning(f"Missing required calendar(s): {', '.join(missing_keywords)}. Please upload all required calendars.")
        else:
            st.success("All required calendars uploaded and verified by content.")
    
    student_file = st.file_uploader("2) Upload Redcap Rotation list CSV (must have a 'legal_name' and 'start_date' column)",type=["csv"])
    
    #record_id = st.text_input("3) Enter the REDCap record_id for this batch", "")
    
    record_id = "peds_clerkship"
    
    # ─── Guard ─────────────────────────────────────────────────────────────────────
    if not schedule_files or not student_file or not record_id:
        st.info("Please upload schedule Excel(s), student CSV")
        st.stop()
    
    # ─── Prep: Date regex & maps ───────────────────────────────────────────────────
    
    date_pat = re.compile(r'^[A-Za-z]+ \d{1,2}, \d{4}$')
    base_map = {
        "hope drive am continuity":    "hd_am_",
        "hope drive pm continuity":    "hd_pm_",
        
        "hope drive am acute precept": "hd_am_acute_",
        "hope drive pm acute precept": "hd_pm_acute_",
    
        "hope drive weekend acute 1": "hd_wknd_acute_1_", # Changed prefix
        "hope drive weekend acute 2": "hd_wknd_acute_2_", # Changed prefix
    
        "hope drive weekend continuity": "hd_wknd_am_",
        
        "etown am continuity":         "etown_am_",
        "etown pm continuity":         "etown_pm_",
        
        "nyes rd am continuity":       "nyes_am_",
        "nyes rd pm continuity":       "nyes_pm_",
        
        "nursery weekday 8a-6p":       ["nursery_am_", "nursery_pm_"],
        
        "rounder 1 7a-7p":             ["ward_a_am_","ward_a_pm_"],
        "rounder 2 7a-7p":             ["ward_a_am_","ward_a_pm_"],
        "rounder 3 7a-7p":             ["ward_a_am_","ward_a_pm_"],
    
        "hope drive clinic am":        "complex_am_",
        "hope drive clinic pm":        "complex_pm_",
        
        "briarcrest clinic am":       "adol_med_am_",
        "briarcrest clinic pm":       "adol_med_pm_",

        "lancaster am":       "lancaster_am_",
        "lancaster pm":       "lancaster_pm_",
    
    }
    
    # Which groups need at least 2 providers?
    min_required = {
        "hope drive am acute precept": 2,
        "hope drive pm acute precept": 2,
        
        "nursery weekday 8a-6p":       2,
        
        "rounder 1 7a-7p":             2,
        "rounder 2 7a-7p":             2,
        "rounder 3 7a-7p":             2,
    }
    
    file_configs = {"HAMPDEN_NURSERY.xlsx": {"title": "HAMPDEN_NURSERY","custom_text": "CUSTOM_PRINT","names": ["Folaranmi, Oluwamayoda","Alur, Pradeep","Nanda, Sharmilarani","HAMPDEN_NURSERY"]},
                    "SJR_HOSP.xlsx": {"title": "SJR_HOSPITALIST","custom_text": "CUSTOM_PRINT","names": ["Spangola, Haley","Gubitosi, Terry","SJR_1","SJR_2"]},
                    "AAC.xlsx": {"title": "AAC","custom_text": "CUSTOM_PRINT","names": ["Vaishnavi Harding","Abimbola Ajayi","Shilu Joshi","Desiree Webb","Amy Zisa","Abdullah Sakarcan","Anna Karasik","AAC_1","AAC_2","AAC_3",]},
                    "LANCASTER_CMG.xlsx": {"title": "LANCASTER_CMG","custom_text": "CUSTOM_PRINT","names": ["Ashleigh Sobotka","Susannah Christman"]},
                    "MAHOUSSI_AHOLOUKPE.xlsx": {"title": "MAHOUSSI_AHOLOUKPE","custom_text": "CUSTOM_PRINT","names": ["Mahoussi Aholoukpe"]},
                    #"REPLACE.xlsx": {"title": "REPLACE","custom_text": "CUSTOM_PRINT","names": ["ReplaceFirstName ReplaceLastName"]},
                   }
    
    # ─── HERE: generate sheet‐specific custom_print entries for the configss...  ────────────────────
    for cfg in file_configs.values():
        sheet = cfg["title"]              # e.g. "HAMPDEN_NURSERY"
        key   = sheet.lower() + "_print"  # e.g. "hampden_nursery_print"
        prefix = f"{cfg['custom_text'].lower()}_{sheet.lower()}_"
        base_map[key] = prefix
        
    # ─── 1. Aggregate schedule assignments by date ────────────────────────────────
    assignments_by_date = {}
    for file in schedule_files:
        df = pd.read_excel(file, header=None, dtype=str)
    
        # find all date cells
        date_positions = []
        for r in range(df.shape[0]):
            for c in range(df.shape[1]):
                val = str(df.iat[r,c]).replace("\xa0"," ").strip()
                if date_pat.match(val):
                    try:
                        d = pd.to_datetime(val).date()
                        date_positions.append((d,r,c))
                    except:
                        pass
    
        # dedupe to the topmost row per date
        unique = {}
        for d,r,c in date_positions:
            if d not in unique or r < unique[d][0]:
                unique[d] = (r,c)
    
        # before the loop, define:
        day_names = {"monday","tuesday","wednesday","thursday","friday","saturday","sunday"}
        
        # collect providers under each date
        for d, (row0,col0) in unique.items():
            grp = assignments_by_date.setdefault(d, {des:[] for des in base_map})
            
            for r in range(row0+1, df.shape[0]):
                raw = str(df.iat[r, col0]).replace("\xa0", " ").strip()
                # stop if we hit a blank row
                if raw == "":
                    break
                # stop if we hit another date header
                if date_pat.match(raw):
                    break
    
                desc = raw.lower()
                prov = str(df.iat[r, col0+1]).strip()
                if desc in grp and prov:
                    grp[desc].append(prov)
    
    # ─── Provider filter UI ──────────────────────────────────────────────────────
    all_providers = sorted({
        p.strip()
        for day in assignments_by_date.values()
        for provs in day.values()
        for p in provs
        if isinstance(p, str) and p.strip()
    })
    
    # Multiselect persists in session; start empty by design
    if "provider_filter" not in st.session_state:
        st.session_state["provider_filter"] = []
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("Select All Providers", key="prov_select_all"):
            st.session_state["provider_filter"] = all_providers
    with col2:
        if st.button("Clear Providers", key="prov_clear_all"):
            st.session_state["provider_filter"] = []
    with col3:
        # Switch to actually apply the filter. Off = treat as 'All'
        apply_provider_filter = st.checkbox(
            "Apply provider filter",
            value=False,
            key="prov_apply_filter",
            help="When OFF, everyone is included even if the multiselect is blank."
        )
    
    allowed_providers = st.multiselect(
        "Limit providers included in OPD",
        options=all_providers,
        key="provider_filter",
        help="Only selected providers will be written when 'Apply provider filter' is ON.",
    )
    
    # Effective allow-list:
    effective_allowed = (
        set(allowed_providers) if (apply_provider_filter and allowed_providers) else set(all_providers)
    )

    # ─── 2. Read student list and prepare s1, s2, … ───────────────────────────────
    students_df = pd.read_csv(student_file, dtype=str)
    legal_names = students_df["legal_name"].dropna().tolist()
    
    # ─── 3. Build the single REDCap row ───────────────────────────────────────────
    redcap_row = {"record_id": record_id}
    sorted_dates = sorted(assignments_by_date.keys())
    
    for idx, date in enumerate(sorted_dates, start=1):
        redcap_row[f"hd_day_date{idx}"] = date
        suffix = f"d{idx}_"
    
        # build day‑specific prefixes
        des_map = {
            des: ([p + suffix for p in prefs] if isinstance(prefs, list)
                  else [prefs + suffix])
            for des, prefs in base_map.items()
        }
    
        # 3a) schedule providers (respect provider filter)
        for des, provs in assignments_by_date[date].items():
            # Do not mutate the original list
            filtered = [p for p in provs if p in effective_allowed]
        
            # If the group has a minimum requirement, pad by repeating the first allowed provider
            req = min_required.get(des, len(filtered))
            if filtered and len(filtered) < req:
                filtered = filtered + [filtered[0]] * (req - len(filtered))
        
            # If nothing allowed and no minimum → skip write
            if not filtered:
                continue
        
            if des.startswith("rounder"):
                # rounder N 7a-7p → slot math
                # NOTE: req here is the number of providers per team (usually 2)
                team_idx = int(des.split()[1]) - 1  # 0-based team index
                for i, name in enumerate(filtered, start=1):
                    slot = team_idx * req + i  # team1→1..req, team2→req+1..2*req, etc.
                    for prefix in des_map[des]:
                        redcap_row[f"{prefix}{i if prefix.endswith('_am_') or prefix.endswith('_pm_') else slot}"] = name
                        # ^ If your rounder prefixes are lists like ["ward_a_am_","ward_a_pm_"],
                        #   they'll be in des_map[des] already; the index logic above preserves slots.
            else:
                for i, name in enumerate(filtered, start=1):
                    for prefix in des_map[des]:
                        redcap_row[f"{prefix}{i}"] = name

        # 3b) custom_print names — once per date, using the SAME suffix
        for fname, cfg in file_configs.items():
            sheet = cfg["title"]          
            key   = sheet.lower() + "_print"
            prefix = base_map[key]       # e.g. "custom_print_hampden_nursery_"
            
            for i, person in enumerate(cfg["names"], start=1):
                # note the suffix goes BEFORE the slot index
                redcap_row[f"{prefix}{suffix}{i}"] = person
                
    # append student slots s1,s2,...
    for i,name in enumerate(legal_names, start=1):
        redcap_row[f"s{i}"] = name
    
    # ─── 4. Display & slice out dates/am/acute and students ─────────────────────
    out_df = pd.DataFrame([redcap_row])
    
    # 1) shuffle
    students = legal_names.copy()
    random.shuffle(students)
    
    # 2) define slot sequence
    slot_seq = [1, 3, 5, 2, 4, 6]
    
    # 3) assign
    ward_a_assignment = {}
    
    for idx, student in enumerate(students):
        slot_group = idx // 4                  # every 4 students move to next slot
        slot       = slot_seq[slot_group % len(slot_seq)]
        week_idx   = idx % 4                   # 0→week1,1→week2,2→week3,3→week4
    
        ward_a_assignment[student] = week_idx
        
        # for their week, each Mon–Fri (days 1–5 + 7*week_idx)
        for day in range(1, 6):
            day_num = day + 7 * week_idx
            for shift in ("am", "pm"):
                key  = f"ward_a_{shift}_d{day_num}_{slot}"
                orig = redcap_row.get(key, "")
                redcap_row[key] = f"{orig} ~ {student}" if orig else f"~ {student}"
                
    # ─── track who’s already grabbed a nursery slot ─────────────────────────────
    nursery_assigned = set()
    
    # ─── HAMPDEN_NURSERY: max 1 student for week1 and 1 for week3, into slot _4 ##FOCUSES ON SLOT 4!!! ─────
    for week_idx in (0, 2):  # 0→week1, 2→week3
        pool = [
            s for s in legal_names
            if s not in nursery_assigned
            and ward_a_assignment.get(s, -1) != week_idx
        ]
        if not pool:
            continue
        student = random.choice(pool)
        nursery_assigned.add(student)    # ← mark them as “used”!
    
        for day in range(1, 6):
            d   = day + 7 * week_idx
            key = f"custom_print_hampden_nursery_d{d}_4"        
            orig = redcap_row.get(key, "")
            redcap_row[key] = f"{orig} ~ {student}" if orig else f"~ {student}"
    
    # ─── 2) SJR_HOSPITALIST (max 2 students, any weeks ≠ their Ward A week) ─────
    for week_idx in range(4):  # 0→wk1,1→wk2,2→wk3,3→wk4
        # build pool excluding Hampden and anyone on Ward A that week
        pool = [
            s for s in legal_names
            if s not in nursery_assigned
            and ward_a_assignment.get(s, -1) != week_idx
        ]
        random.shuffle(pool)
        # assign up to two students: first to slot 3, next to slot 4
        for slot_idx in (3, 4):
            if not pool:
                break
            student = pool.pop()
            nursery_assigned.add(student)
            # Mon–Fri of this week
            for day in range(1, 6):
                d   = day + 7 * week_idx
                key = f"custom_print_sjr_hospitalist_d{d}_{slot_idx}"
                orig = redcap_row.get(key, "")
                redcap_row[key] = f"{orig} ~ {student}" if orig else f"~ {student}"
    
    
    # ─── 3) PSHCH_NURSERY (everyone else, up to 8 slots: slot1 weeks1–4, then slot2 wks1–4) ─────────
    leftovers = [s for s in legal_names if s not in nursery_assigned]
    # build (week_idx, slot) in the desired order
    psch_slots = [(wk,1) for wk in range(4)] + [(wk,2) for wk in range(4)]
    for student in leftovers:
        for wk, slot in psch_slots:
            # skip if conflicts with Ward A week
            if ward_a_assignment.get(student, -1) == wk:
                continue
            # build key once (AM & PM) to test existence and avoid duping
            key_am = f"nursery_am_d{day}_ {slot}"
            # assign across Mon–Fri
            for day in range(1, 6):
                d = day + wk * 7
                for prefix in ("nursery_am_","nursery_pm_"):
                    key  = f"{prefix}d{d}_{slot}"
                    orig = redcap_row.get(key, "")
                    redcap_row[key] = f"{orig} ~ {student}" if orig else f"~ {student}"
            # remove this slot so no one else uses it
            psch_slots.remove((wk,slot))
            nursery_assigned.add(student)
            break
        # if no slot left, the student remains unassigned in PSHCH_NURSERY
    
    # format date columns
    for c in out_df.columns:
        if c.startswith("hd_day_date"):
            out_df[c] = pd.to_datetime(out_df[c]).dt.strftime("%m-%d-%Y")
    
    
    out_df = pd.DataFrame([redcap_row])
    csv_full = out_df.to_csv(index=False).encode("utf-8")
    
    def generate_opd_workbook(full_df: pd.DataFrame) -> bytes:
        import io
        import xlsxwriter
    
        output = io.BytesIO()
        workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    
        # ─── Formats ─────────────────────────────────────────────────────────────────
        format1     = workbook.add_format({'font_size':18,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#FEFFCC','border':1})
        format4     = workbook.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#8ccf6f','border':1})
        format4a    = workbook.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#9fc5e8','border':1})
        format5     = workbook.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#FEFFCC','border':1})
        format5a    = workbook.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#d0e9ff','border':1})
        format11    = workbook.add_format({'font_size':18,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#FEFFCC','border':1})
        formate     = workbook.add_format({'font_size':12,'bold':0,'align':'center','valign':'vcenter','font_color':'white','border':0})
        format3     = workbook.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#FFC7CE','border':1})
        format2     = workbook.add_format({'bg_color':'black'})
        format_date = workbook.add_format({'num_format':'m/d/yyyy','font_size':12,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#FFC7CE','border':1})
        format_label= workbook.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter','font_color':'black','bg_color':'#FFC7CE','border':1})
        merge_format= workbook.add_format({'bold':1,'align':'center','valign':'vcenter','text_wrap':True,'font_color':'red','bg_color':'#FEFFCC','border':1})
    
        # ─── Worksheets ─────────────────────────────────────────────────────────────
        worksheet_names = ['HOPE_DRIVE','ETOWN','NYES','LANCASTER','LANCASTER_CMG','COMPLEX','WARD A','PSHCH_NURSERY','HAMPDEN_NURSERY','SJR_HOSP','AAC','AHOLOUKPE','ADOLMED']
        
        sheets = {name: workbook.add_worksheet(name) for name in worksheet_names}
    
        # ─── Site headers ────────────────────────────────────────────────────────────
        site_list = ['Hope Drive','Elizabethtown','Nyes Road','Lancaster','Lancaster CMG','Complex Care','WARD A','PSHCH NURSERY','HAMPDEN NURSERY','SJR HOSPITALIST','AAC','AHOLOUKPE','ADOLMED']
        
        for ws, site in zip(sheets.values(), site_list):
            ws.write(0, 0, 'Site:', format1)
            ws.write(0, 1, site,   format1)
    
        # ─── HOPE_DRIVE specific ────────────────────────────────────────────────────
        hd = sheets['HOPE_DRIVE']
        for cr in ['A8:H15','A32:H39','A56:H63','A80:H87']:
            hd.conditional_format(cr, {'type':'cell','criteria':'>=','value':0,'format':format1})
        for cr in ['A18:H25','A42:H49','A66:H73','A90:H97']:
            hd.conditional_format(cr, {'type':'cell','criteria':'>=','value':0,'format':format5a})
        for cr in ['A6:H6','A7:H7','A30:H30','A31:H31','A54:H54','A55:H55','A78:H78','A79:H79']:
            hd.conditional_format(cr, {'type':'cell','criteria':'>=','value':0,'format':format4})
        for cr in ['A16:H16','A17:H17','A40:H40','A41:H41','A64:H64','A65:H65','A88:H88','A89:H89']:
            hd.conditional_format(cr, {'type':'cell','criteria':'>=','value':0,'format':format4a})

        # how many acute vs continuity rows per block
        ACUTE_COUNT      = 2
        CONTINUITY_COUNT = 8
        BLOCK_SIZE       = ACUTE_COUNT + CONTINUITY_COUNT  # should be 10
        AM_COUNT         = BLOCK_SIZE
        PM_COUNT         = BLOCK_SIZE
        
        # e.g. [6,16,30,40,54,64,78,88]
        BLOCK_STARTS = [6, 30, 54, 78]
        
        for start in BLOCK_STARTS:
            zero_row = start - 1
            # — AM half of the block —
            for i in range(AM_COUNT):
                # first 2 → ACUTES, rest → Continuity
                if i < ACUTE_COUNT:
                    label = 'AM - ACUTES'
                else:
                    label = 'AM - Continuity'
                hd.write(zero_row + i, 0, label, format5a)
            
            # — PM half of the block —
            for i in range(PM_COUNT):
                if i < ACUTE_COUNT:
                    label = 'PM - ACUTES'
                else:
                    label = 'PM - Continuity'
                hd.write(zero_row + AM_COUNT + i, 0, label, format5a)
            
        # ─── GENERIC SHEETS ─────────────────────────────────────────────────────────
        others       = [ws for name, ws in sheets.items() if name != 'HOPE_DRIVE']
        AM_COUNT     = 10
        PM_COUNT     = 10
        BLOCK_STARTS = [6, 30, 54, 78]
    
        for ws in others:
            # 1) conditional formats
            for cr in ['A6:H15','A30:H39','A54:H63','A78:H87']:
                ws.conditional_format(cr, {
                    'type':'cell','criteria':'>=','value':0,'format':format1
                })
            for cr in ['A16:H25','A40:H49','A64:H73','A88:H97']:
                ws.conditional_format(cr, {
                    'type':'cell','criteria':'>=','value':0,'format':format5a
                })
            for cr, fmt in [
                ('B6:H6',   format4), ('B16:H16', format4a),
                ('B30:H30', format4), ('B40:H40', format4a),
                ('B54:H54', format4), ('B64:H64', format4a),
                ('B78:H78', format4), ('B88:H88', format4a)
            ]:
                ws.conditional_format(cr, {
                    'type':'cell','criteria':'>=','value':0,'format':fmt
                })
    
            # 2) Write exactly 10 AM then 10 PM in column A
            for start in BLOCK_STARTS:
                zero_row = start - 1
                for i in range(AM_COUNT):
                    ws.write(zero_row + i, 0, 'AM', format5a)
                for i in range(PM_COUNT):
                    ws.write(zero_row + AM_COUNT + i, 0, 'PM', format5a)
    
        # ─── Universal formatting & dates ────────────────────────────────────────────
        date_cols = [f"hd_day_date{i}" for i in range(1,29)]
        dates     = pd.to_datetime(full_df[date_cols].iloc[0]).tolist()
        weeks     = [dates[i*7:(i+1)*7] for i in range(4)]
        days      = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
    
        for ws in workbook.worksheets():
            ws.set_zoom(80)
            ws.set_column('A:A', 10)
            ws.set_column('B:H', 65)
            ws.set_row(0, 37.25)
    
            for idx, start in enumerate([2,26,50,74]):
                        # day names
                        for c, d in enumerate(days):
                            ws.write(start, 1+c, d, format3)
                        # dates
                        for c, val in enumerate(weeks[idx]):
                            ws.write(start+1, 1+c, val, format_date)
                            
                        # padding formula bars
                        ws.write_formula(f'A{start}',   '""', format_label)
                        
                        ws.conditional_format(
                            f'A{start+3}:H{start+3}',
                            {'type':'cell','criteria':'>=','value':0,'format':format_label}
                        )
    
    
            # black bars every 24 rows
            step = 24
            for row in range(2, 98, step):
                ws.merge_range(f'A{row}:H{row}', ' ', format2)
    
            # merge CRTS message on every sheet
            text1 = (
                'Students are to alert their preceptors when they have a Clinical '
                'Reasoning Teaching Session (CRTS).  Please allow the students to '
                'leave approximately 15 minutes prior to the start of their session '
                'so they can be prepared to actively participate.  - Thank you!'
            )
            ws.merge_range('C1:F1', text1, merge_format)
            ws.write('G1', '', merge_format)
            ws.write('H1', '', merge_format)
    
            #PAINT Empty White Spaces
            ws.write('A3', '', format_date)
            ws.write('A4', '', format_date)
            ws.write('A27', '', format_date)
            ws.write('A28', '', format_date)
            ws.write('A51', '', format_date)
            ws.write('A52', '', format_date)
            ws.write('A75', '', format_date)
            ws.write('A76', '', format_date)
    
    
        workbook.close()
        output.seek(0)
        return output.read()
    
    excel_bytes = generate_opd_workbook(out_df)
    #st.download_button(label="⬇️ Download OPD.xlsx",data=excel_bytes,file_name="OPD.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    
    import pandas as pd
    import io
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment # <--- NEW IMPORT
    
    # --- MODIFIED update_excel_from_csv function to work with bytes ---
    def update_excel_from_csv(excel_template_bytes: bytes, csv_data_bytes: bytes, mappings: list) -> bytes | None:
        """
        Updates an Excel file (from bytes) with values from a CSV (from bytes)
        based on provided mappings, and returns the updated Excel as bytes.
    
        Args:
            excel_template_bytes (bytes): The bytes content of the Excel file to be updated.
            csv_data_bytes (bytes): The bytes content of the CSV file containing the data.
            mappings (list of dict): A list of dictionaries, where each dictionary
                                     defines a mapping:
                                     {'csv_column': 'name_of_csv_column',
                                      'excel_sheet': 'Sheet Name',
                                      'excel_cell': 'Cell Address (e.g., B8)'}
        Returns:
            bytes: The bytes content of the updated Excel workbook, or None if an error occurs.
        """
        try:
            # Load the CSV data from bytes using BytesIO
            df_csv = pd.read_csv(io.BytesIO(csv_data_bytes))
    
            if df_csv.empty:
                # Assuming st is available in the Streamlit environment
                # If running outside Streamlit, you might use print() or logging
                # st.error("Error: The CSV data is empty. Cannot update Excel.")
                return None
    
            # Load the Excel workbook from bytes using BytesIO
            wb = load_workbook(io.BytesIO(excel_template_bytes))
    
            # Iterate through the mappings and update the Excel file
            for mapping in mappings:
                csv_column = mapping['csv_column']
                excel_sheet_name = mapping['excel_sheet']
                excel_cell = mapping['excel_cell']
    
                if csv_column not in df_csv.columns:
                    # st.warning(f"Warning: CSV column '{csv_column}' not found in CSV data. Skipping this mapping.")
                    continue
    
                # Get the value from the first row of the specified CSV column
                value_to_transfer = df_csv.loc[0, csv_column]
    
                if excel_sheet_name not in wb.sheetnames:
                    # st.warning(f"Warning: Excel sheet '{excel_sheet_name}' not found in the Excel template. Skipping this mapping.")
                    continue
    
                ws = wb[excel_sheet_name]
    
                # --- APPLY THE REQUESTED FORMATTING HERE ---
                # 1. Convert to string and append " ~ "
                # This ensures that even if value_to_transfer is a number, it can be concatenated
                #formatted_value = str(value_to_transfer) + ' ~ '

                orig = str(value_to_transfer).strip()
                # if it already contains a student‐delimiter, don’t add another
                if ' ~ ' in orig:
                    formatted_value = orig
                else:
                    formatted_value = orig + ' ~ '
    
                # 2. Write the formatted value to the cell
                ws[excel_cell] = formatted_value
    
                # 3. Set the alignment for the cell using openpyxl's Alignment
                cell = ws[excel_cell] # Get the cell object
                cell.alignment = Alignment(horizontal='center', vertical='center') # Set horizontal and vertical centering
    
                # st.info(f"Successfully wrote '{formatted_value}' from CSV column '{csv_column}' to '{excel_sheet_name}'!{excel_cell}")
    
            # Save the modified Excel workbook to a BytesIO object
            output_excel_bytes_io = io.BytesIO()
            wb.save(output_excel_bytes_io)
            output_excel_bytes_io.seek(0) # Rewind the buffer to the beginning
    
            return output_excel_bytes_io.getvalue()
    
        except Exception as e:
            # st.error(f"An error occurred during Excel update: {e}")
            return None
            
    # --- Configuration for update_excel_from_csv (your mappings) ---
    data_mappings        = []
    excel_column_letters = ['B','C','D','E','F','G','H']
    num_weeks            = 4
    
    # HOPE_DRIVE acute + continuity row offsets
    hd_row_defs = {
        'AM': {'acute_start': 6,  'cont_start': 8},
        'PM': {'acute_start': 16, 'cont_start': 18},
    }
    
    # Other sheets only need continuity (rows 6–13 for AM, 16–23 for PM)
    cont_row_defs = {
        'AM':  6,
        'PM': 16,
    }
    
    # your prefix map
    base_map = {
        "hope drive am continuity":      "hd_am_",
        "hope drive pm continuity":      "hd_pm_",
        "hope drive am acute precept":   "hd_am_acute_",
        "hope drive pm acute precept":   "hd_pm_acute_",
        "hope drive weekend acute 1":    "hd_wknd_acute_1_",
        "hope drive weekend acute 2":    "hd_wknd_acute_2_",
        "hope drive weekend continuity": "hd_wknd_am_",
        
        "etown am continuity":           "etown_am_",
        "etown pm continuity":           "etown_pm_",
        
        "nyes rd am continuity":         "nyes_am_",
        "nyes rd pm continuity":         "nyes_pm_",
        
        "nursery weekday 8a-6p":         ["nursery_am_","nursery_pm_"],
        
        "rounder 1 7a-7p":               ["ward_a_am_","ward_a_pm_"],
        "rounder 2 7a-7p":               ["ward_a_am_","ward_a_pm_"],
        "rounder 3 7a-7p":               ["ward_a_am_","ward_a_pm_"],
        
        "hope drive clinic am":          "complex_am_",
        "hope drive clinic pm":          "complex_pm_",
        
        "briarcrest clinic am":          "adol_med_am_",
        "briarcrest clinic pm":          "adol_med_pm_",

        "lancaster am":          "lancaster_am_",
        "lancaster pm":          "lancaster_pm_",
    
        'hampden_nursery_print':    'custom_print_hampden_nursery_',
        'sjr_hospitalist_print':    'custom_print_sjr_hospitalist_',
        'aac_print':                'custom_print_aac_',
        'lancaster_cmg_print':      'custom_print_lancaster_cmg_',
    
        'mahoussi_aholoukpe_print': 'custom_print_mahoussi_aholoukpe_',
        
    }
    
    # which keys from base_map for each sheet
    sheet_map = {
        'ETOWN':           ('etown am continuity','etown pm continuity'),
        'NYES':            ('nyes rd am continuity','nyes rd pm continuity'),
        'LANCASTER':            ('lancaster am','lancaster pm'),
        'LANCASTER_CMG':        ('lancaster_cmg_print',),
        
        'COMPLEX':         ('hope drive clinic am','hope drive clinic pm'),
        'WARD A':             ('rounder 1 7a-7p','rounder 2 7a-7p','rounder 3 7a-7p'),
        'PSHCH_NURSERY':    ("nursery weekday 8a-6p","nursery weekday 8a-6p"),
        
        'HAMPDEN_NURSERY': ('hampden_nursery_print',),
        'SJR_HOSP':        ('sjr_hospitalist_print',),
        'AAC':             ('aac_print',),
        'AHOLOUKPE':        ('mahoussi_aholoukpe_print',),
        
        'ADOLMED':             ('briarcrest clinic am','briarcrest clinic pm'),
    }
    
    worksheet_names = ['HOPE_DRIVE','ETOWN','NYES','LANCASTER', 'LANCASTER_CMG', 'COMPLEX','WARD A','PSHCH_NURSERY','HAMPDEN_NURSERY','SJR_HOSP','AAC','AHOLOUKPE','ADOLMED']
    
    for ws in worksheet_names:
        # ─── HOPE_DRIVE ───────────────────────────────────────────
        if ws == 'HOPE_DRIVE':
                    # ─── HOPE_DRIVE: exact same 4‑week AM/PM acute+cont logic ───
            for week_idx in range(1, num_weeks + 1):
                week_base  = (week_idx - 1) * 24
                day_offset = (week_idx - 1) * 7
    
                for day_idx, col in enumerate(excel_column_letters, start=1):
                    is_weekday = day_idx <= 5
                    day_num    = day_idx + day_offset
    
                    # AM acute + continuity
                    if is_weekday:
                        # acute (_1–2)
                        for prov in range(1, 3):
                            row = week_base + hd_row_defs['AM']['acute_start'] + (prov - 1)
                            data_mappings.append({
                                'csv_column':  f'hd_am_acute_d{day_num}_{prov}',
                                'excel_sheet': 'HOPE_DRIVE',
                                'excel_cell':  f'{col}{row}',
                            })
                        # continuity (_1–8)
                        for prov in range(1, 9):
                            row = week_base + hd_row_defs['AM']['cont_start'] + (prov - 1)
                            data_mappings.append({
                                'csv_column':  f'hd_am_d{day_num}_{prov}',
                                'excel_sheet': 'HOPE_DRIVE',
                                'excel_cell':  f'{col}{row}',
                            })
                    else:
                        # weekend acute 1 & 2
                        for acute_type in (1, 2):
                            row = week_base + hd_row_defs['AM']['acute_start'] + (acute_type - 1)
                            data_mappings.append({
                                'csv_column':  f'hd_wknd_acute_{acute_type}_d{day_num}_1',
                                'excel_sheet': 'HOPE_DRIVE',
                                'excel_cell':  f'{col}{row}',
                            })
                        # weekend continuity
                        for prov in range(1, 9):
                            row = week_base + hd_row_defs['AM']['cont_start'] + (prov - 1)
                            data_mappings.append({
                                'csv_column':  f'hd_wknd_am_d{day_num}_{prov}',
                                'excel_sheet': 'HOPE_DRIVE',
                                'excel_cell':  f'{col}{row}',
                            })
    
                    # PM acute + continuity
                    if is_weekday:
                        for prov in range(1, 3):
                            row = week_base + hd_row_defs['PM']['acute_start'] + (prov - 1)
                            data_mappings.append({
                                'csv_column':  f'hd_pm_acute_d{day_num}_{prov}',
                                'excel_sheet': 'HOPE_DRIVE',
                                'excel_cell':  f'{col}{row}',
                            })
                        for prov in range(1, 9):
                            row = week_base + hd_row_defs['PM']['cont_start'] + (prov - 1)
                            data_mappings.append({
                                'csv_column':  f'hd_pm_d{day_num}_{prov}',
                                'excel_sheet': 'HOPE_DRIVE',
                                'excel_cell':  f'{col}{row}',
                            })
                    else:
                        for acute_type in (1, 2):
                            row = week_base + hd_row_defs['PM']['acute_start'] + (acute_type - 1)
                            data_mappings.append({
                                'csv_column':  f'hd_wknd_pm_acute_{acute_type}_d{day_num}_1',
                                'excel_sheet': 'HOPE_DRIVE',
                                'excel_cell':  f'{col}{row}',
                            })
                        for prov in range(1, 9):
                            row = week_base + hd_row_defs['PM']['cont_start'] + (prov - 1)
                            data_mappings.append({
                                'csv_column':  f'hd_wknd_pm_d{day_num}_{prov}',
                                'excel_sheet': 'HOPE_DRIVE',
                                'excel_cell':  f'{col}{row}',
                            })
            # done with HOPE_DRIVE
            continue
    
    
        # ─── W_A (rounders) ───────────────────────────────────────
        if ws == 'W_A':
            mapping_keys = sheet_map[ws]  # ('rounder 1…','rounder 2…','rounder 3…')
            for week_idx in range(1, num_weeks+1):
                week_base  = (week_idx - 1) * 24
                day_offset = (week_idx - 1) * 7
    
                for day_idx, col in enumerate(excel_column_letters, start=1):
                    day_num = day_idx + day_offset
    
                    # AM block → rows 6–…
                    row = week_base + cont_row_defs['AM']
                    for team_idx, key in enumerate(mapping_keys):
                        am_pref = base_map[key][0]  # e.g. "ward_a_am_"
                        provs   = assignments_by_date[date][key]
                        req     = min_required.get(key, len(provs))
                        # pad to exactly 2 providers
                        while len(provs) < req:
                            provs.append(provs[0])
                        offset = team_idx * req
                        for i, name in enumerate(provs, start=1):
                            slot = offset + i     # team1→1,2; team2→3,4; team3→5,6
                            data_mappings.append({
                                'csv_column': f"{am_pref}d{day_num}_{slot}",
                                'excel_sheet': ws,
                                'excel_cell': f"{col}{row}",
                            })
                            row += 1
    
                    # PM block → rows 16–…
                    row = week_base + cont_row_defs['PM']
                    for team_idx, key in enumerate(mapping_keys):
                        pm_pref = base_map[key][1]  # e.g. "ward_a_pm_"
                        provs   = assignments_by_date[date][key]
                        req     = min_required.get(key, len(provs))
                        while len(provs) < req:
                            provs.append(provs[0])
                        offset = team_idx * req
                        for i, name in enumerate(provs, start=1):
                            slot = offset + i
                            data_mappings.append({
                                'csv_column': f"{pm_pref}d{day_num}_{slot}",
                                'excel_sheet': ws,
                                'excel_cell': f"{col}{row}",
                            })
                            row += 1
    
            continue  # skip the generic logic below
    
        # ─── ALL OTHER SHEETS ──────────────────────────────────────
        mapping_keys = sheet_map.get(ws, ())
        if not mapping_keys:
            continue
    
        for key in mapping_keys:
            val = base_map[key]
        
            # Decide which side(s) this key applies to
            am_prefix = pm_prefix = None
            if isinstance(val, list):
                am_prefix, pm_prefix = val
            else:
                k = key.lower()
                if " am " in k:
                    am_prefix = val
                elif " pm " in k:
                    pm_prefix = val
                else:
                    # keys that don't encode AM/PM (rare) write to both
                    am_prefix = pm_prefix = val
    
            for week_idx in range(1, num_weeks + 1):
                week_base  = (week_idx - 1) * 24
                day_offset = (week_idx - 1) * 7
    
                for day_idx, col in enumerate(excel_column_letters, start=1):
                    day_num = day_idx + day_offset
    
                    # AM continuity (_1–10)
                    for prov in range(1, 11):
                        row = week_base + cont_row_defs['AM'] + (prov - 1)
                        data_mappings.append({
                            'csv_column': f"{am_prefix}d{day_num}_{prov}",
                            'excel_sheet': ws,
                            'excel_cell': f"{col}{row}",
                        })
    
                    # PM continuity (_1-10)
                    for prov in range(1, 11):
                        row = week_base + cont_row_defs['PM'] + (prov - 1)
                        data_mappings.append({
                            'csv_column': f"{pm_prefix}d{day_num}_{prov}",
                            'excel_sheet': ws,
                            'excel_cell': f"{col}{row}",
                        })
    
        def hide_blank_rows_all_sheets(excel_bytes: bytes):
            """
            Hide rows where col A starts with AM/PM and ALL of B..H are empty.
            Works for every sheet. Preserves row indices so conditional formats stay aligned.
        
            Returns (new_excel_bytes, per_sheet_hidden_counts, total_hidden)
            """
            import io, re
            from openpyxl import load_workbook
        
            def _empty(v):
                return v is None or (isinstance(v, str) and v.strip() == "")
        
            wb = load_workbook(io.BytesIO(excel_bytes))
            per_sheet = {}
            total = 0
        
            for ws in wb.worksheets:
                hidden = 0
                for r in range(1, ws.max_row + 1):
                    a1 = ws.cell(row=r, column=1).value
                    if not (isinstance(a1, str) and re.match(r"^\s*(AM|PM)\b", a1, re.IGNORECASE)):
                        continue
                    if all(_empty(ws.cell(row=r, column=c).value) for c in range(2, 9)):
                        ws.row_dimensions[r].hidden = True
                        ws.row_dimensions[r].height = 0
                        hidden += 1
                per_sheet[ws.title] = hidden
                total += hidden
        
            out = io.BytesIO()
            wb.save(out)
            out.seek(0)
            return out.getvalue(), per_sheet, total
                
    # --- Main execution flow for generating and then updating the workbook ---
    st.subheader("Generate & Update OPD.xlsx + Summary")
    
    if st.button("Generate OPD File For Sarah to Load Students"):
        # 1) Generate the initial OPD workbook
        excel_template_bytes = generate_opd_workbook(out_df)
        if not excel_template_bytes:
            st.error("Failed to generate OPD template.")
            st.stop()
    
        # 2) Update it with your CSV data
        updated_excel_bytes = update_excel_from_csv(excel_template_bytes, csv_full, data_mappings)
        if not updated_excel_bytes:
            st.error("Failed to update OPD.xlsx with data.")
            st.stop()

        cleaned_bytes, hidden_map, hidden_total = hide_blank_rows_all_sheets(updated_excel_bytes)
        st.success("✅ OPD.xlsx updated successfully!")
    
        # 3) Build your summary DataFrame (reuse your df_summary logic)
        summary = []
        for student in legal_names:
            entry = {"Student": student}
            for w in range(4):
                days = [d + w*7 for d in range(1,6)]
                assigns = []
                # Ward A
                ward_found = False
                for shift in ("am","pm"):
                    for slot in range(1,7):
                        for d in days:
                            key = f"ward_a_{shift}_d{d}_{slot}"
                            if student in redcap_row.get(key,""):
                                assigns.append("Ward A")
                                ward_found = True
                                break
                        if ward_found: break
                    if ward_found: break
                # Hampden
                if not ward_found:
                    for d in days:
                        key = f"custom_print_hampden_nursery_d{d}_4"
                        if student in redcap_row.get(key,""):
                            assigns.append("Hampden")
                            break
                # SJR
                sjr_found = False
                for slot in (3,4):
                    for d in days:
                        key = f"custom_print_sjr_hospitalist_d{d}_{slot}"
                        if student in redcap_row.get(key,""):
                            assigns.append("SJR")
                            sjr_found = True
                            break
                    if sjr_found: break
                # PSHCH
                pshch_found = False
                for slot in (1,2):
                    for d in days:
                        for pref in ("nursery_am_","nursery_pm_"):
                            key = f"{pref}d{d}_{slot}"
                            if student in redcap_row.get(key,""):
                                assigns.append("PSHCH")
                                pshch_found = True
                                break
                        if pshch_found: break
                    if pshch_found: break
    
                entry[f"Week {w+1}"] = ", ".join(assigns) or ""
            summary.append(entry)
        df_summary = pd.DataFrame(summary)
    
        # 4) Build a Word doc with the summary table
        doc = Document()
        # make landscape
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        
        doc.add_heading("Assignment Summary by Week", level=1)
        
        cols  = df_summary.columns.tolist()
        table = doc.add_table(rows=1, cols=len(cols), style="Table Grid")
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr_cells[i].text = c
        
        for _, row in df_summary.iterrows():
            row_cells = table.add_row().cells
            for i, c in enumerate(cols):
                row_cells[i].text = str(row[c])
        
        # **Save** into bytes
        word_io = io.BytesIO()
        doc.save(word_io)
        word_io.seek(0)
        word_bytes = word_io.read()
        
        # 5) Package into a ZIP
        zip_io = io.BytesIO()
        with zipfile.ZipFile(zip_io, "w") as z:
            z.writestr("Updated_OPD.xlsx", cleaned_bytes)
            z.writestr("Assignment_Summary.docx", word_bytes)
        zip_io.seek(0)
        
        # 6) Single download
        st.download_button(label="⬇️ Download OPD.xlsx + Summary (zip)",data=zip_io.read(),file_name="Batch_Output.zip",mime="application/zip")

elif mode == "Create Student Schedule":
    st.subheader("Create Student Schedule")
    def save_to_session(filename, fileobj, namespace="uploaded_files"):
        st.session_state.setdefault(namespace, {})[filename] = fileobj
        
    # ───────── Helper to load & stash uploads ─────────
    def load_workbook_df(label, types, key):
        upload = st.file_uploader(label, type=types, key=key)
        if not upload:
            st.info(f"Please upload {label}.")
            return None
        # stash the raw upload under a known session key
        st.session_state[f"{key}_file"] = upload
        try:
            if upload.name.lower().endswith(".csv"):
                return pd.read_csv(upload)
            else:
                return pd.read_excel(upload)
        except Exception as e:
            st.error(f"Error loading {upload.name}: {e}")
            return None

    def create_ms_schedule_template(students, dates):
        buf = io.BytesIO()
        wb = xlsxwriter.Workbook(buf, {'in_memory': True})
    
        # — Formats —
        f1 = wb.add_format({'font_size':14,'bold':1,'align':'center','valign':'vcenter',
                            'font_color':'black','text_wrap':True,'bg_color':'#FEFFCC','border':1})
        f2 = wb.add_format({'font_size':10,'bold':1,'align':'center','valign':'vcenter',
                            'font_color':'yellow','bg_color':'black','border':1,'text_wrap':True})
        f3 = wb.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter',
                            'font_color':'black','bg_color':'#FFC7CE','border':1})
        f4 = wb.add_format({'num_format':'mm/dd/yyyy','font_size':12,'bold':1,'align':'center',
                            'valign':'vcenter','font_color':'black','bg_color':'#F4F6F7','border':1})
        f5 = wb.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter',
                            'font_color':'black','bg_color':'#F4F6F7','border':1})
        f6 = wb.add_format({'bg_color':'black','border':1})
        f7 = wb.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter',
                            'font_color':'black','bg_color':'#90EE90','border':1})
        f8 = wb.add_format({'font_size':12,'bold':1,'align':'center','valign':'vcenter',
                            'font_color':'black','bg_color':'#89CFF0','border':1})
    
        days = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
        start_rows = [2, 10, 18, 26]
        weeks = ['Week 1','Week 2','Week 3','Week 4']
        due_texts = [
            'Quiz 1 Due',
            'Quiz 2, Pediatric Documentation #1, 1 Clinical Encounter Log Due',
            'Quiz 3 Due',
            'Quiz 4, Pediatric Documentation #2, Social Drivers of Health Assessment Form, Developmental Assessment of Pediatric Patient Form, All Clinical Encounter Logs are Due!'
        ]
    
        for name in students:
            title = name[:31].replace('/','-').replace('\\','-')
            ws = wb.add_worksheet(title)
            ws.set_zoom(70)
    
            # Header
            ws.merge_range('A1:A2','Student Name:', f1)
            ws.merge_range('B1:B2',      title,      f1)
            note = ("*Note* Asynchronous time is for coursework only. During this time period, "
                    "we expect students to do coursework, be available for any additional educational "
                    "activities, and any extra clinical time that may be available. If the student is not "
                    "available during this time period and has not made an absence request, the student "
                    "will be cited for unprofessionalism and will risk failing the course.")
            ws.merge_range('C1:H2', note, f2)
    
            # Column widths & row height
            ws.set_column('A:A', 20)
            ws.set_column('B:B', 30)
            ws.set_column('C:G', 40)
            ws.set_column('H:H',155)
            ws.set_row(0, 37.25)
    
            # Days headers and dates
            date_idx = 0
            for block, row in enumerate(start_rows):
                # 1) write the days on row `row`, cols B–H
                for col_offset, day in enumerate(days, start=1):
                    ws.write(row, col_offset, day, f3)
        
            # 2) write the dates directly beneath in B–H (row+1)
                for col_offset in range(7):
                    if date_idx < len(dates):
                        # note the +1 here instead of +2
                        ws.write(row+1, col_offset+1, dates[date_idx], f4)
                        date_idx += 1
    
            # Week labels
            for i, week in enumerate(weeks):
                row = 4 + (i * 8)
                ws.write(f'A{row}', week, f3)
    
            # AM / PM labels
            for i in range(4):
                ws.write(f'A{6 + i*8}', 'AM', f3)
                ws.write(f'A{7 + i*8}', 'PM', f3)
    
            # Fill AM/PM blocks with Asynchronous Time (cols C–J)
            for block in range(4):
                am_row = 5 + block*8
                pm_row = 6 + block*8
                for col in range(1, 8):
                    ws.write(am_row, col, "Asynchronous Time", f5)
                    ws.write(pm_row, col, "Asynchronous Time", f5)
    
            # Separators
            for sep in [10, 18, 26, 34]:
                ws.merge_range(f'A{sep}:H{sep}', '', f6)
    
            # Green filler rows
            for filler in [8, 16, 24, 32]:
                for col in range(8):
                    ws.write(filler, col, ' ', f7)
    
            # Assignment‑due rows
            for i, base in enumerate([8, 16, 24, 32]):
                ws.write(f'A{base}', 'ASSIGNMENT DUE:', f8)
                for col in range(1, 8):
                    if col == 5:
                        ws.write(base-1, col, 'Ask for Feedback!', f8)
                    elif col == 7:
                        ws.write(base-1, col, due_texts[i], f8)
                    else:
                        ws.write(base-1, col, ' ', f8)
    
        wb.close()
        buf.seek(0)
        return buf
    
    def assign_preceptors_all_weeks_am(opd_file, ms_file):
        """
        For each OPD sheet:
          1) Collect all rows where col A starts with "AM".
          2) Cluster those rows into contiguous week‐blocks.
          3) Map each week_block i to MS_Schedule row [6,14,22,30][i].
          4) Copy any "Preceptor ~ Student" in OPD cols B–H within that block
             into the student's sheet at that row.
        Returns an in‑memory BytesIO of the populated MS_Schedule.
        """
        # Open workbooks
        opd_wb = load_workbook(opd_file, data_only=True)
        ms_wb  = load_workbook(ms_file)
    
        # Fixed target rows in MS template for Week1–4 AM
        target_ms_rows = [6, 14, 22, 30]
    
        for site in opd_wb.sheetnames:
            ws_opd = opd_wb[site]
    
            # 1) Find all AM marker rows in col A
            am_rows = [
                cell.row
                for cell in ws_opd['A']
                if isinstance(cell.value, str) and re.match(r"^\s*AM\b", cell.value, re.IGNORECASE)
            ]
    
            if not am_rows:
                continue
    
            # 2) Cluster contiguous AM rows into blocks
            am_rows.sort()
            blocks = []
            current = [am_rows[0]]
            for r in am_rows[1:]:
                if r == current[-1] + 1:
                    current.append(r)
                else:
                    blocks.append(current)
                    current = [r]
            blocks.append(current)  # last block
    
            # 3) Process up to 4 week‐blocks
            for week_idx, block in enumerate(blocks[:4]):
                ms_row = target_ms_rows[week_idx]
    
                # 4) Copy assignments in B–H for every row in this block
                for col in range(2, 9):  # B=2 … H=8
                    for r in block:
                        val = ws_opd.cell(row=r, column=col).value
                        if not val or "~" not in str(val):
                            continue
                        pre, student = [s.strip() for s in str(val).split("~", 1)]
                        if student not in ms_wb.sheetnames:
                            continue
                        ws_ms = ms_wb[student]
                        ws_ms.cell(row=ms_row, column=col).value = f"{pre} - [{site}]"
    
        # Save back to a BytesIO buffer
        out = io.BytesIO()
        ms_wb.save(out)
        out.seek(0)
        return out

    def assign_preceptors_all_weeks_pm(opd_file, ms_file):
        """
        For each OPD sheet:
          1) Collect all rows where col A starts with "PM".
          2) Cluster those rows into contiguous week‑blocks.
          3) Map each week_block i to MS_Schedule row [7,15,23,31][i].
          4) Copy any "Preceptor ~ Student" in OPD cols B–H within that block
             into the student's sheet at that row.
        Returns an in‑memory BytesIO of the populated MS_Schedule.
        """
        # Open workbooks
        opd_wb = load_workbook(opd_file, data_only=True)
        ms_wb  = load_workbook(ms_file)
    
        # Fixed target rows in MS template for Week1–4 PM
        target_ms_rows = [7, 15, 23, 31]
    
        for site in opd_wb.sheetnames:
            ws_opd = opd_wb[site]
    
            # 1) Find all PM marker rows in col A
            pm_rows = [
                cell.row
                for cell in ws_opd['A']
                if isinstance(cell.value, str) and re.match(r"^\s*PM\b", cell.value, re.IGNORECASE)
            ]
    
            if not pm_rows:
                continue
    
            # 2) Cluster contiguous PM rows into blocks
            pm_rows.sort()
            blocks = []
            current = [pm_rows[0]]
            for r in pm_rows[1:]:
                if r == current[-1] + 1:
                    current.append(r)
                else:
                    blocks.append(current)
                    current = [r]
            blocks.append(current)
    
            # 3) Process up to 4 week‐blocks
            for week_idx, block in enumerate(blocks[:4]):
                ms_row = target_ms_rows[week_idx]
    
                # 4) Copy assignments in B–H for every row in this block
                for col in range(2, 9):  # B=2 … H=8
                    for r in block:
                        val = ws_opd.cell(row=r, column=col).value
                        if not val or "~" not in str(val):
                            continue
                        pre, student = [s.strip() for s in str(val).split("~", 1)]
                        if student not in ms_wb.sheetnames:
                            continue
                        ws_ms = ms_wb[student]
                        ws_ms.cell(row=ms_row, column=col).value = f"{pre} - [{site}]"

        # Save back to a BytesIO buffer
        out = io.BytesIO()
        ms_wb.save(out)
        out.seek(0)
        return out
    
    def detect_shift_conflicts(opd_file):
        """
        Scans both AM and PM shifts, week 1–4, day Mon–Sun, and flags any student
        who appears more than once in the same shift/day/week.
        Only ignores entries where there is no text after the '~'.
        """
        wb = load_workbook(opd_file, data_only=True)
        days      = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
        am_marker = re.compile(r"^\s*AM\b", re.IGNORECASE)
        pm_marker = re.compile(r"^\s*PM\b", re.IGNORECASE)
        conflicts = []
    
        def find_blocks(ws, marker_re):
            rows = [c.row for c in ws['A']
                    if isinstance(c.value, str) and marker_re.match(c.value)]
            rows.sort()
            blocks, curr = [], []
            for r in rows:
                if not curr or r == curr[-1] + 1:
                    curr.append(r)
                else:
                    blocks.append(curr)
                    curr = [r]
            if curr:
                blocks.append(curr)
            return blocks[:4]
    
        # derive AM/PM blocks from first sheet
        tpl      = wb[wb.sheetnames[0]]
        am_blocks = find_blocks(tpl, am_marker)
        pm_blocks = find_blocks(tpl, pm_marker)
    
        for shift, blocks in (("AM", am_blocks), ("PM", pm_blocks)):
            for week_idx, block_rows in enumerate(blocks, start=1):
                for day_idx, day_name in enumerate(days):
                    col = 2 + day_idx
                    locs = defaultdict(list)
    
                    # collect all assignments in this shift
                    for sheet in wb.sheetnames:
                        ws = wb[sheet]
                        for r in block_rows:
                            raw = ws.cell(row=r, column=col).value
                            text = str(raw or "")
                            if "~" not in text:
                                continue
                            pre, student = [s.strip() for s in text.split("~",1)]
                            if not student:
                                continue
                            coord = ws.cell(row=r, column=col).coordinate
                            locs[student].append((sheet, coord))
    
                    # flag duplicates
                    for student, occ in locs.items():
                        if len(occ) > 1:
                            conflicts.append({
                                "student":     student,
                                "week":        week_idx,
                                "day":         day_name,
                                "shift":       shift,
                                "occurrences": occ
                            })
    
        return conflicts


    # ───────── Load OPD & Rotation Schedule ─────────
    df_opd = load_workbook_df("Upload OPD.xlsx file", ["xlsx"], key="opd_main")
    df_rot = load_workbook_df("Upload Rotation Schedule (.xlsx or .csv)", ["xlsx", "csv"], key="rot_main")

        # ───────── Check for duplicates ─────────
    if df_opd is not None:
        conflicts = detect_shift_conflicts(st.session_state["opd_main_file"])
        if conflicts:
            for c in conflicts:
                occ_str = "; ".join(f"{sheet}@{coord}" for sheet, coord in c["occurrences"])
                st.warning(
                    f"⚠️ Week {c['week']} {c['day']} {c['shift']}: "
                    f"{c['student']} double‑booked ({occ_str})"
                )
        else:
            st.success("No AM/PM shift conflicts detected.")


    # ───────── Build, Assign & Download ─────────
    if df_opd is not None and df_rot is not None:
        # compute dates
        df_rot["start_date"] = pd.to_datetime(df_rot["start_date"])
        monday = df_rot["start_date"].min() - pd.Timedelta(days=df_rot["start_date"].min().weekday())
        dates  = pd.date_range(start=monday, periods=28, freq="D").tolist()

        # students from OPD
        students = df_rot["legal_name"].dropna().unique().tolist()

        if st.button("Create & Download Fully‑Populated MS_Schedule"):
            # 1) Build the blank 4‑week calendar
            blank_buf = create_ms_schedule_template(students, dates)
        
            # 2) Populate AM slots from OPD
            am_buf = assign_preceptors_all_weeks_am(opd_file = st.session_state["opd_main_file"],ms_file  = blank_buf)
        
            # 3) Populate PM slots on top of the AM‑populated file
            full_buf = assign_preceptors_all_weeks_pm(opd_file = st.session_state["opd_main_file"],ms_file  = am_buf)
        
            # 4) Offer the final workbook for download
            st.download_button("Download MS_Schedule.xlsx",data = full_buf.getvalue(),file_name = "MS_Schedule.xlsx",mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    else:
        st.info("Please upload both OPD.xlsx and the rotation schedule above to proceed.")

elif mode == "Create Individual Schedules":
    st.subheader("Individual Schedule Creator")

    uploaded = st.file_uploader(
        "Upload the master Excel (.xlsx) with one tab per person",
        type=["xlsx"]
    )

    def copy_sheet_to_new_wb(src_ws):
        """Return a BytesIO of a new .xlsx containing src_ws with formatting."""
        from openpyxl import Workbook
        from io import BytesIO
    
        wb_new = Workbook()
        ws_new = wb_new.active
        ws_new.title = src_ws.title[:31]
    
        # Column widths & visibility
        for col_letter, dim in src_ws.column_dimensions.items():
            if dim.width is not None:
                ws_new.column_dimensions[col_letter].width = dim.width
            ws_new.column_dimensions[col_letter].hidden = dim.hidden
    
        # Row heights & visibility
        for idx, dim in src_ws.row_dimensions.items():
            if dim.height is not None:
                ws_new.row_dimensions[idx].height = dim.height
            ws_new.row_dimensions[idx].hidden = dim.hidden
    
        # Sheet settings (best-effort)
        try:
            ws_new.sheet_format.defaultColWidth = src_ws.sheet_format.defaultColWidth
            ws_new.sheet_format.defaultRowHeight = src_ws.sheet_format.defaultRowHeight
        except Exception:
            pass
        ws_new.freeze_panes = src_ws.freeze_panes
        try:
            ws_new.page_setup.orientation = src_ws.page_setup.orientation
            ws_new.page_setup.fitToWidth = src_ws.page_setup.fitToWidth
            ws_new.page_setup.fitToHeight = src_ws.page_setup.fitToHeight
            ws_new.page_margins = src_ws.page_margins
            ws_new.print_options.horizontalCentered = src_ws.print_options.horizontalCentered
            ws_new.print_options.verticalCentered = src_ws.print_options.verticalCentered
            ws_new.print_area = src_ws.print_area
        except Exception:
            pass

        try:
            # Set zoom to 70%
            ws_new.sheet_view.zoomScale = 70
    
            # Set column widths
            ws_new.column_dimensions["A"].width = 20
            ws_new.column_dimensions["B"].width = 30
            for col in ["C", "D", "E", "F", "G"]:
                ws_new.column_dimensions[col].width = 40
            ws_new.column_dimensions["H"].width = 155
        except Exception:
            pass
    
        # Copy cells: values + (copied) styles
        for row in src_ws.iter_rows():
            for cell in row:
                # Skip non-master cells from merged ranges
                if isinstance(cell, MergedCell):
                    continue
        
                # Get a reliable numeric column index
                col_idx = getattr(cell, "col_idx", None)
                if col_idx is None:
                    col = cell.column  # may be int or letter depending on version
                    col_idx = col if isinstance(col, int) else column_index_from_string(col)
        
                # Create target cell with value (formula preserved if present)
                tgt = ws_new.cell(row=cell.row, column=col_idx, value=cell.value)
        
                # Copy style safely
                if getattr(cell, "has_style", False):
                    try:
                        from copy import copy
                        if cell.font:        tgt.font        = copy(cell.font)
                        if cell.fill:        tgt.fill        = copy(cell.fill)
                        if cell.border:      tgt.border      = copy(cell.border)
                        if cell.alignment:   tgt.alignment   = copy(cell.alignment)
                        if cell.protection:  tgt.protection  = copy(cell.protection)
                        tgt.number_format = cell.number_format
                    except Exception:
                        pass
    
        # Copy merged cell ranges (after values)
        for merged in list(src_ws.merged_cells.ranges):
            try:
                ws_new.merge_cells(str(merged))
            except Exception:
                pass
    
        # Copy data validations (best-effort)
        try:
            if src_ws.data_validations and src_ws.data_validations.dataValidation:
                from openpyxl.worksheet.datavalidation import DataValidation
                for dv in src_ws.data_validations.dataValidation:
                    dv_new = DataValidation(
                        type=dv.type,
                        formula1=dv.formula1,
                        formula2=dv.formula2,
                        allow_blank=dv.allow_blank,
                        operator=dv.operator,
                        showDropDown=dv.showDropDown,
                        showErrorMessage=dv.showErrorMessage,
                        errorTitle=dv.errorTitle,
                        error=dv.error,
                        promptTitle=dv.promptTitle,
                        prompt=dv.prompt
                    )
                    # Copy cell refs
                    for sqref in getattr(dv, "sqref", []):
                        dv_new.add(sqref)
                    ws_new.add_data_validation(dv_new)
        except Exception:
            pass
    
        # Filters
        try:
            ws_new.auto_filter.ref = getattr(src_ws.auto_filter, "ref", None)
        except Exception:
            pass
    
        # Save to buffer
        buf = BytesIO()
        wb_new.save(buf)
        buf.seek(0)
        return buf
        
    if uploaded is not None:
        # Keep formulas/formatting -> data_only=False
        wb = load_workbook(uploaded, data_only=False)
        st.write(f"Found **{len(wb.sheetnames)}** tabs.")

        if st.button("Split tabs and build ZIP"):
            zip_buf = BytesIO()
            with ZipFile(zip_buf, mode="w", compression=ZIP_DEFLATED) as zf:
                used_names = defaultdict(int)

                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]

                    # Skip truly empty sheets (no cells with value)
                    has_any_value = any(cell.value is not None for row in ws.iter_rows() for cell in row)
                    if not has_any_value:
                        continue

                    # Safe file name
                    base = re.sub(r"[^A-Za-z0-9._-]+", "_", sheet_name).strip("_") or "sheet"
                    used_names[base] += 1
                    safe_name = base if used_names[base] == 1 else f"{base}_{used_names[base]}"

                    # Copy this sheet into its own new workbook (preserving formatting)
                    out_buf = copy_sheet_to_new_wb(ws)

                    # Add to ZIP
                    zf.writestr(f"{safe_name}.xlsx", out_buf.getvalue())

            zip_buf.seek(0)
            st.download_button(
                label="Download individual schedules (ZIP)",
                data=zip_buf,
                file_name="individual_schedules_formatted.zip",
                mime="application/zip",
            )

elif mode == "OPD MD PA Conflict Detector":
    import streamlit as st
    import pandas as pd
    from collections import defaultdict

    st.title("OPD MD/PA Double-Booking & Availability")
    st.write(
        "Upload the MD and PA OPD Excel files to scan for double-booked preceptors and to list availability "
        "by site/date/period (including other sites)."
    )

    # -----------------------------
    # Config & Constants
    # -----------------------------
    DAYS = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
    DEFAULT_FOCUS = ['HOPE_DRIVE','NYES','ETOWN','LANCASTER']

    TRUST_ONLY_AM_PM = True          # Only parse rows with Column A starting AM/PM
    REQUIRE_VALID_DATE = True        # Bookings require a valid date anchor

    # -----------------------------
    # Helpers
    # -----------------------------
    @st.cache_data(show_spinner=False)
    def read_sheet_names(file):
        try:
            xl = pd.ExcelFile(file)
            return xl.sheet_names
        except Exception as e:
            st.error(f"Failed to read sheet names: {e}")
            return []

    @st.cache_data(show_spinner=False)
    def load_sheet(file, sheet_name):
        return pd.read_excel(file, sheet_name=sheet_name, header=None)

    def _try_parse_date(x):
        try:
            d = pd.to_datetime(x, errors='coerce')
            if pd.notna(d):
                return d
            return None
        except Exception:
            return None

    def find_week_headers(df: pd.DataFrame):
        """
        Robust week header finder.
        - Find 'Monday' (case-insensitive) in Column B.
        - Within next 10 rows, pick first row where B..H has >=2 parseable dates -> dates row.
        - monday_date is B's parsed date; if missing, infer from any parsed day in that row.
        Returns [(monday_row, date_row, monday_date)]
        """
        col = 1
        s = df.iloc[:, col].astype(str).str.strip().str.lower()
        day_rows = df.index[s.eq('monday')].tolist()

        headers = []
        for dr in day_rows:
            date_r, monday_date = None, None
            for look_ahead in range(1, 11):
                r = dr + look_ahead
                if r >= len(df):
                    break
                parsed_dates = {}
                for i in range(7):
                    c = 1 + i
                    if c >= df.shape[1]:
                        continue
                    parsed = _try_parse_date(df.iat[r, c])
                    if parsed is not None:
                        parsed_dates[i] = parsed
                if len(parsed_dates) >= 2:
                    date_r = r
                    if 0 in parsed_dates:
                        monday_date = parsed_dates[0].date()
                    else:
                        i0 = sorted(parsed_dates.keys())[0]
                        monday_date = (parsed_dates[i0] - pd.Timedelta(days=i0)).date()
                    break
            headers.append((dr, date_r, monday_date))
        return headers

    def row_to_week_monday(row_idx: int, headers):
        prev = [h for h in headers if h[0] <= row_idx]
        if not prev:
            return None
        prev.sort(key=lambda x: x[0])
        return prev[-1][2]

    def detect_am_pm_runs(df: pd.DataFrame, start_row: int = 0):
        """Scan Col A for AM/PM rows and group consecutive runs."""
        runs, current, run_start, prev_idx = [], None, None, None
        for idx in range(start_row, len(df)):
            label = None
            raw = df.iat[idx, 0]
            if isinstance(raw, str):
                ru = raw.strip().upper()
                if ru.startswith('AM'):
                    label = 'AM'
                elif ru.startswith('PM'):
                    label = 'PM'
            if TRUST_ONLY_AM_PM and label is None:
                continue
            if label is None:
                continue
            if current is None:
                current, run_start = label, idx
            elif label != current or (prev_idx is not None and idx != prev_idx + 1):
                runs.append((current, run_start, prev_idx))
                current, run_start = label, idx
            prev_idx = idx
        if current is not None:
            runs.append((current, run_start, prev_idx))
        return runs

    def build_maps_and_roster(df: pd.DataFrame):
        """
        Returns:
          mapping_by_week: {(monday_date, period, day, preceptor) -> {'student','cell','date'}}
                           (only when a real student exists)
          index_by_date:   {(date, period, preceptor) -> {'student','cell','day'}}  # for date-based matching
          roster_week:     {(monday_date, period) -> set(preceptors)}               # week-level roster
          day_roster:      {(monday_date, period, day) -> set(preceptors)}          # day-level roster
          week_dates:      {monday_date -> {day -> date}}
          occupied:        set((monday_date, period, day, preceptor)) even w/o student
          diag_weeks:      diagnostics list
        """
        import re, unicodedata

        def is_placeholder_preceptor(text: str) -> bool:
            if not text: return True
            t = text.strip().upper()
            EXCLUDE_PREFIXES = [
                'CLOSED','CLOSE','BLOCK','VACATION','ADMIN','MEETING','NO CLINIC',
                'CLINIC CANCELLED','CANCELLED','HOLIDAY','OFF','PTO','SICK',
                'NOTE','NOTES','REFERENCE','INFO','FYI','ORIENTATION'
            ]
            return any(t.startswith(pfx) for pfx in EXCLUDE_PREFIXES)

        def _norm(s: str) -> str:
            s = unicodedata.normalize("NFKC", s)
            s = s.replace("\u00A0", " ")                    # NBSP -> space
            s = s.replace("\u2013", "-").replace("\u2014", "-")  # en/em dash -> '-'
            s = re.sub(r"\s+", " ", s.strip())
            return s

        def parse_cell(val: str):
            """Robust split on '~'. RHS counts as student only if alphanumeric & not a placeholder."""
            if not isinstance(val, str): return None, None
            raw = _norm(val)
            if "~" not in raw:
                pre = _norm(raw)
                return (pre if pre else None), None
            pre, rhs = re.split(r"\s*~\s*", raw, maxsplit=1)
            pre = _norm(pre)
            rhs = _norm(rhs)

            BLANK_TOKENS = {"", "nan", "n/a", "na", "-", "--", "—", "none", "null"}
            PLACEHOLDER_HINTS = {"note", "notes", "ref", "reference", "info", "fyi"}

            rhs_l = rhs.lower()
            # treat as empty unless it has at least one letter/number and is not a placeholder
            if (rhs_l in BLANK_TOKENS) or (not re.search(r"[a-z0-9]", rhs_l)) or any(h in rhs_l for h in PLACEHOLDER_HINTS):
                rhs = None
            return (pre if pre else None), rhs

        headers = find_week_headers(df)
        runs = detect_am_pm_runs(df, start_row=0)

        mapping_by_week = {}
        index_by_date = {}
        roster_week = defaultdict(set)
        day_roster = defaultdict(set)
        week_dates = defaultdict(dict)
        occupied = set()
        diag_weeks = []

        # Build date anchors with fallback
        for (day_row, date_row, monday_date) in headers:
            if monday_date is None or date_row is None:
                continue
            inferred_days = []
            for i, day in enumerate(DAYS):
                col_idx = 1 + i  # B..H
                val = df.iat[date_row, col_idx] if col_idx < df.shape[1] else None
                parsed = _try_parse_date(val)
                if parsed is not None:
                    week_dates[monday_date][day] = parsed.date()
                else:
                    # fallback Monday + i days
                    try:
                        fallback = (pd.to_datetime(monday_date) + pd.Timedelta(days=i)).date()
                        week_dates[monday_date][day] = fallback
                        inferred_days.append(day)
                    except Exception:
                        week_dates[monday_date][day] = None
                        inferred_days.append(day)
            diag_weeks.append({
                'monday_date': monday_date,
                'date_row': date_row,
                'inferred_days': inferred_days
            })

        # Parse inside AM/PM runs
        for period, rstart, rend in runs:
            monday_date = row_to_week_monday(rstart, headers)
            if monday_date is None:
                continue
            for col_idx, day in enumerate(DAYS, start=1):
                date_anchor = week_dates[monday_date].get(day)
                for row in range(rstart, rend+1):
                    if col_idx >= df.shape[1]:
                        continue
                    val = df.iat[row, col_idx]
                    if pd.isna(val) or not isinstance(val, str):
                        continue
                    pre, stu = parse_cell(val)
                    if not pre or is_placeholder_preceptor(pre):
                        continue

                    # Present in week & specific day
                    roster_week[(monday_date, period)].add(pre)
                    day_roster[(monday_date, period, day)].add(pre)
                    occupied.add((monday_date, period, day, pre))  # presence marker

                    # Keep booking only if real student + valid date
                    if stu is None:
                        continue
                    if REQUIRE_VALID_DATE and _try_parse_date(date_anchor) is None:
                        continue

                    cell = f"{chr(ord('A')+col_idx)}{row+1}"
                    wk_key = (monday_date, period, day, pre)
                    mapping_by_week.setdefault(wk_key, {'student': stu, 'cell': cell, 'date': date_anchor})
                    # date-based index for cross-file matching
                    dt_key = (pd.to_datetime(date_anchor).date(), period, pre)
                    # prefer first seen student for stability
                    index_by_date.setdefault(dt_key, {'student': stu, 'cell': cell, 'day': day})

        return mapping_by_week, index_by_date, roster_week, day_roster, week_dates, occupied, diag_weeks

    # -----------------------------
    # UI - File uploads
    # -----------------------------
    col1, col2 = st.columns(2)
    with col1:
        md_file = st.file_uploader("Upload MD OPD (xlsx)", type=["xlsx"], key="md")
    with col2:
        pa_file = st.file_uploader("Upload PA OPD (xlsx)", type=["xlsx"], key="pa")

    if md_file and pa_file:
        md_sheets = read_sheet_names(md_file)
        pa_sheets = read_sheet_names(pa_file)

        common_sheets = sorted([s for s in DEFAULT_FOCUS if s in md_sheets and s in pa_sheets])
        selected_sheets = st.multiselect(
            "Sites (tabs) to compare",
            options=sorted(list(set(md_sheets) & set(pa_sheets))),
            default=common_sheets or sorted(list(set(md_sheets) & set(pa_sheets)))
        )

        if not selected_sheets:
            st.warning("No common sheets selected.")
            st.stop()

        # Keep per-site context so we can search "other sites"
        site_ctx = {}

        conflict_rows = []
        diagnostics = []

        for sheet in selected_sheets:
            df_md = load_sheet(md_file, sheet)
            df_pa = load_sheet(pa_file, sheet)

            (md_map_wk, md_idx_date, md_roster_wk, md_day_roster, md_week_dates,
             md_occupied, md_diag) = build_maps_and_roster(df_md)
            (pa_map_wk, pa_idx_date, pa_roster_wk, pa_day_roster, pa_week_dates,
             pa_occupied, pa_diag) = build_maps_and_roster(df_pa)
            diagnostics.append({'site': sheet, 'md': md_diag, 'pa': pa_diag})

            # Save for cross-site availability
            site_ctx[sheet] = dict(
                md_idx_date=md_idx_date,
                pa_idx_date=pa_idx_date,
                md_week_dates=md_week_dates,
                pa_week_dates=pa_week_dates,
                md_day_roster=md_day_roster,
                pa_day_roster=pa_day_roster
            )

            # --------- CONFLICTS by actual date ---------
            md_keys = set(md_idx_date.keys())
            pa_keys = set(pa_idx_date.keys())
            for (date_obj, period, pre) in sorted(md_keys & pa_keys):
                md_entry = md_idx_date[(date_obj, period, pre)]
                pa_entry = pa_idx_date[(date_obj, period, pre)]
                conflict_rows.append({
                    'site': sheet,
                    'date': date_obj,
                    'day': pd.to_datetime(date_obj).strftime('%A'),
                    'period': period,
                    'preceptor': pre,
                    'md_student': md_entry['student'],
                    'pa_student': pa_entry['student']
                })

        # Conflicts dataframe
        conflicts_df = pd.DataFrame(conflict_rows)

        # --------- helpers to build pools ---------
        def pool_for_site_day(site, day_name, period, date_obj):
            """Union of preceptors present in THIS site for (day, period, date)."""
            ctx = site_ctx[site]
            pool = set()
            md_week_dates, pa_week_dates = ctx['md_week_dates'], ctx['pa_week_dates']
            md_day_roster, pa_day_roster = ctx['md_day_roster'], ctx['pa_day_roster']
            # MD
            for m in md_week_dates.keys():
                if md_week_dates[m].get(day_name) == date_obj:
                    pool |= (md_day_roster.get((m, period, day_name)) or set())
            # PA
            for m in pa_week_dates.keys():
                if pa_week_dates[m].get(day_name) == date_obj:
                    pool |= (pa_day_roster.get((m, period, day_name)) or set())
            return pool

        def pool_for_other_sites(current_site, day_name, period, date_obj):
            """Union of preceptors present in ALL OTHER sites for (day, period, date)."""
            pool = set()
            for site in site_ctx.keys():
                if site == current_site:
                    continue
                pool |= pool_for_site_day(site, day_name, period, date_obj)
            return pool

        def count_assigned_any_site(pre, date_obj, period):
            """How many students (MD+PA) does preceptor have across all sites at this date/period?"""
            total = 0
            for s, ctx in site_ctx.items():
                if (date_obj, period, pre) in ctx['md_idx_date']:
                    total += 1
                if (date_obj, period, pre) in ctx['pa_idx_date']:
                    total += 1
            return total

        # --------- AVAILABILITY (same-site & other-sites) for conflict slots ---------
        availability_same_rows = []
        availability_other_rows = []
        suggestions_rows = []

        if not conflicts_df.empty:
            for _, r in conflicts_df.iterrows():
                site   = r['site']
                date_o = r['date']
                day_nm = r['day']         # 'Monday'...'Sunday'
                period = r['period']

                # Pools
                same_pool  = pool_for_site_day(site, day_nm, period, date_o)
                other_pool = pool_for_other_sites(site, day_nm, period, date_o)

                # Build availability function
                def add_pool(pool, dest_list, pool_site_label):
                    for pre in sorted(pool):
                        total_assigned = count_assigned_any_site(pre, date_o, period)
                        is_acute = ("ACUTE" in str(pre).upper())
                        capacity = 2 if is_acute else 1
                        seats_left = max(0, capacity - total_assigned)
                        if seats_left > 0:
                            dest_list.append({
                                'site_of_conflict': site,
                                'candidate_site': pool_site_label,
                                'date': date_o,
                                'day': day_nm,
                                'period': period,
                                'conflict_preceptor': r['preceptor'],
                                'preceptor': pre,
                                'is_acute': is_acute,
                                'current_students': total_assigned,
                                'capacity': capacity,
                                'seats_left': seats_left,
                                'status': 'available'
                            })

                add_pool(same_pool, availability_same_rows, site)
                # For other sites, keep which site each candidate belongs to.
                for other_site in site_ctx.keys():
                    if other_site == site:
                        continue
                    pool = pool_for_site_day(other_site, day_nm, period, date_o)
                    add_pool(pool, availability_other_rows, other_site)

        avail_same_df = pd.DataFrame(availability_same_rows)
        avail_other_df = pd.DataFrame(availability_other_rows)

        # --------- SUGGESTIONS (top-3) prefer same-site, then other-sites ---------
        if not conflicts_df.empty:
            for _, r in conflicts_df.iterrows():
                in_slot_same  = avail_same_df[
                    (avail_same_df['site_of_conflict'] == r['site']) &
                    (avail_same_df['date'] == r['date']) &
                    (avail_same_df['period'] == r['period'])
                ].copy()

                in_slot_other = avail_other_df[
                    (avail_other_df['site_of_conflict'] == r['site']) &
                    (avail_other_df['date'] == r['date']) &
                    (avail_other_df['period'] == r['period'])
                ].copy()

                # Put same preceptor first if eligible (Acute w/ 1 student)
                def order(df):
                    df['_self'] = (df['preceptor'] == r['preceptor'])
                    return df.sort_values(['_self','candidate_site','preceptor'], ascending=[False, True, True]).drop(columns=['_self'])

                ordered = pd.concat([order(in_slot_same), order(in_slot_other)], ignore_index=True)

                if not ordered.empty:
                    for _, a in ordered.head(3).iterrows():
                        label = a['preceptor']
                        if a['preceptor'] == r['preceptor']:
                            label = f"{a['preceptor']} (currently assigned)"
                        suggestions_rows.append({
                            'conflict_site': r['site'],
                            'date': r['date'],
                            'day': r['day'],
                            'period': r['period'],
                            'conflict_preceptor': r['preceptor'],
                            'md_student': r['md_student'],
                            'pa_student': r['pa_student'],
                            'suggested_preceptor': label,
                            'suggested_site': a['candidate_site'],
                            'suggested_is_acute': bool(a['is_acute']),
                            'suggested_current_students': int(a['current_students']),
                            'suggested_capacity': int(a['capacity']),
                            'suggested_seats_left': int(a['seats_left'])
                        })
                else:
                    suggestions_rows.append({
                        'conflict_site': r['site'],
                        'date': r['date'],
                        'day': r['day'],
                        'period': r['period'],
                        'conflict_preceptor': r['preceptor'],
                        'md_student': r['md_student'],
                        'pa_student': r['pa_student'],
                        'suggested_preceptor': '⚠️ No alternative preceptors available',
                        'suggested_site': None,
                        'suggested_is_acute': None,
                        'suggested_current_students': None,
                        'suggested_capacity': None,
                        'suggested_seats_left': None
                    })

        suggestions_df = pd.DataFrame(suggestions_rows)

        # -----------------------------
        # Results UI (conflict-focused)
        # -----------------------------
        st.subheader("Results (conflict-focused)")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.metric("Sites compared", len(selected_sheets))
        with c2:
            st.metric("Double bookings found", 0 if conflicts_df.empty else len(conflicts_df))
        with c3:
            st.metric("Avail. (same-site)", 0 if avail_same_df.empty else len(avail_same_df))
        with c4:
            st.metric("Avail. (other-sites)", 0 if avail_other_df.empty else len(avail_other_df))

        st.markdown("**Double-booked preceptors (MD & PA in same slot)**")
        if conflicts_df.empty:
            st.info("No double-bookings found for the selected sites.")
        else:
            st.dataframe(conflicts_df[['site','date','day','period','preceptor','md_student','pa_student']], use_container_width=True)
            st.download_button(
                label="Download double-bookings CSV",
                data=conflicts_df[['site','date','day','period','preceptor','md_student','pa_student']].to_csv(index=False).encode('utf-8'),
                file_name="opd_double_bookings.csv",
                mime="text/csv"
            )

        # Availability (same site)
        show_same  = st.toggle("Show available preceptors in the SAME site (Acutes can take 2)",value=False, key="tog_same_site")
        if show_same:
            if avail_same_df.empty:
                st.info("No same-site availability for the conflicted slots.")
            else:
                st.markdown("**Available preceptors (same site as conflict)** — Acutes shown if <2 students; others only if unbooked.")
                st.dataframe(
                    avail_same_df[['site_of_conflict','candidate_site','date','day','period',
                                   'conflict_preceptor','preceptor','is_acute',
                                   'current_students','capacity','seats_left','status']],
                    use_container_width=True
                )
                st.download_button(
                    label="Download same-site availability CSV",
                    data=avail_same_df.to_csv(index=False).encode('utf-8'),
                    file_name="opd_availability_same_site.csv",
                    mime="text/csv"
                )

        # Availability (other sites)
        show_other = st.toggle("Show available preceptors in OTHER sites (Acutes can take 2)",value=False, key="tog_other_site")
        if show_other:
            if avail_other_df.empty:
                st.info("No other-site availability for the conflicted slots.")
            else:
                st.markdown("**Available preceptors (other sites)** — same date & AM/PM, different site.")
                st.dataframe(
                    avail_other_df[['site_of_conflict','candidate_site','date','day','period',
                                    'conflict_preceptor','preceptor','is_acute',
                                    'current_students','capacity','seats_left','status']],
                    use_container_width=True
                )
                st.download_button(
                    label="Download other-site availability CSV",
                    data=avail_other_df.to_csv(index=False).encode('utf-8'),
                    file_name="opd_availability_other_sites.csv",
                    mime="text/csv"
                )

        # Suggestions
        show_sugg  = st.toggle("Show suggestions to resolve each conflict (prefers same site, then other sites)",
                       value=False, key="tog_suggestions")
        if show_sugg:
            if suggestions_df.empty:
                st.info("No suggestions available.")
            else:
                st.markdown("**Targeted suggestions** — same site first; if none, suggests from other sites on the same date & AM/PM.")
                st.dataframe(
                    suggestions_df[['conflict_site','date','day','period','conflict_preceptor',
                                    'md_student','pa_student','suggested_preceptor','suggested_site',
                                    'suggested_is_acute','suggested_current_students',
                                    'suggested_capacity','suggested_seats_left']],
                    use_container_width=True
                )
                st.download_button(
                    label="Download suggestions CSV",
                    data=suggestions_df.to_csv(index=False).encode('utf-8'),
                    file_name="opd_targeted_suggestions_cross_site.csv",
                    mime="text/csv"
                )

        # -----------------------------
        # Optional Annotated Downloads Toggle
        # -----------------------------
        show_annotated = st.toggle("Generate annotated OPD files (highlight conflicts in RED)",
                           value=False, key="tog_annotated_downloads")
        
        if show_annotated:
            st.markdown("---")
            st.subheader("Download annotated OPDs (conflicts highlighted in RED)")
            st.caption("Cells are red when the *other* OPD already has that preceptor booked for the same site, date, and AM/PM.")
        
            from io import BytesIO
            from openpyxl import load_workbook
            from openpyxl.styles import Font, Color, Border, Side, PatternFill
            from openpyxl.comments import Comment
            
            def _annot_make_copy(uploaded_file, other_idx_by_site: dict, selected_sheets: list) -> bytes:
                """
                Annotate: red font (if visible), THICK RED BORDER, and a small note so conflicts
                are obvious even when Conditional Formatting overrides font color.
                """
                raw = uploaded_file.getvalue()
                wb = load_workbook(BytesIO(raw))
            
                # --- helpers matching your main parser ---
                import re, unicodedata, pandas as pd
                def _norm(s: str) -> str:
                    s = unicodedata.normalize("NFKC", s).replace("\u00A0", " ")
                    s = s.replace("\u2013", "-").replace("\u2014", "-")
                    return re.sub(r"\s+", " ", s.strip())
            
                def _parse_cell(val: str):
                    if not isinstance(val, str): return None, None
                    raw = _norm(val)
                    if "~" not in raw:
                        pre = _norm(raw); return (pre if pre else None), None
                    pre, rhs = re.split(r"\s*~\s*", raw, maxsplit=1)
                    pre = _norm(pre); rhs = _norm(rhs)
                    if rhs.lower() in {"", "nan", "n/a", "na", "-", "--", "—", "none", "null"}:
                        rhs = None
                    return (pre if pre else None), rhs
            
                def _is_placeholder_preceptor(text: str) -> bool:
                    if not text: return True
                    t = text.strip().upper()
                    return any(t.startswith(pfx) for pfx in [
                        'CLOSED','CLOSE','BLOCK','VACATION','ADMIN','MEETING','NO CLINIC',
                        'CLINIC CANCELLED','CANCELLED','HOLIDAY','OFF','PTO','SICK',
                        'NOTE','NOTES','REFERENCE','INFO','FYI','ORIENTATION'
                    ])
            
                # opaque ARGB
                OPAQUE_RED = Color(rgb="FFFF0000")
                RED_SIDE   = Side(style="thick", color="FFFF0000")
                RED_BORDER = Border(left=RED_SIDE, right=RED_SIDE, top=RED_SIDE, bottom=RED_SIDE)
            
                for sheet in selected_sheets:
                    if sheet not in wb.sheetnames:
                        continue
                    ws = wb[sheet]
            
                    # rebuild date map from the uploaded file (aligns weeks/days to dates)
                    df = load_sheet(uploaded_file, sheet)
                    headers = find_week_headers(df)
                    runs = detect_am_pm_runs(df, start_row=0)
            
                    week_dates = {}
                    for (_day_row, date_row, monday_date) in headers:
                        if monday_date is None or date_row is None:
                            continue
                        week_dates.setdefault(monday_date, {})
                        for i, day in enumerate(DAYS):
                            c = 1 + i  # B..H
                            if c >= df.shape[1]: continue
                            parsed = pd.to_datetime(df.iat[date_row, c], errors='coerce')
                            if pd.notna(parsed):
                                week_dates[monday_date][day] = parsed.date()
            
                    other_idx = other_idx_by_site.get(sheet, {})  # keys: (date, period, preceptor)
            
                    for period, rstart, rend in runs:
                        monday_date = row_to_week_monday(rstart, headers)
                        if monday_date is None:
                            continue
                        for c_idx, day in enumerate(DAYS, start=1):  # B..H
                            dt = week_dates.get(monday_date, {}).get(day)
                            if dt is None:
                                continue
                            for row in range(rstart, rend+1):
                                if c_idx >= df.shape[1]: continue
                                val = df.iat[row, c_idx]
                                if pd.isna(val) or not isinstance(val, str): continue
                                pre, _stu = _parse_cell(val)
                                if not pre or _is_placeholder_preceptor(pre): continue
            
                                if (dt, period, pre) in other_idx:
                                    addr = f"{chr(ord('A')+c_idx)}{row+1}"
                                    cell = ws[addr]
            
                                    # Try to ensure red font (CF may still override)
                                    f = cell.font or Font()
                                    try:
                                        cell.font = f.copy(color="FFFF0000")
                                    except Exception:
                                        cell.font = Font(
                                            name=f.name, size=f.size or 11, bold=f.bold,
                                            italic=f.italic, underline=f.underline, color=OPAQUE_RED
                                        )
            
                                    # Add thick red border (highly visible even with CF)
                                    cell.border = RED_BORDER
            
                                    # Add a small note/comment (red triangle)
                                    if cell.comment is None:
                                        txt = f"Booked in other OPD\n{sheet} — {day} {dt} — {period}\nPreceptor: {pre}"
                                        try:
                                            cell.comment = Comment(txt, "MD↔PA conflict")
                                        except Exception:
                                            pass
            
                out = BytesIO()
                wb.save(out)
                out.seek(0)
                return out.getvalue()

        
            # Compare across files
            md_compare_against_pa = {s: site_ctx[s]['pa_idx_date'] for s in site_ctx}
            pa_compare_against_md = {s: site_ctx[s]['md_idx_date'] for s in site_ctx}
            
            col_md, col_pa = st.columns(2)
            with col_md:
                md_bytes = _annot_make_copy(md_file, md_compare_against_pa, selected_sheets)
                st.download_button("⬇️ MD annotated (RED = booked in PA)", md_bytes,
                                   "md_opd_annotated.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            with col_pa:
                pa_bytes = _annot_make_copy(pa_file, pa_compare_against_md, selected_sheets)
                st.download_button("⬇️ PA annotated (RED = booked in MD)", pa_bytes,
                                   "pa_opd_annotated.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    else:
        st.info("Upload both the MD and PA OPD files to begin.")

elif mode == "Shift Availability Tracker":
    st.title("Shift Availability Tracker")
        
    opd_file = st.file_uploader("Upload md_opd.xlsx", type=["xlsx"])
       
    if opd_file:
            excel = pd.ExcelFile(opd_file)
            shift_summary = []
    
            for sheet in excel.sheet_names:
                df = pd.read_excel(excel, sheet_name=sheet, header=None)
                for i, row in df.iterrows():
                    shift_label = str(row[0]).strip().upper()
                    if sheet == "HOPE_DRIVE":
                        valid_shifts = ["AM - ACUTES", "AM - CONTINUITY", "PM - ACUTES", "PM - CONTINUITY"]
                    else:
                        valid_shifts = ["AM", "PM"]
                    if shift_label in valid_shifts:
                        dates = df.iloc[2, 1:]  # date row
                        for col, date in enumerate(dates, start=1):
                            cell_values = df.iloc[i, col]
                            if isinstance(cell_values, str) and "~" in cell_values:
                                shift_summary.append({
                                    "Site": sheet,
                                    "Date": pd.to_datetime(date, errors="coerce"),
                                    "Shift": shift_label,
                                    "Preceptor": cell_values.strip()
                                })
    
            summary_df = pd.DataFrame(shift_summary)
            summary_count = summary_df.groupby(["Site", "Date", "Shift"]).size().reset_index(name="Preceptor_Count")
    
            st.write("### Summary Table")
            st.dataframe(summary_count)
    
            am_pm_toggle = st.toggle("Show AM Only")
            if am_pm_toggle:
                st.dataframe(summary_count[summary_count["Shift"].str.startswith("AM")])
            else:
                st.dataframe(summary_count[summary_count["Shift"].str.startswith("PM")])
