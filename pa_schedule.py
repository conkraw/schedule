import io
import streamlit as st
import pandas as pd
import numpy as np 
import re
import xlsxwriter
import random
from openpyxl import load_workbook # Ensure load_workbook is imported
import io, zipfile
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from datetime import timedelta
from xlsxwriter import Workbook as Workbook
from collections import defaultdict
from datetime import datetime, timedelta
from collections import Counter
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED
import re
from collections import defaultdict
from copy import copy
from openpyxl import load_workbook, Workbook
from openpyxl.utils import get_column_letter
from openpyxl.cell.cell import MergedCell
from openpyxl.utils import column_index_from_string
import math

st.set_page_config(page_title="PSUCOM PA SCHEDULE CREATOR", layout="wide")
st.title("PSUCOM PA SCHEDULE CREATOR")

# ─── Sidebar mode selector ─────────────────────────────────────────────────────
#mode = st.sidebar.radio("What do you want to do?",("Instructions", "Format OPD + Summary (4-sheet, 5-week)", "Create Student Schedule","OPD Check","Create Individual Schedules"))
mode = st.sidebar.radio("What do you want to do?",("PA OPD Creator", "Create Student Schedule","Create Individual Schedules"))
# ─── Sidebar mode selector ─────────────────────────────────────────────────────

if mode == "OPD Check":
    DAYS = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
    
    def detect_am_pm_blocks(df):
        """
        Scan down column A and group runs whose cell text begins with 'AM' or 'PM',
        ignoring everything else.
        Returns a list of (label, start_row, end_row) tuples.
        """
        runs = []
        current_label = None
        run_start = None
        prev_row = None
    
        # Excel row 6 is df index 5
        for idx in range(5, len(df)):
            raw = df.iat[idx, 0]
            if isinstance(raw, str):
                ru = raw.upper()
                if ru.startswith("AM"):
                    label = "AM"
                elif ru.startswith("PM"):
                    label = "PM"
                else:
                    continue
            else:
                continue
    
            row = idx + 1  # convert back to 1‑based Excel row
            if current_label is None:
                # start new run
                current_label = label
                run_start = row
            elif label != current_label or row != prev_row + 1:
                # close out previous run
                runs.append((current_label, run_start, prev_row))
                current_label = label
                run_start = row
    
            prev_row = row
    
        # finish last run
        if current_label is not None:
            runs.append((current_label, run_start, prev_row))
    
        return runs
    
    st.title("OPD Preceptor GI‑al Check")
    
    baseline_file = st.file_uploader(
        "1) Upload Latest Updated OPD", 
        type=["xlsx"], key="baseline"
    )
    assigned_file = st.file_uploader(
        "2) Upload OPD with Student Assignments", 
        type=["xlsx"], key="assigned"
    )
    
    if baseline_file and assigned_file:
        SHEETS = [
            'HOPE_DRIVE','ETOWN','NYES','COMPLEX','WARD A',
            'PSHCH_NURSERY','HAMPDEN_NURSERY','SJR_HOSP','AAC',
            'AHOLOUKPE','ADOLMED'
        ]
    
        # Read all relevant sheets at once
        base_sheets = pd.read_excel(baseline_file, sheet_name=SHEETS, header=None)
        assn_sheets = pd.read_excel(assigned_file, sheet_name=SHEETS, header=None)

        # 1) Remove student suffix from every baseline sheet
        for sheet_name, df in base_sheets.items():
            for col in df.columns[1:]:
                df[col] = (
                    df[col]
                      .where(df[col].notna(), np.nan)      # keep NaN
                      .astype(str)
                      .str.partition('~')[0]               # take text before "~"
                      .replace({'nan': np.nan})            # restore NaN
                )
            base_sheets[sheet_name] = df
    
        results = {}
        
        for sheet in SHEETS:
            df_base = base_sheets[sheet]
            df_assn = assn_sheets[sheet]
        
            runs = detect_am_pm_blocks(df_base)
            week_pairs = [(runs[i], runs[i+1]) for i in range(0, len(runs), 2)]
        
            base_map = {}    # (w,period,day,pre) → (cell, student)
            assn_map = {}
        
            for w_idx, ((am_lbl, am_s, am_e), (pm_lbl, pm_s, pm_e)) in enumerate(week_pairs, start=1):
                for period, (lbl, start, end) in [('AM',(am_lbl,am_s,am_e)), ('PM',(pm_lbl,pm_s,pm_e))]:
                    for col, day in enumerate(DAYS, start=1):
                        for row in range(start, end+1):
                            cell = f"{chr(ord('A')+col)}{row}"
        
                            # baseline
                            vb = df_base.iat[row-1, col]
                            if pd.notna(vb) and isinstance(vb, str):
                                parts = str(vb).split('~', 1)
                                pre = parts[0].strip()
                                stu = parts[1].strip() if len(parts)==2 else None
                                base_map.setdefault((w_idx,period,day,pre), (cell, stu))
        
                            # assigned
                            va = df_assn.iat[row-1, col]
                            if pd.notna(va) and isinstance(va, str):
                                parts = str(va).split('~', 1)
                                pre = parts[0].strip()
                                stu = parts[1].strip() if len(parts)==2 else None
                                assn_map.setdefault((w_idx,period,day,pre), (cell, stu))
        
            dropped = []
            added   = []
        
            # drops: in base not in assigned
            for key, (cell, stu) in base_map.items():
                if key not in assn_map:
                    w,p,d,pre = key
                    dropped.append((w,p,d,pre,cell,stu))
        
            # adds: in assigned not in base
            for key, (cell, stu) in assn_map.items():
                if key not in base_map:
                    w,p,d,pre = key
                    added.append((w,p,d,pre,cell,stu))
        
            # sort exactly as before...
            dropped_sorted = sorted(
                dropped,
                key=lambda x: (x[0], {'AM':0,'PM':1}[x[1]], DAYS.index(x[2]), x[4])
            )
            added_sorted = sorted(
                added,
                key=lambda x: (x[0], {'AM':0,'PM':1}[x[1]], DAYS.index(x[2]), x[4])
            )
        
            results[sheet] = {"dropped": dropped_sorted, "added": added_sorted}



        
        #
    
    doc = Document()
    doc.add_heading('Change Report', level=1)
    
    for sheet, change in results.items():
        doc.add_heading(sheet, level=2)
    
        # build week→day map (collect both AM & PM under each day)
        week_map = defaultdict(lambda: defaultdict(lambda: {'dropped': [], 'added': []}))
        for w,p,d,pre,cell,stu in change['dropped']:
            week_map[w][d]['dropped'].append((p, pre, cell, stu))
        for w,p,d,pre,cell,stu in change['added']:
            week_map[w][d]['added'].append((p, pre, cell, stu))
    
        # emit in week order
        for w in sorted(week_map):
            doc.add_heading(f'Week {w}', level=3)
    
            for day in DAYS:
                slot = week_map[w].get(day)
                if not slot or (not slot['dropped'] and not slot['added']):
                    continue
    
                doc.add_heading(day, level=4)
    
                # DROPS
                for p, pre, cell, stu in slot['dropped']:
                    line = f"- Dropped: {pre} — was at {cell}"
                    if stu:
                        line += f"  (Student impacted: {stu})"
                    doc.add_paragraph(line, style='List Bullet')
    
                # ADDS
                for p, pre, cell, stu in slot['added']:
                    line = f"- Added: {pre} — now at {cell}"
                    if stu:
                        line += f"  (Student impacted: {stu})"
                    doc.add_paragraph(line, style='List Bullet')
    
            doc.add_paragraph()  # blank line between weeks




    # Save to in-memory buffer
    word_file = io.BytesIO()
    doc.save(word_file)
    word_file.seek(0)
    
    # Download button
    st.download_button(
        label="📄 Download Word Report",
        data=word_file,
        file_name="change_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


elif mode == "Instructions":
    d = st.text_input('Start date (m/d/yyyy)')
    if d:
        try:
            s = datetime.strptime(d, '%m/%d/%Y')
            e = s + timedelta(days=41)
            st.write(f"{s:%B %d, %Y} → {e:%B %d, %Y}")
    
            st.write('Please go to https://login.qgenda.com/')
    
            # Display on‐screen instructions
            st.markdown(f"""Download four files and create reports based on **{s:%B %d, %Y}** → **{e:%B %d, %Y}**.""")
            st.write("Download instructions here:")


        except ValueError:
            st.error('Invalid format – use m/d/yyyy (e.g. 7/6/2021)')

        # --- Generate a Word document with the same instructions ---
        doc = Document()
        doc.add_heading('Qgenda Report Instructions', level=1)
        doc.styles['Normal'].font.size = Pt(8)
        
        # Bold the date range
        p = doc.add_paragraph()
        p.add_run('Date range: ')
        run_start = p.add_run(f'{s:%B %d, %Y}')
        run_start.bold = True
        p.add_run(' → ')
        run_end = p.add_run(f'{e:%B %d, %Y}')
        run_end.bold = True
        
        
        doc.add_paragraph('1. Go to https://login.qgenda.com/')
        
        # Helper to add each report block
        def add_report(title, steps):
            doc.add_heading(title, level=2)
            for step in steps:
                # If the step contains dates, bold them
                if 'Enter Start Date:' in step:
                    p = doc.add_paragraph(style='List Bullet')
                    prefix, dates = step.split(':', 1)
                    p.add_run(prefix + ': ')
                    # split the two dates on " and End Date:"
                    start_part, end_part = dates.strip().split(' and End Date:')
                    r1 = p.add_run(start_part.strip())
                    r1.bold = True
                    p.add_run(' and End Date: ')
                    r2 = p.add_run(end_part.strip())
                    r2.bold = True
                else:
                    doc.add_paragraph(step, style='List Bullet')
        
        add_report(
            'Report 1 – Penn State Health Hershey Medical Center - Academic General Pediatrics',
            [
                'Click Penn State Health Hershey Medical Center - Academic General Pediatrics → Schedule → Reports',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Short Name',
                'Click Run Report'
            ]
        )
        add_report(
            "Report 2 – Penn State Health Children's Hospital – Hospitalists",
            [
                'Click Penn State Health Children\'s Hospital → Schedule → Reports',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )
        add_report(
            'Report 3 – Department of Pediatrics (Admin - Adolescent Med)',
            [
                'Click Department of Pediatrics → Schedule → Reports',
                'Select Admin - Adolescent Med in top-right corner',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )
        add_report(
            'Report 4 – Department of Pediatrics (Complex Care)',
            [
                'Click Department of Pediatrics → Schedule → Reports',
                'Select Complex Care in top-right corner',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )


        # Save to bytes and offer download
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(label="📄 Download Instructions (Word)",data=buf.getvalue(),file_name="Qgenda_Report_Instructions.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif mode == "PA OPD Creator":
    # ─────────────────────────────────────────────────────────────────────────────
    # Imports
    # ─────────────────────────────────────────────────────────────────────────────
    import io
    import re
    from collections import defaultdict
    import pandas as pd
    import streamlit as st

    # ─────────────────────────────────────────────────────────────────────────────
    # Helpers & Defaults
    # ─────────────────────────────────────────────────────────────────────────────
    AM_PM_RE = re.compile(r"\b(am|pm)\b", re.IGNORECASE)
    date_pat = re.compile(r"^[A-Za-z]+ \d{1,2}, \d{4}$")

    # purely for the "missing keyword" heads-up; adjust as you like
    REQUIRED_KEYWORDS = sorted({
        "academic general pediatrics",
        "complex care",
        "neurology",
        "endo",
        "hemonc",
    })

    def slugify(name: str) -> str:
        return re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")

    def guess_site_bucket(des: str):
        d = des.lower()
        if "hope" in d:
            return "HOPE_DRIVE"
        if "nyes" in d:
            return "NYES"
        # keep Lancaster gen peds (avoid catching Endo Lancaster, etc.)
        if "lancaster" in d and all(x not in d for x in ["endo", "endocrin", "neuro", "hemonc", "onc", "spec"]):
            return "LANCASTER"
        if "etown" in d or "elizabethtown" in d:
            return "ETOWN"
        return "SUBSPECIALTY"

    def guess_slot(des: str):
        m = AM_PM_RE.search(des)
        if not m:
            return "NONE"
        return m.group(1).upper()  # "AM"/"PM"

    # ─────────────────────────────────────────────────────────────────────────────
    # Inputs
    # ─────────────────────────────────────────────────────────────────────────────
    schedule_files = st.file_uploader(
        "1) Upload one or more QGenda calendar Excel(s)",
        type=["xlsx", "xls"],
        accept_multiple_files=True,
    )
    student_file = st.file_uploader(
        "2) Upload Redcap Rotation list CSV (must have a 'legal_name' column)",
        type=["csv"],
    )

    record_id = "peds_clerkship"

    if not schedule_files or not student_file or not record_id:
        st.info("Please upload schedule Excel(s) and the student CSV.")
        st.stop()

    # ─────────────────────────────────────────────────────────────────────────────
    # Parse QGenda files → discover (date → {designation: [providers]}) + designation list
    # ─────────────────────────────────────────────────────────────────────────────
    assignments_by_date: dict = {}
    designation_counts = defaultdict(int)
    designation_prov_sample = defaultdict(set)
    found_keywords = set()

    for file in schedule_files:
        try:
            df = pd.read_excel(file, header=None, dtype=str)
        except Exception as e:
            st.error(f"Error reading {getattr(file, 'name', 'uploaded file')}: {e}")
            continue

        flat_vals = df.astype(str).apply(lambda s: s.str.lower()).values.flatten().tolist()
        for kw in REQUIRED_KEYWORDS:
            if any(kw in v for v in flat_vals):
                found_keywords.add(kw)

        # find date headers
        date_positions = []
        for r in range(df.shape[0]):
            for c in range(df.shape[1]):
                val = str(df.iat[r, c]).replace("\xa0", " ").strip()
                if date_pat.match(val):
                    try:
                        d = pd.to_datetime(val).date()
                        date_positions.append((d, r, c))
                    except Exception:
                        pass

        # Deduplicate to top-most for each date
        unique = {}
        for d, r, c in date_positions:
            if d not in unique or r < unique[d][0]:
                unique[d] = (r, c)

        for d, (row0, col0) in unique.items():
            grp = assignments_by_date.setdefault(d, defaultdict(list))
            for r in range(row0 + 1, df.shape[0]):
                raw = str(df.iat[r, col0]).replace("\xa0", " ").strip()
                if raw == "" or date_pat.match(raw):
                    break
                desc = raw.lower()
                prov = str(df.iat[r, col0 + 1]).strip()
                if desc and prov:
                    grp[desc].append(prov)
                    designation_counts[desc] += 1
                    if len(designation_prov_sample[desc]) < 5:
                        designation_prov_sample[desc].add(prov)

    missing_keywords = [kw for kw in REQUIRED_KEYWORDS if kw not in found_keywords]
    if missing_keywords:
        st.warning("These site keywords weren’t detected in your uploads: " + ", ".join(missing_keywords))

    unique_designations = sorted(designation_counts.keys())
    if not unique_designations:
        st.error("No designations found. Check that your Excel(s) have date headers like 'September 1, 2025' and designations in the next rows.")
        st.stop()

    # ─────────────────────────────────────────────────────────────────────────────
    # Students list
    # ─────────────────────────────────────────────────────────────────────────────
    students_df = pd.read_csv(student_file, dtype=str)
    if "legal_name" not in students_df.columns:
        st.error("Student CSV must contain a 'legal_name' column.")
        st.stop()
    legal_names = students_df["legal_name"].dropna().tolist()

    # ─────────────────────────────────────────────────────────────────────────────
    # Provider filter UI
    # ─────────────────────────────────────────────────────────────────────────────
    all_providers = sorted({
        p.strip()
        for day in assignments_by_date.values()
        for provs in day.values()
        for p in provs
        if isinstance(p, str) and p.strip()
    })

    if "provider_filter" not in st.session_state:
        st.session_state["provider_filter"] = []

    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("Select All Providers"):
            st.session_state["provider_filter"] = all_providers
    with c2:
        if st.button("Clear Providers"):
            st.session_state["provider_filter"] = []
    with c3:
        apply_provider_filter = st.checkbox(
            "Apply provider filter",
            value=False,
            help="When OFF, everyone is included even if the multiselect is blank.",
        )

    allowed_providers = st.multiselect(
        "Limit providers included in OPD",
        options=all_providers,
        key="provider_filter",
        help="Only selected providers will be written when 'Apply provider filter' is ON.",
    )
    effective_allowed = set(allowed_providers) if (apply_provider_filter and allowed_providers) else set(all_providers)

    # ─────────────────────────────────────────────────────────────────────────────
    # Routing UI (see/cull designations; choose sheet/tab + AM/PM policy)
    # ─────────────────────────────────────────────────────────────────────────────
    st.subheader("Designations → Routing")

    with st.expander("Preview detected designations (toggle what to include)"):
        if "route_map" not in st.session_state:
            st.session_state["route_map"] = {}

        default_slot_for_none = st.radio(
            "Default slot for designations with no 'am'/'pm' in the name",
            options=["AM", "PM", "Split evenly", "Skip"],
            index=0,
            help="Used only if a designation text has no AM/PM and you don’t override it below.",
        )

        # Which sheets are *dedicated* tabs?
        dedicated_tabs_default = {"HOPE_DRIVE", "LANCASTER", "NYES", "ETOWN"}
        dedicated_tabs = st.multiselect(
            "Sheets to keep on separate tabs (default 4)",
            options=["HOPE_DRIVE", "LANCASTER", "NYES", "ETOWN"],
            default=list(dedicated_tabs_default),
            help="Anything not selected here routes to SUBSPECIALTY unless added as an extra tab below.",
        )

        # Any extra named tabs?
        extra_separate_tabs_txt = st.text_input(
            "Optional: create additional separate continuity tabs (comma-separated names)",
            value="",
            help="Example: Neurology Hope Drive, Endocrinology Camp Hill",
        )
        extra_separate_tabs = [s.strip() for s in extra_separate_tabs_txt.split(",") if s.strip()]

        st.write("Select routing for each designation:")
        for i, des in enumerate(unique_designations):
            prev = st.session_state["route_map"].get(des, {})
            slug = slugify(des) or f"item{i}"
            keybase = f"route_{i}_{slug}"
        
            colA, colB, colC, colD = st.columns([0.42, 0.24, 0.2, 0.14])
        
            include = colA.checkbox(
                f"{des}",
                value=prev.get("include", True),
                help="Uncheck to ignore this designation everywhere.",
                key=f"{keybase}_include",
            )
        
            suggested = guess_site_bucket(des)
            bucket_choices = ["HOPE_DRIVE", "LANCASTER", "NYES", "ETOWN", "SUBSPECIALTY"] + extra_separate_tabs
            try_idx = bucket_choices.index(prev.get("sheet", suggested)) if prev.get("sheet") in bucket_choices else bucket_choices.index(suggested)
        
            sheet_choice = colB.selectbox(
                "Sheet",
                options=bucket_choices,
                index=try_idx,
                key=f"{keybase}_sheet",
            )
        
            gslot = guess_slot(des)  # "AM"/"PM"/"NONE"
            slot_default = "Auto (from name)" if gslot in ("AM", "PM") else "Use default"
            slot_choice = colC.selectbox(
                "Slot",
                options=["Auto (from name)", "AM", "PM", "Both", "Use default"],
                index=["Auto (from name)", "AM", "PM", "Both", "Use default"].index(prev.get("slot", slot_default)),
                key=f"{keybase}_slot",
            )
        
            # sample providers
            colD.write("Sample:")
            colD.caption(", ".join(sorted(list(designation_prov_sample[des]))))
        
            st.session_state["route_map"][des] = {
                "include": include,
                "sheet": sheet_choice,
                "slot": slot_choice,
            }


    # ─────────────────────────────────────────────────────────────────────────────
    # Build sheet definitions (prefix registry)
    # ─────────────────────────────────────────────────────────────────────────────
    def build_sheet_prefixes(dedicated_tabs_set, extra_tabs):
        sheet_defs = {}

        # HOPE_DRIVE (special)
        if "HOPE_DRIVE" in dedicated_tabs_set:
            sheet_defs["HOPE_DRIVE"] = {
                "type": "hope_drive",
                "am_prefix": "hd_am_",
                "pm_prefix": "hd_pm_",
                "am_acute_prefix": "hd_am_acute_",
                "pm_acute_prefix": "hd_pm_acute_",
                "wknd_am_prefix": "hd_wknd_am_",
                "wknd_pm_acute_prefix": "hd_wknd_pm_acute_",
                "wknd_acute1_prefix": "hd_wknd_acute_1_",
                "wknd_acute2_prefix": "hd_wknd_acute_2_",
            }

        # Continuity helper
        def add_continuity_sheet(key, am_pref, pm_pref):
            sheet_defs[key] = {"type": "continuity", "am_prefix": am_pref, "pm_prefix": pm_pref}

        if "LANCASTER" in dedicated_tabs_set:
            add_continuity_sheet("LANCASTER", "lancaster_am_", "lancaster_pm_")
        if "NYES" in dedicated_tabs_set:
            add_continuity_sheet("NYES", "nyes_am_", "nyes_pm_")
        if "ETOWN" in dedicated_tabs_set:
            add_continuity_sheet("ETOWN", "etown_am_", "etown_pm_")

        # Extra user-defined tabs
        for tab in extra_tabs:
            base = slugify(tab)
            add_continuity_sheet(tab, f"{base}_am_", f"{base}_pm_")

        # SUBSPECIALTY catch-all with extra capacity (20 rows/slot/day)
        add_continuity_sheet("SUBSPECIALTY", "subspec_am_", "subspec_pm_")
        return sheet_defs

    sheet_defs = build_sheet_prefixes(set(dedicated_tabs), extra_separate_tabs)

    # ─────────────────────────────────────────────────────────────────────────────
    # designation_map: designation → list of (sheet_key, prefix_kind)
    # ─────────────────────────────────────────────────────────────────────────────
    designation_map = defaultdict(list)

    # HOPE Drive acute minimums
    min_required = {
        "hope drive am acute precept": 2,
        "hope drive pm acute precept": 2,
    }

    for des, cfg in st.session_state["route_map"].items():
        if not cfg.get("include", True):
            continue
        target_sheet = cfg["sheet"]
        slot_sel = cfg["slot"]  # "Auto (from name)" | "AM" | "PM" | "Both" | "Use default"

        gslot = guess_slot(des)  # "AM"/"PM"/"NONE"
        if slot_sel == "Auto (from name)":
            slot = gslot if gslot in ("AM", "PM") else "NONE"
        elif slot_sel in ("AM", "PM", "Both"):
            slot = slot_sel
        else:
            # "Use default"
            slot = {"AM": "AM", "PM": "PM", "Split evenly": "Both", "Skip": "NONE"}[default_slot_for_none]

        # Normalize unknown sheets to SUBSPECIALTY
        if target_sheet not in sheet_defs and target_sheet != "HOPE_DRIVE":
            target_sheet = "SUBSPECIALTY"

        # HOPE_DRIVE special acutes (preserve exact strings)
        if target_sheet == "HOPE_DRIVE" and "acute" in des:
            acute_map = {
                "hope drive am acute precept": ("HOPE_DRIVE", "am_acute_prefix"),
                "hope drive pm acute precept": ("HOPE_DRIVE", "pm_acute_prefix"),
                "hope drive weekend acute 1": ("HOPE_DRIVE", "wknd_acute1_prefix"),
                "hope drive weekend acute 2": ("HOPE_DRIVE", "wknd_acute2_prefix"),
                "hope drive weekend continuity": ("HOPE_DRIVE", "wknd_am_prefix"),
            }
            if des in acute_map:
                designation_map[des].append(acute_map[des])
            continue

        # Continuity routing
        if target_sheet in sheet_defs and sheet_defs[target_sheet]["type"] == "continuity":
            if slot in ("AM", "Both"):
                designation_map[des].append((target_sheet, "am_prefix"))
            if slot in ("PM", "Both"):
                designation_map[des].append((target_sheet, "pm_prefix"))
            # slot == "NONE" → skip

    # ─────────────────────────────────────────────────────────────────────────────
    # Build redcap row
    # ─────────────────────────────────────────────────────────────────────────────
    NUM_WEEKS = 5
    NUM_DAYS = NUM_WEEKS * 7
    sorted_dates = sorted(assignments_by_date.keys())

    redcap_row = {"record_id": record_id}

    for idx, date in enumerate(sorted_dates[:NUM_DAYS], start=1):
        redcap_row[f"hd_day_date{idx}"] = date
        suffix = f"d{idx}_"

        for des, provs in assignments_by_date[date].items():
            if des not in designation_map:
                continue

            # provider filter
            filtered = [p for p in provs if p in effective_allowed]

            # HOPE acute mins
            req = min_required.get(des, len(filtered))
            if filtered:
                while len(filtered) < req:
                    filtered.append(filtered[0])

            for (sheet_key, prefix_kind) in designation_map[des]:
                pref = sheet_defs[sheet_key][prefix_kind] + suffix
                for i, name in enumerate(filtered, start=1):
                    redcap_row[f"{pref}{i}"] = name

    # Students passthrough
    for i, name in enumerate(legal_names, start=1):
        redcap_row[f"s{i}"] = name

    out_df = pd.DataFrame([redcap_row])
    for c in list(out_df.columns):
        if c.startswith("hd_day_date"):
            out_df[c] = pd.to_datetime(out_df[c]).dt.strftime("%m-%d-%Y")
    csv_full = out_df.to_csv(index=False).encode("utf-8")

    # ─────────────────────────────────────────────────────────────────────────────
    # Excel generator (dynamic tabs)
    # ─────────────────────────────────────────────────────────────────────────────
    def generate_opd_workbook(full_df: pd.DataFrame, sheet_defs: dict) -> bytes:
        import xlsxwriter

        output = io.BytesIO()
        workbook = xlsxwriter.Workbook(output, {"in_memory": True})

        # formats
        format1   = workbook.add_format({"font_size": 18, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#FEFFCC", "border": 1})
        format4   = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#8ccf6f", "border": 1})
        format4a  = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#9fc5e8", "border": 1})
        format5a  = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#d0e9ff", "border": 1})
        format3   = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#FFC7CE", "border": 1})
        format2   = workbook.add_format({"bg_color": "black"})
        format_dt = workbook.add_format({"num_format": "m/d/yyyy", "font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#FFC7CE", "border": 1})
        format_lb = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#FFC7CE", "border": 1})
        merge_fmt = workbook.add_format({"bold": 1, "align": "center", "valign": "vcenter", "text_wrap": True, "font_color": "red", "bg_color": "#FEFFCC", "border": 1})

        format_am_acute = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#8ccf6f", "border": 1})
        format_pm_acute = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "white", "bg_color": "#1f4e79", "border": 1})

        # Sheet ordering
        worksheet_names = []
        if "HOPE_DRIVE" in sheet_defs: worksheet_names.append("HOPE_DRIVE")
        for k in ("ETOWN", "NYES", "COMPLEX", "LANCASTER"):
            if k in sheet_defs: worksheet_names.append(k)
        # Add any extra continuity tabs user created
        for k in sheet_defs:
            if k not in worksheet_names and k not in ("SUBSPECIALTY",):
                if sheet_defs[k]["type"] == "continuity":
                    worksheet_names.append(k)
        if "SUBSPECIALTY" in sheet_defs: worksheet_names.append("SUBSPECIALTY")

        title_map = {
            "ETOWN": "Elizabethtown",
            "NYES": "Nyes Road",
            "LANCASTER": "Lancaster",
            "HOPE_DRIVE": "Hope Drive",
            "SUBSPECIALTY": "Subspecialty Clinics",
        }

        sheets = {name: workbook.add_worksheet(name) for name in worksheet_names}
        for name, ws in sheets.items():
            ws.set_row(0, 37.25)
            ws.set_column("A:A", 15 if name != "SUBSPECIALTY" else 30)
            ws.set_column("B:H", 65)
            ws.write(0, 0, "Site:", format1)
            ws.write(0, 1, title_map.get(name, name), format1)

        # Dates
        date_cols = [f"hd_day_date{i}" for i in range(1, 5 * 7 + 1)]
        dates = pd.to_datetime(full_df[date_cols].iloc[0], errors="coerce").tolist()
        weeks = [dates[i * 7:(i + 1) * 7] for i in range(5)]
        days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

        # geometry
        DEF_BLOCK_H = 24
        def_hdr_starts = [2 + i * DEF_BLOCK_H for i in range(5)]
        def_blk_starts = [6 + i * DEF_BLOCK_H for i in range(5)]

        # CRTS note
        def write_crts_note(ws):
            text1 = (
                "Students are to alert their preceptors when they have a Clinical "
                "Reasoning Teaching Session (CRTS). Please allow the students to "
                "leave ~15 minutes prior to the start of their session so they can be prepared."
            )
            ws.merge_range("C1:F1", text1, merge_fmt)
            ws.write("G1", "", merge_fmt)
            ws.write("H1", "", merge_fmt)

        # HOPE_DRIVE
        if "HOPE_DRIVE" in sheets:
            hd = sheets["HOPE_DRIVE"]
            write_crts_note(hd)
            ACUTE_COUNT, CONT_COUNT = 2, 8
            for start in def_blk_starts:
                # acutes (top two rows) & continuity
                hd.conditional_format(f"A{start}:H{start+1}", {"type": "no_errors", "format": format_am_acute})
                hd.conditional_format(f"A{start+10}:H{start+11}", {"type": "no_errors", "format": format_pm_acute})

                hd.conditional_format(f"B{start}:H{start}", {"type": "no_errors", "format": format4})
                hd.conditional_format(f"B{start+10}:H{start+10}", {"type": "no_errors", "format": format4a})

                hd.conditional_format(f"A{start}:H{start+9}", {"type": "no_errors", "format": format1})
                hd.conditional_format(f"A{start+10}:H{start+19}", {"type": "no_errors", "format": format5a})

                zero = start - 1
                for i in range(ACUTE_COUNT + CONT_COUNT):
                    hd.write(zero + i, 0, "AM - ACUTES" if i < ACUTE_COUNT else "AM - Continuity", format5a)
                    hd.write(zero + 10 + i, 0, "PM - ACUTES" if i < ACUTE_COUNT else "PM - Continuity", format5a)

        # Continuity (standard 10 rows per AM/PM)
        for name in worksheet_names:
            if name in ("HOPE_DRIVE", "SUBSPECIALTY"):
                continue
            ws = sheets[name]
            write_crts_note(ws)
            ws.set_zoom(80)
            for idx, hstart in enumerate(def_hdr_starts):
                # headers and dates
                for c, dname in enumerate(days):
                    ws.write(hstart, 1 + c, dname, format3)
                for c, val in enumerate(weeks[idx]):
                    if pd.notna(val):
                        ws.write(hstart + 1, 1 + c, val, format_dt)
                ws.write_formula(f"A{hstart}", '""', format_lb)
                ws.conditional_format(f"A{hstart+3}:H{hstart+3}", {"type": "no_errors", "format": format_lb})

            # bars and blocks
            for row in [2 + i * DEF_BLOCK_H for i in range(5)]:
                ws.merge_range(f"A{row}:H{row}", " ", format2)
                ws.write(f"A{row+1}", "", format_dt)
                ws.write(f"A{row+2}", "", format_dt)

            for start in def_blk_starts:
                zero = start - 1
                for i in range(10):
                    ws.write(zero + i, 0, "AM", format5a)
                    ws.write(zero + 10 + i, 0, "PM", format5a)
                ws.conditional_format(f"A{start}:H{start+9}", {"type": "no_errors", "format": format1})
                ws.conditional_format(f"A{start+10}:H{start+19}", {"type": "no_errors", "format": format5a})
                ws.conditional_format(f"B{start}:H{start}", {"type": "no_errors", "format": format4})
                ws.conditional_format(f"B{start+10}:H{start+10}", {"type": "no_errors", "format": format4a})

        # SUBSPECIALTY (larger capacity: 20 rows per AM/PM)
        if "SUBSPECIALTY" in sheets:
            ws = sheets["SUBSPECIALTY"]
            write_crts_note(ws)
            ws.set_zoom(75)
            # geometry for 20-row continuity blocks per AM/PM
            SUB_CONT_ROWS = 20
            SUB_BLOCK_H = 4 + (SUB_CONT_ROWS * 2)  # header rows + AM/PM sections
            sub_hdr_starts = [2 + i * SUB_BLOCK_H for i in range(5)]
            sub_blk_starts = [6 + i * SUB_BLOCK_H for i in range(5)]

            for idx, hstart in enumerate(sub_hdr_starts):
                for c, dname in enumerate(days):
                    ws.write(hstart, 1 + c, dname, format3)
                for c, val in enumerate(weeks[idx]):
                    if pd.notna(val):
                        ws.write(hstart + 1, 1 + c, val, format_dt)
                ws.write_formula(f"A{hstart}", '""', format_lb)
                ws.conditional_format(f"A{hstart+3}:H{hstart+3}", {"type": "no_errors", "format": format_lb})
                ws.merge_range(f"A{hstart}:H{hstart}", " ", format2)
                ws.write(f"A{hstart+1}", "", format_dt)
                ws.write(f"A{hstart+2}", "", format_dt)

            for start in sub_blk_starts:
                zero = start - 1
                # AM labels (20), PM labels (20)
                for i in range(SUB_CONT_ROWS):
                    ws.write(zero + i, 0, "AM", format5a)
                    ws.write(zero + SUB_CONT_ROWS + i, 0, "PM", format5a)
                # backgrounds + header tints
                ws.conditional_format(f"A{start}:H{start+SUB_CONT_ROWS-1}", {"type": "no_errors", "format": format1})
                ws.conditional_format(f"A{start+SUB_CONT_ROWS}:H{start+SUB_CONT_ROWS*2-1}", {"type": "no_errors", "format": format5a})
                ws.conditional_format(f"B{start}:H{start}", {"type": "no_errors", "format": format4})
                ws.conditional_format(f"B{start+SUB_CONT_ROWS}:H{start+SUB_CONT_ROWS}", {"type": "no_errors", "format": format4a})

        workbook.close()
        output.seek(0)
        return output.read()

    # ─────────────────────────────────────────────────────────────────────────────
    # Hide blank rows where A starts with AM/PM and all B..H empty
    # ─────────────────────────────────────────────────────────────────────────────
    def hide_blank_rows_all_sheets(excel_bytes: bytes):
        import io as _io
        import re as _re
        from openpyxl import load_workbook

        def _empty(v):
            return v is None or (isinstance(v, str) and v.strip() == "")

        wb = load_workbook(_io.BytesIO(excel_bytes))
        per_sheet = {}
        total = 0
        for ws in wb.worksheets:
            hidden = 0
            for r in range(1, ws.max_row + 1):
                a1 = ws.cell(row=r, column=1).value
                if not (isinstance(a1, str) and _re.match(r"^\s*(AM|PM)\b", a1, _re.IGNORECASE)):
                    continue
                if all(_empty(ws.cell(row=r, column=c).value) for c in range(2, 9)):
                    ws.row_dimensions[r].hidden = True
                    ws.row_dimensions[r].height = 0
                    hidden += 1
            per_sheet[ws.title] = hidden
            total += hidden
        out = _io.BytesIO()
        wb.save(out)
        out.seek(0)
        return out.getvalue(), per_sheet, total

    # ─────────────────────────────────────────────────────────────────────────────
    # CSV → Excel updater
    # ─────────────────────────────────────────────────────────────────────────────
    def update_excel_from_csv(excel_template_bytes: bytes, csv_data_bytes: bytes, mappings: list) -> bytes | None:
        from openpyxl import load_workbook
        from openpyxl.styles import Alignment
        try:
            df_csv = pd.read_csv(io.BytesIO(csv_data_bytes))
            if df_csv.empty:
                return None
            wb = load_workbook(io.BytesIO(excel_template_bytes))
            for m in mappings:
                col = m["csv_column"]
                if col not in df_csv.columns:
                    continue
                val = str(df_csv.loc[0, col]).strip()
                ws = wb[m["excel_sheet"]]
                ws[m["excel_cell"]] = val if " ~ " in val else (val + " ~ " if val else "")
                ws[m["excel_cell"]].alignment = Alignment(horizontal="center", vertical="center")
            out = io.BytesIO()
            wb.save(out)
            out.seek(0)
            return out.getvalue()
        except Exception:
            return None

    # ─────────────────────────────────────────────────────────────────────────────
    # Build mappings dynamically from sheet_defs
    # ─────────────────────────────────────────────────────────────────────────────
    def build_mappings(NUM_WEEKS: int, sheet_defs: dict) -> list:
        mappings = []
        excel_column_letters = ["B", "C", "D", "E", "F", "G", "H"]
        DEF_BLOCK_H = 24
        cont_row_defs = {"AM": 6, "PM": 16}
        hd_row_defs = {
            "AM": {"acute_start": 6, "cont_start": 8},
            "PM": {"acute_start": 16, "cont_start": 18},
        }

        # HOPE_DRIVE (unchanged from your spec)
        if "HOPE_DRIVE" in sheet_defs:
            for week_idx in range(1, NUM_WEEKS + 1):
                week_base = (week_idx - 1) * DEF_BLOCK_H
                day_offset = (week_idx - 1) * 7
                for day_idx, col in enumerate(excel_column_letters, start=1):
                    is_weekday = day_idx <= 5
                    day_num = day_idx + day_offset
                    # AM
                    if is_weekday:
                        for prov in range(1, 2 + 1):
                            row = week_base + hd_row_defs["AM"]["acute_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_am_acute_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                        for prov in range(1, 8 + 1):
                            row = week_base + hd_row_defs["AM"]["cont_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_am_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                    else:
                        for acute_type in (1, 2):
                            row = week_base + hd_row_defs["AM"]["acute_start"] + (acute_type - 1)
                            mappings.append({"csv_column": f"hd_wknd_acute_{acute_type}_d{day_num}_1", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                        for prov in range(1, 8 + 1):
                            row = week_base + hd_row_defs["AM"]["cont_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_wknd_am_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                    # PM
                    if is_weekday:
                        for prov in range(1, 2 + 1):
                            row = week_base + hd_row_defs["PM"]["acute_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_pm_acute_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                        for prov in range(1, 8 + 1):
                            row = week_base + hd_row_defs["PM"]["cont_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_pm_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                    else:
                        for acute_type in (1, 2):
                            row = week_base + hd_row_defs["PM"]["acute_start"] + (acute_type - 1)
                            mappings.append({"csv_column": f"hd_wknd_pm_acute_{acute_type}_d{day_num}_1", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                        for prov in range(1, 8 + 1):
                            row = week_base + hd_row_defs["PM"]["cont_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_wknd_pm_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})

        # Continuity tabs (all present in sheet_defs except SUBSPECIALTY)
        for sheet_name, meta in sheet_defs.items():
            if sheet_name in ("HOPE_DRIVE", "SUBSPECIALTY") or meta["type"] != "continuity":
                continue
            am_pref = meta["am_prefix"]
            pm_pref = meta["pm_prefix"]
            for week_idx in range(1, NUM_WEEKS + 1):
                week_base = (week_idx - 1) * DEF_BLOCK_H
                day_offset = (week_idx - 1) * 7
                for day_idx, col in enumerate(excel_column_letters, start=1):
                    day_num = day_idx + day_offset
                    for prov in range(1, 10 + 1):  # 10 rows AM
                        row = week_base + cont_row_defs["AM"] + (prov - 1)
                        mappings.append({"csv_column": f"{am_pref}d{day_num}_{prov}", "excel_sheet": sheet_name, "excel_cell": f"{col}{row}"})
                    for prov in range(1, 10 + 1):  # 10 rows PM
                        row = week_base + cont_row_defs["PM"] + (prov - 1)
                        mappings.append({"csv_column": f"{pm_pref}d{day_num}_{prov}", "excel_sheet": sheet_name, "excel_cell": f"{col}{row}"})

        # SUBSPECIALTY: larger capacity (20 rows AM/PM)
        if "SUBSPECIALTY" in sheet_defs:
            am_pref = sheet_defs["SUBSPECIALTY"]["am_prefix"]
            pm_pref = sheet_defs["SUBSPECIALTY"]["pm_prefix"]
            SUB_CONT_ROWS = 20
            SUB_BLOCK_H = 4 + (SUB_CONT_ROWS * 2)
            for week_idx in range(1, NUM_WEEKS + 1):
                week_base = (week_idx - 1) * SUB_BLOCK_H
                day_offset = (week_idx - 1) * 7
                base_start = week_base + 6
                am_row0 = base_start
                pm_row0 = am_row0 + SUB_CONT_ROWS
                for day_idx, col in enumerate(excel_column_letters, start=1):
                    day_num = day_idx + day_offset
                    for prov in range(1, SUB_CONT_ROWS + 1):
                        mappings.append({"csv_column": f"{am_pref}d{day_num}_{prov}", "excel_sheet": "SUBSPECIALTY", "excel_cell": f"{col}{am_row0 + (prov - 1)}"})
                    for prov in range(1, SUB_CONT_ROWS + 1):
                        mappings.append({"csv_column": f"{pm_pref}d{day_num}_{prov}", "excel_sheet": "SUBSPECIALTY", "excel_cell": f"{col}{pm_row0 + (prov - 1)}"})
        return mappings

    # ─────────────────────────────────────────────────────────────────────────────
    # Generate & Download
    # ─────────────────────────────────────────────────────────────────────────────
    st.subheader("Generate OPD.xlsx (configurable, 5-week)")
    if st.button("Generate OPD (5-week, dynamic)"):
        excel_template_bytes = generate_opd_workbook(out_df, sheet_defs)
        mappings = build_mappings(NUM_WEEKS, sheet_defs)
        updated_excel_bytes = update_excel_from_csv(excel_template_bytes, csv_full, mappings)
        if not updated_excel_bytes:
            st.error("Failed to update OPD.xlsx with data.")
            st.stop()
        updated_excel_bytes, hidden_map, hidden_total = hide_blank_rows_all_sheets(updated_excel_bytes)
        st.success(f"✅ OPD.xlsx updated! (hidden rows: {hidden_total})")
        st.download_button(
            label="⬇️ Download OPD.xlsx",
            data=updated_excel_bytes,
            file_name="OPD_dynamic_5_weeks.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )


elif mode == "Create Individual Schedules":
    st.subheader("Individual Schedule Creator")

    uploaded = st.file_uploader(
        "Upload the master Excel (.xlsx) with one tab per person",
        type=["xlsx"]
    )

    def copy_sheet_to_new_wb(src_ws):
        """Return a BytesIO of a new .xlsx containing src_ws with formatting."""
        from openpyxl import Workbook
        from io import BytesIO
    
        wb_new = Workbook()
        ws_new = wb_new.active
        ws_new.title = src_ws.title[:31]
    
        # Column widths & visibility
        for col_letter, dim in src_ws.column_dimensions.items():
            if dim.width is not None:
                ws_new.column_dimensions[col_letter].width = dim.width
            ws_new.column_dimensions[col_letter].hidden = dim.hidden
    
        # Row heights & visibility
        for idx, dim in src_ws.row_dimensions.items():
            if dim.height is not None:
                ws_new.row_dimensions[idx].height = dim.height
            ws_new.row_dimensions[idx].hidden = dim.hidden
    
        # Sheet settings (best-effort)
        try:
            ws_new.sheet_format.defaultColWidth = src_ws.sheet_format.defaultColWidth
            ws_new.sheet_format.defaultRowHeight = src_ws.sheet_format.defaultRowHeight
        except Exception:
            pass
        ws_new.freeze_panes = src_ws.freeze_panes
        try:
            ws_new.page_setup.orientation = src_ws.page_setup.orientation
            ws_new.page_setup.fitToWidth = src_ws.page_setup.fitToWidth
            ws_new.page_setup.fitToHeight = src_ws.page_setup.fitToHeight
            ws_new.page_margins = src_ws.page_margins
            ws_new.print_options.horizontalCentered = src_ws.print_options.horizontalCentered
            ws_new.print_options.verticalCentered = src_ws.print_options.verticalCentered
            ws_new.print_area = src_ws.print_area
        except Exception:
            pass

        try:
            # Set zoom to 70%
            ws_new.sheet_view.zoomScale = 70
    
            # Set column widths
            ws_new.column_dimensions["A"].width = 20
            ws_new.column_dimensions["B"].width = 30
            for col in ["C", "D", "E", "F", "G"]:
                ws_new.column_dimensions[col].width = 40
            ws_new.column_dimensions["H"].width = 155
        except Exception:
            pass
    
        # Copy cells: values + (copied) styles
        for row in src_ws.iter_rows():
            for cell in row:
                # Skip non-master cells from merged ranges
                if isinstance(cell, MergedCell):
                    continue
        
                # Get a reliable numeric column index
                col_idx = getattr(cell, "col_idx", None)
                if col_idx is None:
                    col = cell.column  # may be int or letter depending on version
                    col_idx = col if isinstance(col, int) else column_index_from_string(col)
        
                # Create target cell with value (formula preserved if present)
                tgt = ws_new.cell(row=cell.row, column=col_idx, value=cell.value)
        
                # Copy style safely
                if getattr(cell, "has_style", False):
                    try:
                        from copy import copy
                        if cell.font:        tgt.font        = copy(cell.font)
                        if cell.fill:        tgt.fill        = copy(cell.fill)
                        if cell.border:      tgt.border      = copy(cell.border)
                        if cell.alignment:   tgt.alignment   = copy(cell.alignment)
                        if cell.protection:  tgt.protection  = copy(cell.protection)
                        tgt.number_format = cell.number_format
                    except Exception:
                        pass
    
        # Copy merged cell ranges (after values)
        for merged in list(src_ws.merged_cells.ranges):
            try:
                ws_new.merge_cells(str(merged))
            except Exception:
                pass
    
        # Copy data validations (best-effort)
        try:
            if src_ws.data_validations and src_ws.data_validations.dataValidation:
                from openpyxl.worksheet.datavalidation import DataValidation
                for dv in src_ws.data_validations.dataValidation:
                    dv_new = DataValidation(
                        type=dv.type,
                        formula1=dv.formula1,
                        formula2=dv.formula2,
                        allow_blank=dv.allow_blank,
                        operator=dv.operator,
                        showDropDown=dv.showDropDown,
                        showErrorMessage=dv.showErrorMessage,
                        errorTitle=dv.errorTitle,
                        error=dv.error,
                        promptTitle=dv.promptTitle,
                        prompt=dv.prompt
                    )
                    # Copy cell refs
                    for sqref in getattr(dv, "sqref", []):
                        dv_new.add(sqref)
                    ws_new.add_data_validation(dv_new)
        except Exception:
            pass
    
        # Filters
        try:
            ws_new.auto_filter.ref = getattr(src_ws.auto_filter, "ref", None)
        except Exception:
            pass
    
        # Save to buffer
        buf = BytesIO()
        wb_new.save(buf)
        buf.seek(0)
        return buf
        
    if uploaded is not None:
        # Keep formulas/formatting -> data_only=False
        wb = load_workbook(uploaded, data_only=False)
        st.write(f"Found **{len(wb.sheetnames)}** tabs.")

        if st.button("Split tabs and build ZIP"):
            zip_buf = BytesIO()
            with ZipFile(zip_buf, mode="w", compression=ZIP_DEFLATED) as zf:
                used_names = defaultdict(int)

                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]

                    # Skip truly empty sheets (no cells with value)
                    has_any_value = any(cell.value is not None for row in ws.iter_rows() for cell in row)
                    if not has_any_value:
                        continue

                    # Safe file name
                    base = re.sub(r"[^A-Za-z0-9._-]+", "_", sheet_name).strip("_") or "sheet"
                    used_names[base] += 1
                    safe_name = base if used_names[base] == 1 else f"{base}_{used_names[base]}"

                    # Copy this sheet into its own new workbook (preserving formatting)
                    out_buf = copy_sheet_to_new_wb(ws)

                    # Add to ZIP
                    zf.writestr(f"{safe_name}.xlsx", out_buf.getvalue())

            zip_buf.seek(0)
            st.download_button(
                label="Download individual schedules (ZIP)",
                data=zip_buf,
                file_name="individual_schedules_formatted.zip",
                mime="application/zip",
            )

elif mode == "Create Student Schedule":
    import io
    import re
    import pandas as pd
    import xlsxwriter
    from collections import defaultdict
    from datetime import timedelta
    import streamlit as st
    from openpyxl import load_workbook

    st.subheader("Create Student Schedule")

    # ========= Configure here =========
    NUM_WEEKS = 5  # ← change to 4/5/6... if ever needed
    BLOCK_HEIGHT = 8  # visual rows per week segment in your template pattern
    # ==================================

    def save_to_session(filename, fileobj, namespace="uploaded_files"):
        st.session_state.setdefault(namespace, {})[filename] = fileobj

    # ───────── Helper to load & stash uploads ─────────
    def load_workbook_df(label, types, key):
        upload = st.file_uploader(label, type=types, key=key)
        if not upload:
            st.info(f"Please upload {label}.")
            return None
        st.session_state[f"{key}_file"] = upload
        try:
            if upload.name.lower().endswith(".csv"):
                return pd.read_csv(upload)
            else:
                return pd.read_excel(upload)
        except Exception as e:
            st.error(f"Error loading {upload.name}: {e}")
            return None

    def _seq(start, step, n):
        """Helper: n values starting at start, step apart (1‑based row numbers)."""
        return [start + i * step for i in range(n)]

    def create_ms_schedule_template(students, dates):
        buf = io.BytesIO()
        wb = xlsxwriter.Workbook(buf, {"in_memory": True})

        # — Formats — (unchanged palette)
        f1 = wb.add_format({"font_size": 14, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "text_wrap": True, "bg_color": "#FEFFCC", "border": 1})
        f2 = wb.add_format({"font_size": 10, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "yellow", "bg_color": "black", "border": 1, "text_wrap": True})
        f3 = wb.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "bg_color": "#FFC7CE", "border": 1})
        f4 = wb.add_format({"num_format": "mm/dd/yyyy", "font_size": 12, "bold": 1, "align": "center",
                            "valign": "vcenter", "font_color": "black", "bg_color": "#F4F6F7", "border": 1})
        f5 = wb.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "bg_color": "#F4F6F7", "border": 1})
        f6 = wb.add_format({"bg_color": "black", "border": 1})
        f7 = wb.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "bg_color": "#90EE90", "border": 1})
        f8 = wb.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "bg_color": "#89CFF0", "border": 1})

        days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        # Derive all row anchors for N weeks (1‑based for A1 merges; numeric writes may be slightly offset like your original)
        start_rows = _seq(2, BLOCK_HEIGHT, NUM_WEEKS)           # day headers row per week (2,10,18,26,34)
        week_label_rows = _seq(4, BLOCK_HEIGHT, NUM_WEEKS)      # "Week i" labels (4,12,20,28,36)
        am_rows = _seq(6, BLOCK_HEIGHT, NUM_WEEKS)              # AM label row (6,14,22,30,38)
        pm_rows = _seq(7, BLOCK_HEIGHT, NUM_WEEKS)              # PM label row (7,15,23,31,39)
        green_fillers = _seq(8, BLOCK_HEIGHT, NUM_WEEKS)        # green row anchors (8,16,24,32,40)
        separators = _seq(10, BLOCK_HEIGHT, NUM_WEEKS)          # black bar rows (10,18,26,34,42)

        # Optional: extend your due text list to 5, with a safe fallback
        due_texts = [
            "Quiz 1 Due",
            "Quiz 2, Pediatric Documentation #1, 1 Clinical Encounter Log Due",
            "Quiz 3 Due",
            "Quiz 4, Pediatric Documentation #2, Social Drivers of Health Assessment Form, Developmental Assessment of Pediatric Patient Form, All Clinical Encounter Logs are Due!",
            "Final Evaluations & Wrap‑Up Due",  # Week 5 (new)
        ]

        for name in students:
            title = name[:31].replace("/", "-").replace("\\", "-")
            ws = wb.add_worksheet(title)
            ws.set_zoom(70)

            # Header
            ws.merge_range("A1:A2", "Student Name:", f1)
            ws.merge_range("B1:B2", title, f1)
            note = (
                "*Note* Asynchronous time is for coursework only. During this time period, "
                "we expect students to do coursework, be available for any additional educational "
                "activities, and any extra clinical time that may be available. If the student is not "
                "available during this time period and has not made an absence request, the student "
                "will be cited for unprofessionalism and will risk failing the course."
            )
            ws.merge_range("C1:H2", note, f2)

            # Column widths & row height
            ws.set_column("A:A", 20)
            ws.set_column("B:B", 30)
            ws.set_column("C:G", 40)
            ws.set_column("H:H", 155)
            ws.set_row(0, 37.25)

            # Day names + dates
            date_idx = 0
            for row in start_rows:
                # day names (row, B..H). Using numeric write like your original
                for col_offset, day in enumerate(days, start=1):
                    ws.write(row, col_offset, day, f3)
                # dates directly beneath
                for col_offset in range(7):
                    if date_idx < len(dates):
                        ws.write(row + 1, col_offset + 1, dates[date_idx], f4)
                        date_idx += 1

            # Week labels
            for i, row in enumerate(week_label_rows):
                ws.write(f"A{row}", f"Week {i+1}", f3)

            # AM / PM labels
            for r in am_rows:
                ws.write(f"A{r}", "AM", f3)
            for r in pm_rows:
                ws.write(f"A{r}", "PM", f3)

            # Fill AM/PM blocks with Asynchronous Time (cols B..H)
            for i in range(NUM_WEEKS):
                am_r = 5 + i * BLOCK_HEIGHT  # keep your original indexing style
                pm_r = 6 + i * BLOCK_HEIGHT
                for col in range(1, 8):
                    ws.write(am_r, col, "Asynchronous Time", f5)
                    ws.write(pm_r, col, "Asynchronous Time", f5)

            # Separators (black bars)
            for sep in separators:
                ws.merge_range(f"A{sep}:H{sep}", "", f6)

            # Green filler rows
            for filler in green_fillers:
                for col in range(8):
                    ws.write(filler, col, " ", f7)

            # Assignment‑due rows
            for i, base in enumerate(green_fillers):
                ws.write(f"A{base}", "ASSIGNMENT DUE:", f8)
                for col in range(1, 8):
                    if col == 5:
                        ws.write(base - 1, col, "Ask for Feedback!", f8)
                    elif col == 7:
                        ws.write(base - 1, col, due_texts[i] if i < len(due_texts) else "", f8)
                    else:
                        ws.write(base - 1, col, " ", f8)

        wb.close()
        buf.seek(0)
        return buf

    def _cluster_blocks(marker_rows):
        marker_rows = sorted(marker_rows)
        if not marker_rows:
            return []
        blocks, curr = [], [marker_rows[0]]
        for r in marker_rows[1:]:
            if r == curr[-1] + 1:
                curr.append(r)
            else:
                blocks.append(curr)
                curr = [r]
        blocks.append(curr)
        return blocks
    
    # --- helpers ---------------------------------------------------------------
    
    _date_re = re.compile(r"^\s*\d{1,2}/\d{1,2}/\d{2,4}\s*$")
    
    def _is_date_like(v):
        """True for datetime/date, mm/dd/yyyy-ish strings, or Excel serial dates."""
        from datetime import date, datetime
        if isinstance(v, (date, datetime)):
            return True
        if isinstance(v, str) and _date_re.match(v):
            return True
        # Excel serial numbers (optional)
        if isinstance(v, (int, float)):
            try:
                from openpyxl.utils.datetime import from_excel
                _ = from_excel(v)  # just test convert
                return True
            except Exception:
                return False
        return False
    
    def _week_headers_by_column(ws, col_index=2):
        """
        Return sorted row numbers of the header *date* cells in a given column.
        NOTE: iter_rows yields tuples; grab the first element.
        """
        rows = []
        for row_tuple in ws.iter_rows(min_col=col_index, max_col=col_index, values_only=False):
            cell = row_tuple[0]
            if _is_date_like(cell.value):
                rows.append(cell.row)
        return sorted(rows)
    
    def _week_index_for_row(row_idx, header_rows):
        """
        Given a row and a sorted list of header date row indices, return 0-based week index.
        The content for a week lies strictly below its header row and above the next header row.
        """
        for i, hr in enumerate(header_rows):
            nxt = header_rows[i+1] if i+1 < len(header_rows) else 10**9
            if row_idx > hr and row_idx < nxt:
                return i
        return None
    
    def _row_label_for_bracket(site_name: str, ws_opd, row_idx: int) -> str:
        """Bracket label: sheet name except SUBSPECIALTY, which uses the text in column A for that row."""
        if site_name != "SUBSPECIALTY":
            return site_name
        val = ws_opd.cell(row=row_idx, column=1).value  # col A
        label = (str(val).strip() if val is not None else "") or "SUBSPECIALTY"
        return label
    
    # --- AM copy ---------------------------------------------------------------
    
    def assign_preceptors_all_weeks_am(opd_file, ms_file):
        """
        Copy AM assignments from OPD (B..H) to each student's tab.
        Week is determined by the nearest header *date* row above (column B),
        so SUBSPECIALTY extra rows still map to the correct week.
        Brackets show sheet name, except SUBSPECIALTY uses the column A designation.
        """
        opd_wb = load_workbook(opd_file, data_only=True)
        ms_wb  = load_workbook(ms_file)
    
        target_ms_rows = _seq(6, BLOCK_HEIGHT, NUM_WEEKS)  # [6,14,22,30,38]
    
        for site in opd_wb.sheetnames:
            ws_opd = opd_wb[site]
    
            # header date rows (use Monday column B)
            header_rows = _week_headers_by_column(ws_opd, col_index=2)
            if not header_rows:
                # No header dates found; skip this sheet to avoid wrong mapping
                continue
    
            # All AM-labeled rows
            am_rows = [c.row for c in ws_opd['A']
                       if isinstance(c.value, str) and re.match(r"^\s*AM\b", c.value, re.IGNORECASE)]
            if not am_rows:
                continue
    
            for r in am_rows:
                wk = _week_index_for_row(r, header_rows)
                if wk is None or wk >= NUM_WEEKS:
                    continue
                ms_row = target_ms_rows[wk]
    
                for col in range(2, 9):  # B..H
                    raw = ws_opd.cell(row=r, column=col).value
                    if not raw or "~" not in str(raw):
                        continue
                    pre, student = [s.strip() for s in str(raw).split("~", 1)]
                    if not student or student not in ms_wb.sheetnames:
                        continue
                    ws_ms = ms_wb[student]
                    label = _row_label_for_bracket(site, ws_opd, r)
                    ws_ms.cell(row=ms_row, column=col).value = f"{pre} - [{label}]"
    
        out = io.BytesIO()
        ms_wb.save(out)
        out.seek(0)
        return out
    
    # --- PM copy ---------------------------------------------------------------
    
    def assign_preceptors_all_weeks_pm(opd_file, ms_file):
        """
        Copy PM assignments; week rows [7,15,23,31,39] for 5 weeks.
        Uses header date rows (column B) to determine week boundaries.
        """
        opd_wb = load_workbook(opd_file, data_only=True)
        ms_wb  = load_workbook(ms_file)
    
        target_ms_rows = _seq(7, BLOCK_HEIGHT, NUM_WEEKS)  # [7,15,23,31,39]
    
        for site in opd_wb.sheetnames:
            ws_opd = opd_wb[site]
    
            header_rows = _week_headers_by_column(ws_opd, col_index=2)
            if not header_rows:
                continue
    
            pm_rows = [c.row for c in ws_opd['A']
                       if isinstance(c.value, str) and re.match(r"^\s*PM\b", c.value, re.IGNORECASE)]
            if not pm_rows:
                continue
    
            for r in pm_rows:
                wk = _week_index_for_row(r, header_rows)
                if wk is None or wk >= NUM_WEEKS:
                    continue
                ms_row = target_ms_rows[wk]
    
                for col in range(2, 9):  # B..H
                    raw = ws_opd.cell(row=r, column=col).value
                    if not raw or "~" not in str(raw):
                        continue
                    pre, student = [s.strip() for s in str(raw).split("~", 1)]
                    if not student or student not in ms_wb.sheetnames:
                        continue
                    ws_ms = ms_wb[student]
                    label = _row_label_for_bracket(site, ws_opd, r)
                    ws_ms.cell(row=ms_row, column=col).value = f"{pre} - [{label}]"
    
        out = io.BytesIO()
        ms_wb.save(out)
        out.seek(0)
        return out
    

    def detect_shift_conflicts(opd_file):
        """
        Finds students double-booked in the same WEEK/DAY/SHIFT across any sheet.
        Week boundaries are determined by the header date rows in column B, so
        SUBSPECIALTY (taller layout) works too. For SUBSPECIALTY, the 'sheet'
        label in occurrences uses column A's designation (e.g., 'AM - ENDO_HOPE').
        """
        from datetime import date, datetime
        from openpyxl.utils.datetime import from_excel
        wb = load_workbook(opd_file, data_only=True)
    
        days = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
        am_re = re.compile(r"^\s*AM\b", re.IGNORECASE)
        pm_re = re.compile(r"^\s*PM\b", re.IGNORECASE)
    
        # ---- helpers (safe to keep local even if you already defined similar ones) ----
        _date_re = re.compile(r"^\s*\d{1,2}/\d{1,2}/\d{2,4}\s*$")
        def _is_date_like(v):
            if isinstance(v, (date, datetime)):
                return True
            if isinstance(v, str) and _date_re.match(v):
                return True
            if isinstance(v, (int, float)):
                try:
                    _ = from_excel(v)
                    return True
                except Exception:
                    return False
            return False
    
        def _week_headers_by_column(ws, col_index=2):
            rows = []
            for row_tuple in ws.iter_rows(min_col=col_index, max_col=col_index, values_only=False):
                cell = row_tuple[0]
                if _is_date_like(cell.value):
                    rows.append(cell.row)
            return sorted(rows)
    
        def _week_index_for_row(row_idx, header_rows):
            # return 0-based week index based on nearest header date row above
            for i, hr in enumerate(header_rows):
                nxt = header_rows[i+1] if i+1 < len(header_rows) else 10**9
                if row_idx > hr and row_idx < nxt:
                    return i
            return None
    
        def _label_for_occurrence(sheet_name, ws, row_idx):
            # Use sheet name normally; for SUBSPECIALTY use the designation in column A
            if sheet_name != "SUBSPECIALTY":
                return sheet_name
            v = ws.cell(row=row_idx, column=1).value
            return (str(v).strip() if v else "SUBSPECIALTY")
    
        # (week, day_name, shift, student) -> list[(label, coord)]
        from collections import defaultdict
        occurrences = defaultdict(list)
    
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            header_rows = _week_headers_by_column(ws, col_index=2)
            if not header_rows:
                continue
    
            for shift, rx in (("AM", am_re), ("PM", pm_re)):
                # find all AM/PM marker rows in column A
                mark_rows = [c.row for c in ws["A"] if isinstance(c.value, str) and rx.match(c.value)]
                for r in mark_rows:
                    wk = _week_index_for_row(r, header_rows)
                    if wk is None or wk >= NUM_WEEKS:
                        continue
                    # scan B..H for Monday..Sunday
                    for day_idx in range(7):
                        col = 2 + day_idx
                        raw = ws.cell(row=r, column=col).value
                        if not raw:
                            continue
                        text = str(raw)
                        if "~" not in text:
                            continue
                        pre, student = [s.strip() for s in text.split("~", 1)]
                        if not student:
                            continue
                        coord = ws.cell(row=r, column=col).coordinate
                        label = _label_for_occurrence(sheet_name, ws, r)
                        occurrences[(wk+1, days[day_idx], shift, student)].append((label, coord))
    
        # build conflict list in your existing shape
        conflicts = []
        for (week, day_name, shift, student), occ in occurrences.items():
            if len(occ) > 1:
                conflicts.append({
                    "student": student,
                    "week": week,
                    "day": day_name,
                    "shift": shift,
                    "occurrences": occ,  # list of (label, A1coord)
                })
        return conflicts


    # ───────── Load OPD & Rotation Schedule ─────────
    df_opd = load_workbook_df("Upload OPD.xlsx file", ["xlsx"], key="opd_main")
    df_rot = load_workbook_df("Upload Rotation Schedule (.xlsx or .csv)", ["xlsx", "csv"], key="rot_main")

    # ───────── Check for duplicates ─────────
    if df_opd is not None:
        conflicts = detect_shift_conflicts(st.session_state["opd_main_file"])
        if conflicts:
            for c in conflicts:
                occ_str = "; ".join(f"{sheet}@{coord}" for sheet, coord in c["occurrences"])
                st.warning(
                    f"⚠️ Week {c['week']} {c['day']} {c['shift']}: "
                    f"{c['student']} double‑booked ({occ_str})"
                )
        else:
            st.success("No AM/PM shift conflicts detected.")

    # ───────── Build, Assign & Download ─────────
    if df_opd is not None and df_rot is not None:
        df_rot["start_date"] = pd.to_datetime(df_rot["start_date"])
        monday = df_rot["start_date"].min()
        monday = monday - pd.Timedelta(days=monday.weekday())  # normalize to Monday
        dates = pd.date_range(start=monday, periods=NUM_WEEKS * 7, freq="D").tolist()

        students = df_rot["legal_name"].dropna().unique().tolist()

        if st.button("Create & Download Fully‑Populated MS_Schedule"):
            blank_buf = create_ms_schedule_template(students, dates)
            am_buf = assign_preceptors_all_weeks_am(
                opd_file=st.session_state["opd_main_file"], ms_file=blank_buf
            )
            full_buf = assign_preceptors_all_weeks_pm(
                opd_file=st.session_state["opd_main_file"], ms_file=am_buf
            )
            st.download_button(
                "Download MS_Schedule.xlsx",
                data=full_buf.getvalue(),
                file_name="MS_Schedule.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
    else:
        st.info("Please upload both OPD.xlsx and the rotation schedule above to proceed.")
