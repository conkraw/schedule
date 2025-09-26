import io
import streamlit as st
import pandas as pd
import numpy as np 
import re
import xlsxwriter
import random
from openpyxl import load_workbook # Ensure load_workbook is imported
import io, zipfile
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from datetime import timedelta
from xlsxwriter import Workbook as Workbook
from collections import defaultdict
from datetime import datetime, timedelta
from collections import Counter
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED
import re
from collections import defaultdict
from copy import copy
from openpyxl import load_workbook, Workbook
from openpyxl.utils import get_column_letter
from openpyxl.cell.cell import MergedCell
from openpyxl.utils import column_index_from_string
import math

st.set_page_config(page_title="Batch Preceptor → REDCap Import", layout="wide")
st.title("Batch Preceptor → REDCap Import Generator")

# ─── Sidebar mode selector ─────────────────────────────────────────────────────
#mode = st.sidebar.radio("What do you want to do?",("Instructions", "Format OPD + Summary (4-sheet, 5-week)", "Create Student Schedule","OPD Check","Create Individual Schedules"))
mode = st.sidebar.radio("What do you want to do?",("Format OPD + Summary (4-sheet, 5-week)", "Create Student Schedule","Create Individual Schedules"))
# ─── Sidebar mode selector ─────────────────────────────────────────────────────

if mode == "OPD Check":
    DAYS = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
    
    def detect_am_pm_blocks(df):
        """
        Scan down column A and group runs whose cell text begins with 'AM' or 'PM',
        ignoring everything else.
        Returns a list of (label, start_row, end_row) tuples.
        """
        runs = []
        current_label = None
        run_start = None
        prev_row = None
    
        # Excel row 6 is df index 5
        for idx in range(5, len(df)):
            raw = df.iat[idx, 0]
            if isinstance(raw, str):
                ru = raw.upper()
                if ru.startswith("AM"):
                    label = "AM"
                elif ru.startswith("PM"):
                    label = "PM"
                else:
                    continue
            else:
                continue
    
            row = idx + 1  # convert back to 1‑based Excel row
            if current_label is None:
                # start new run
                current_label = label
                run_start = row
            elif label != current_label or row != prev_row + 1:
                # close out previous run
                runs.append((current_label, run_start, prev_row))
                current_label = label
                run_start = row
    
            prev_row = row
    
        # finish last run
        if current_label is not None:
            runs.append((current_label, run_start, prev_row))
    
        return runs
    
    st.title("OPD Preceptor GI‑al Check")
    
    baseline_file = st.file_uploader(
        "1) Upload Latest Updated OPD", 
        type=["xlsx"], key="baseline"
    )
    assigned_file = st.file_uploader(
        "2) Upload OPD with Student Assignments", 
        type=["xlsx"], key="assigned"
    )
    
    if baseline_file and assigned_file:
        SHEETS = [
            'HOPE_DRIVE','ETOWN','NYES','COMPLEX','WARD A',
            'PSHCH_NURSERY','HAMPDEN_NURSERY','SJR_HOSP','AAC',
            'AHOLOUKPE','ADOLMED'
        ]
    
        # Read all relevant sheets at once
        base_sheets = pd.read_excel(baseline_file, sheet_name=SHEETS, header=None)
        assn_sheets = pd.read_excel(assigned_file, sheet_name=SHEETS, header=None)

        # 1) Remove student suffix from every baseline sheet
        for sheet_name, df in base_sheets.items():
            for col in df.columns[1:]:
                df[col] = (
                    df[col]
                      .where(df[col].notna(), np.nan)      # keep NaN
                      .astype(str)
                      .str.partition('~')[0]               # take text before "~"
                      .replace({'nan': np.nan})            # restore NaN
                )
            base_sheets[sheet_name] = df
    
        results = {}
        
        for sheet in SHEETS:
            df_base = base_sheets[sheet]
            df_assn = assn_sheets[sheet]
        
            runs = detect_am_pm_blocks(df_base)
            week_pairs = [(runs[i], runs[i+1]) for i in range(0, len(runs), 2)]
        
            base_map = {}    # (w,period,day,pre) → (cell, student)
            assn_map = {}
        
            for w_idx, ((am_lbl, am_s, am_e), (pm_lbl, pm_s, pm_e)) in enumerate(week_pairs, start=1):
                for period, (lbl, start, end) in [('AM',(am_lbl,am_s,am_e)), ('PM',(pm_lbl,pm_s,pm_e))]:
                    for col, day in enumerate(DAYS, start=1):
                        for row in range(start, end+1):
                            cell = f"{chr(ord('A')+col)}{row}"
        
                            # baseline
                            vb = df_base.iat[row-1, col]
                            if pd.notna(vb) and isinstance(vb, str):
                                parts = str(vb).split('~', 1)
                                pre = parts[0].strip()
                                stu = parts[1].strip() if len(parts)==2 else None
                                base_map.setdefault((w_idx,period,day,pre), (cell, stu))
        
                            # assigned
                            va = df_assn.iat[row-1, col]
                            if pd.notna(va) and isinstance(va, str):
                                parts = str(va).split('~', 1)
                                pre = parts[0].strip()
                                stu = parts[1].strip() if len(parts)==2 else None
                                assn_map.setdefault((w_idx,period,day,pre), (cell, stu))
        
            dropped = []
            added   = []
        
            # drops: in base not in assigned
            for key, (cell, stu) in base_map.items():
                if key not in assn_map:
                    w,p,d,pre = key
                    dropped.append((w,p,d,pre,cell,stu))
        
            # adds: in assigned not in base
            for key, (cell, stu) in assn_map.items():
                if key not in base_map:
                    w,p,d,pre = key
                    added.append((w,p,d,pre,cell,stu))
        
            # sort exactly as before...
            dropped_sorted = sorted(
                dropped,
                key=lambda x: (x[0], {'AM':0,'PM':1}[x[1]], DAYS.index(x[2]), x[4])
            )
            added_sorted = sorted(
                added,
                key=lambda x: (x[0], {'AM':0,'PM':1}[x[1]], DAYS.index(x[2]), x[4])
            )
        
            results[sheet] = {"dropped": dropped_sorted, "added": added_sorted}



        
        #
    
    doc = Document()
    doc.add_heading('Change Report', level=1)
    
    for sheet, change in results.items():
        doc.add_heading(sheet, level=2)
    
        # build week→day map (collect both AM & PM under each day)
        week_map = defaultdict(lambda: defaultdict(lambda: {'dropped': [], 'added': []}))
        for w,p,d,pre,cell,stu in change['dropped']:
            week_map[w][d]['dropped'].append((p, pre, cell, stu))
        for w,p,d,pre,cell,stu in change['added']:
            week_map[w][d]['added'].append((p, pre, cell, stu))
    
        # emit in week order
        for w in sorted(week_map):
            doc.add_heading(f'Week {w}', level=3)
    
            for day in DAYS:
                slot = week_map[w].get(day)
                if not slot or (not slot['dropped'] and not slot['added']):
                    continue
    
                doc.add_heading(day, level=4)
    
                # DROPS
                for p, pre, cell, stu in slot['dropped']:
                    line = f"- Dropped: {pre} — was at {cell}"
                    if stu:
                        line += f"  (Student impacted: {stu})"
                    doc.add_paragraph(line, style='List Bullet')
    
                # ADDS
                for p, pre, cell, stu in slot['added']:
                    line = f"- Added: {pre} — now at {cell}"
                    if stu:
                        line += f"  (Student impacted: {stu})"
                    doc.add_paragraph(line, style='List Bullet')
    
            doc.add_paragraph()  # blank line between weeks




    # Save to in-memory buffer
    word_file = io.BytesIO()
    doc.save(word_file)
    word_file.seek(0)
    
    # Download button
    st.download_button(
        label="📄 Download Word Report",
        data=word_file,
        file_name="change_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


elif mode == "Instructions":
    d = st.text_input('Start date (m/d/yyyy)')
    if d:
        try:
            s = datetime.strptime(d, '%m/%d/%Y')
            e = s + timedelta(days=41)
            st.write(f"{s:%B %d, %Y} → {e:%B %d, %Y}")
    
            st.write('Please go to https://login.qgenda.com/')
    
            # Display on‐screen instructions
            st.markdown(f"""Download four files and create reports based on **{s:%B %d, %Y}** → **{e:%B %d, %Y}**.""")
            st.write("Download instructions here:")


        except ValueError:
            st.error('Invalid format – use m/d/yyyy (e.g. 7/6/2021)')

        # --- Generate a Word document with the same instructions ---
        doc = Document()
        doc.add_heading('Qgenda Report Instructions', level=1)
        doc.styles['Normal'].font.size = Pt(8)
        
        # Bold the date range
        p = doc.add_paragraph()
        p.add_run('Date range: ')
        run_start = p.add_run(f'{s:%B %d, %Y}')
        run_start.bold = True
        p.add_run(' → ')
        run_end = p.add_run(f'{e:%B %d, %Y}')
        run_end.bold = True
        
        
        doc.add_paragraph('1. Go to https://login.qgenda.com/')
        
        # Helper to add each report block
        def add_report(title, steps):
            doc.add_heading(title, level=2)
            for step in steps:
                # If the step contains dates, bold them
                if 'Enter Start Date:' in step:
                    p = doc.add_paragraph(style='List Bullet')
                    prefix, dates = step.split(':', 1)
                    p.add_run(prefix + ': ')
                    # split the two dates on " and End Date:"
                    start_part, end_part = dates.strip().split(' and End Date:')
                    r1 = p.add_run(start_part.strip())
                    r1.bold = True
                    p.add_run(' and End Date: ')
                    r2 = p.add_run(end_part.strip())
                    r2.bold = True
                else:
                    doc.add_paragraph(step, style='List Bullet')
        
        add_report(
            'Report 1 – Penn State Health Hershey Medical Center - Academic General Pediatrics',
            [
                'Click Penn State Health Hershey Medical Center - Academic General Pediatrics → Schedule → Reports',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Short Name',
                'Click Run Report'
            ]
        )
        add_report(
            "Report 2 – Penn State Health Children's Hospital – Hospitalists",
            [
                'Click Penn State Health Children\'s Hospital → Schedule → Reports',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )
        add_report(
            'Report 3 – Department of Pediatrics (Admin - Adolescent Med)',
            [
                'Click Department of Pediatrics → Schedule → Reports',
                'Select Admin - Adolescent Med in top-right corner',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )
        add_report(
            'Report 4 – Department of Pediatrics (Complex Care)',
            [
                'Click Department of Pediatrics → Schedule → Reports',
                'Select Complex Care in top-right corner',
                'Set Report Type to Calendar by Task',
                'Set Format to Excel',
                f'Enter Start Date: {s:%m/%d/%Y} and End Date: {e:%m/%d/%Y}',
                'Ensure Calendar starts on Monday',
                'Show Staff by Last Name, First Name',
                'Show Tasks by Long Name',
                'Click Run Report'
            ]
        )


        # Save to bytes and offer download
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(label="📄 Download Instructions (Word)",data=buf.getvalue(),file_name="Qgenda_Report_Instructions.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif mode == "Format OPD + Summary (4-sheet, 5-week)":
    import math
    import re
    import io
    import zipfile
    import pandas as pd
    import streamlit as st

    # ─────────────────────────────────────────────────────────────────────────────
    # Parameterized site registry
    # Add new sheets by appending entries to SITE_CONFIGS below — no other code edits.
    # Each entry defines: Excel sheet name, title, type ("hope_drive" with acutes or
    # generic "continuity"), the labels as they appear in QGenda (aliases allowed),
    # and the REDCap column prefixes to write into.
    # ─────────────────────────────────────────────────────────────────────────────
    SITE_CONFIGS = {
        # Special template with AM/PM acutes + continuity and weekend handling
        "HOPE_DRIVE": {
            "title": "Hope Drive",
            "type": "hope_drive",
            "keywords": ["academic general pediatrics"],
            "designations": {
                # aliases seen in QGenda → (canonical_key, redcap_prefix)
                # Continuity
                "hope drive am continuity": ("hd_am", "hd_am_"),
                "hope drive pm continuity": ("hd_pm", "hd_pm_"),
                # Weekday acutes
                "hope drive am acute precept": ("hd_am_acute", "hd_am_acute_"),
                "hope drive pm acute precept": ("hd_pm_acute", "hd_pm_acute_"),
                # Weekend
                "hope drive weekend acute 1": ("hd_wknd_acute_1", "hd_wknd_acute_1_"),
                "hope drive weekend acute 2": ("hd_wknd_acute_2", "hd_wknd_acute_2_"),
                "hope drive weekend continuity": ("hd_wknd_am", "hd_wknd_am_"),
            },
        },
        # Generic continuity-only templates (10 AM rows, 10 PM rows)
        "ETOWN": {
            "title": "Elizabethtown",
            "type": "continuity",
            "keywords": ["academic general pediatrics"],
            "am": {"aliases": ["etown am continuity"], "prefix": "etown_am_"},
            "pm": {"aliases": ["etown pm continuity"], "prefix": "etown_pm_"},
        },
        "NYES": {
            "title": "Nyes Road",
            "type": "continuity",
            "keywords": ["academic general pediatrics"],
            "am": {"aliases": ["nyes rd am continuity"], "prefix": "nyes_am_"},
            "pm": {"aliases": ["nyes rd pm continuity"], "prefix": "nyes_pm_"},
        },
        "LANCASTER": {
            "title": "Lancaster",
            "type": "continuity",
            "keywords": ["academic general pediatrics"],
            "am": {"aliases": ["lancaster am"], "prefix": "lancaster_am_"},
            "pm": {"aliases": ["lancaster pm"], "prefix": "lancaster_pm_"},
        },
        "COMPLEX": {
            "title": "Complex Care",
            "type": "continuity",
            "keywords": ["complex care"],
            "am": {"aliases": ["hope drive clinic am"], "prefix": "complex_am_"},
            "pm": {"aliases": ["hope drive clinic pm"], "prefix": "complex_pm_"},
        },
        "NEURO_HOPE": {
            "title": "Neurology Hope Drive",
            "type": "continuity",
            "keywords": ["neurology"],
            "am": {"aliases": ["35 hope drive am"], "prefix": "neurohd_am_"},
            "pm": {"aliases": ["35 hope drive pm"], "prefix": "neurohd_pm_"},
        },
        "NEURO_LANCASTER": {
            "title": "Neurology Lancaster",
            "type": "continuity",
            "keywords": ["neurology"],
            "am": {"aliases": ["lanc spec am"], "prefix": "neurolanc_am_"},
            "pm": {"aliases": ["lanc spec pm"], "prefix": "neurolanc_pm_"},
        },
        "NEURO_READING": {
            "title": "Neurology Reading",
            "type": "continuity",
            "keywords": ["neurology"],
            "am": {"aliases": ["reading am"], "prefix": "neuroread_am_"},
            "pm": {"aliases": ["reading pm"], "prefix": "neuroread_pm_"},
        },

        "ENDO_HOPE": {
            "title": "Endocrinology Hope Drive",
            "type": "continuity",
            "keywords": ["endo"],
            "am": {"aliases": ["hope drive am"], "prefix": "endohope_am_"},
            "pm": {"aliases": ["hope drive pm"], "prefix": "endohope_pm_"},
        },

        "ENDO_CAMPHILL": {
            "title": "Endocrinology Camp Hill",
            "type": "continuity",
            "keywords": ["endo"],
            "am": {"aliases": ["camp hill clinic am"], "prefix": "endocamp_am_"},
            "pm": {"aliases": ["camp hill clinic pm"], "prefix": "endocamp_pm_"},
        },

        "ENDO_LANCASTER": {
            "title": "Endocrinology Lancaster",
            "type": "continuity",
            "keywords": ["endo"],
            "am": {"aliases": ["lanc ped center am"], "prefix": "endolanc_am_"},
            "pm": {"aliases": ["lanc ped center pm"], "prefix": "endolanc_pm_"},
        },

        "ENDO_SJR": {
            "title": "Endocrinology SJR",
            "type": "continuity",
            "keywords": ["endo"],
            "am": {"aliases": ["st. joseph am"], "prefix": "endosjr_am_"},
            "pm": {"aliases": ["st. joseph pm"], "prefix": "endosjr_pm_"},
        },
        
        # Example to add later (just copy and adjust aliases/prefixes):
        #"ADOLMED": {
        #     "title": "Adolescent Medicine",
        #     "type": "continuity",
        #     "keywords": ["adol med", "adolescent"],
        #     "am": {"aliases": ["briarcrest clinic am"], "prefix": "adol_med_am_"},
        #     "pm": {"aliases": ["briarcrest clinic pm"], "prefix": "adol_med_pm_"},
        # },
    }

    # Build required keywords dynamically from the included sites
    REQUIRED_KEYWORDS = sorted({kw for cfg in SITE_CONFIGS.values() for kw in cfg.get("keywords", [])})

    # ─── Inputs ──────────────────────────────────────────────────────────────────
    schedule_files = st.file_uploader(
        "1) Upload one or more QGenda calendar Excel(s)",
        type=["xlsx", "xls"],
        accept_multiple_files=True,
    )

    student_file = st.file_uploader(
        "2) Upload Redcap Rotation list CSV (must have a 'legal_name' column)",
        type=["csv"],
    )

    record_id = "peds_clerkship"

    if not schedule_files or not student_file or not record_id:
        st.info("Please upload schedule Excel(s), student CSV")
        st.stop()

    # ─── Parse files ─────────────────────────────────────────────────────────────
    date_pat = re.compile(r"^[A-Za-z]+ \d{1,2}, \d{4}$")

    # Build designation→prefix map from SITE_CONFIGS (supports aliases)
    designation_map = {}
    for sheet_name, cfg in SITE_CONFIGS.items():
        if cfg["type"] == "continuity":
            for alias in cfg["am"]["aliases"]:
                designation_map[alias.lower()] = cfg["am"]["prefix"]
            for alias in cfg["pm"]["aliases"]:
                designation_map[alias.lower()] = cfg["pm"]["prefix"]
        elif cfg["type"] == "hope_drive":
            for alias, (_canon, prefix) in cfg["designations"].items():
                designation_map[alias.lower()] = prefix

    # Minimum headcount rules (per designation key as it appears in QGenda)
    min_required = {
        "hope drive am acute precept": 2,
        "hope drive pm acute precept": 2,
    }

    # Aggregate schedule assignments by date
    assignments_by_date = {}
    found_keywords = set()

    for file in schedule_files:
        try:
            df = pd.read_excel(file, header=None, dtype=str)
        except Exception as e:
            st.error(f"Error reading {file.name}: {e}")
            continue

        # Detect keywords present in the file (for friendly warning if any missing)
        flat_vals = df.astype(str).apply(lambda s: s.str.lower()).values.flatten().tolist()
        for kw in REQUIRED_KEYWORDS:
            if any(kw in v for v in flat_vals):
                found_keywords.add(kw)

        # Find date columns and collect provider rows under each date
        date_positions = []
        for r in range(df.shape[0]):
            for c in range(df.shape[1]):
                val = str(df.iat[r, c]).replace("\xa0", " ").strip()
                if date_pat.match(val):
                    try:
                        d = pd.to_datetime(val).date()
                        date_positions.append((d, r, c))
                    except Exception:
                        pass

        # Deduplicate to the topmost row per date value
        unique = {}
        for d, r, c in date_positions:
            if d not in unique or r < unique[d][0]:
                unique[d] = (r, c)

        for d, (row0, col0) in unique.items():
            # Pre-seed with all known designations (so missing ones exist as empty lists)
            grp = assignments_by_date.setdefault(d, {des: [] for des in designation_map})
            for r in range(row0 + 1, df.shape[0]):
                raw = str(df.iat[r, col0]).replace("\xa0", " ").strip()
                if raw == "":
                    break
                if date_pat.match(raw):
                    break
                desc = raw.lower()
                prov = str(df.iat[r, col0 + 1]).strip()
                if desc in grp and prov:
                    grp[desc].append(prov)

    # Friendly notice if expected site keywords are missing across uploads
    missing_keywords = [kw for kw in REQUIRED_KEYWORDS if kw not in found_keywords]
    if missing_keywords:
        st.warning("These site keywords weren’t detected in your uploads: " + ", ".join(missing_keywords))

    students_df = pd.read_csv(student_file, dtype=str)
    legal_names = students_df["legal_name"].dropna().tolist()

    # ─── Provider filter UI ──────────────────────────────────────────────────────
    # Build list of all providers seen; let user whitelist a subset before we fill
    all_providers = sorted({
        p.strip() for day in assignments_by_date.values() for provs in day.values() for p in provs if isinstance(p, str) and p.strip()
    })

    if "provider_filter" not in st.session_state:
        st.session_state["provider_filter"] = []

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Select All Providers"):
            st.session_state["provider_filter"] = all_providers
    with col2:
        if st.button("Clear Providers"):
            st.session_state["provider_filter"] = []

    allowed_providers = st.multiselect(
        "Limit providers included in OPD",
        options=all_providers,
        key="provider_filter",
        help="Only selected providers will be written into the OPD sheets (others will be left blank).",
    )

    # ─── Build redcap_row (5 weeks fixed) ────────────────────────────────────────
    redcap_row = {"record_id": record_id}
    sorted_dates = sorted(assignments_by_date.keys())
    NUM_WEEKS = 5
    NUM_DAYS = NUM_WEEKS * 7

    for idx, date in enumerate(sorted_dates[:NUM_DAYS], start=1):
        redcap_row[f"hd_day_date{idx}"] = date
        suffix = f"d{idx}_"
        # Expand designation_map into actual column prefixes with current day suffix
        des_map = {des: ([prefix + suffix]) for des, prefix in designation_map.items()}
        for des, provs in assignments_by_date[date].items():
            # Filter by allowed providers
            filtered = [p for p in provs if p in allowed_providers]
            req = min_required.get(des, len(filtered))
            if filtered:
                while len(filtered) < req:
                    filtered.append(filtered[0])
            for i, name in enumerate(filtered, start=1):
                for prefix in des_map[des]:
                    redcap_row[f"{prefix}{i}"] = name

    # Students list stored but not appended into cells in this 4-sheet build
    for i, name in enumerate(legal_names, start=1):
        redcap_row[f"s{i}"] = name

    out_df = pd.DataFrame([redcap_row])
    for c in list(out_df.columns):
        if c.startswith("hd_day_date"):
            out_df[c] = pd.to_datetime(out_df[c]).dt.strftime("%m-%d-%Y")

    csv_full = out_df.to_csv(index=False).encode("utf-8")

    # ─── Workbook ────────────────────────────────────────────────────────────────
    def generate_opd_workbook(full_df: pd.DataFrame) -> bytes:
        import xlsxwriter
        output = io.BytesIO()
        workbook = xlsxwriter.Workbook(output, {"in_memory": True})

        # formats (subset from your originals)
        format1 = workbook.add_format({"font_size": 18, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#FEFFCC", "border": 1})
        format4 = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#8ccf6f", "border": 1})
        format4a = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#9fc5e8", "border": 1})
        format5a = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#d0e9ff", "border": 1})
        format3 = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#FFC7CE", "border": 1})
        format2 = workbook.add_format({"bg_color": "black"})
        format_date = workbook.add_format({"num_format": "m/d/yyyy", "font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#FFC7CE", "border": 1})
        format_label = workbook.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter", "font_color": "black", "bg_color": "#FFC7CE", "border": 1})
        merge_format = workbook.add_format({"bold": 1, "align": "center", "valign": "vcenter", "text_wrap": True, "font_color": "red", "bg_color": "#FEFFCC", "border": 1})

        # Acute highlight formats (HOPE_DRIVE)
        format_am_acute = workbook.add_format({
            "font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
            "font_color": "black", "bg_color": "#8ccf6f", "border": 1  # green
        })
        format_pm_acute = workbook.add_format({
            "font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
            "font_color": "white", "bg_color": "#1f4e79", "border": 1  # dark blue
        })

        # Create only the configured sheets (order is registry order)
        worksheet_names = list(SITE_CONFIGS.keys())
        site_list = [SITE_CONFIGS[name]["title"] for name in worksheet_names]
        sheets = {name: workbook.add_worksheet(name) for name in worksheet_names}
        for ws, site in zip(sheets.values(), site_list):
            ws.write(0, 0, "Site:", format1)
            ws.write(0, 1, site, format1)

        # Block math shared by all
        BLOCK_HEIGHT = 24
        BLOCK_STARTS = [6 + i * BLOCK_HEIGHT for i in range(NUM_WEEKS)]
        HDR_STARTS = [2 + i * BLOCK_HEIGHT for i in range(NUM_WEEKS)]

        # HOPE_DRIVE styling (if present)
        if "HOPE_DRIVE" in sheets and SITE_CONFIGS["HOPE_DRIVE"]["type"] == "hope_drive":
            hd = sheets["HOPE_DRIVE"]
            ACUTE_COUNT = 2
            CONTINUITY_COUNT = 8
            AM_COUNT = ACUTE_COUNT + CONTINUITY_COUNT  # 10
            PM_COUNT = AM_COUNT

            for start in BLOCK_STARTS:
                # Priority 1: acutes (include column A for labels)
                hd.conditional_format(f"A{start}:H{start+1}", {"type": "no_errors", "format": format_am_acute})
                hd.conditional_format(f"A{start+10}:H{start+11}", {"type": "no_errors", "format": format_pm_acute})
                # Priority 2: header tints
                hd.conditional_format(f"B{start}:H{start}", {"type": "no_errors", "format": format4})
                hd.conditional_format(f"B{start+10}:H{start+10}", {"type": "no_errors", "format": format4a})
                # Priority 3: broad AM/PM fills
                hd.conditional_format(f"A{start}:H{start+9}", {"type": "no_errors", "format": format1})
                hd.conditional_format(f"A{start+10}:H{start+19}", {"type": "no_errors", "format": format5a})
                # Column A labels for AM/PM sections
                zero_row = start - 1
                for i in range(AM_COUNT):
                    label = "AM - ACUTES" if i < ACUTE_COUNT else "AM - Continuity"
                    hd.write(zero_row + i, 0, label, format5a)
                for i in range(PM_COUNT):
                    label = "PM - ACUTES" if i < ACUTE_COUNT else "PM - Continuity"
                    hd.write(zero_row + AM_COUNT + i, 0, label, format5a)

        # Continuity-only sheets (auto for anything in registry with type=continuity)
        for name, cfg in SITE_CONFIGS.items():
            if cfg["type"] != "continuity":
                continue
            ws = sheets[name]
            for start in BLOCK_STARTS:
                zero_row = start - 1
                # AM/PM column A labels
                for i in range(10):
                    ws.write(zero_row + i, 0, "AM", format5a)
                for i in range(10):
                    ws.write(zero_row + 10 + i, 0, "PM", format5a)
                # Background bands + headers
                ws.conditional_format(f"A{start}:H{start+9}", {"type": "no_errors", "format": format1})
                ws.conditional_format(f"A{start+10}:H{start+19}", {"type": "no_errors", "format": format5a})
                ws.conditional_format(f"B{start}:H{start}", {"type": "no_errors", "format": format4})
                ws.conditional_format(f"B{start+10}:H{start+10}", {"type": "no_errors", "format": format4a})

        # Headers, dates, bars, CRTS — universal for all sheets
        date_cols = [f"hd_day_date{i}" for i in range(1, NUM_DAYS + 1)]
        dates = pd.to_datetime(full_df[date_cols].iloc[0], errors="coerce").tolist()
        weeks = [dates[i * 7 : (i + 1) * 7] for i in range(NUM_WEEKS)]
        days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

        for ws in workbook.worksheets():
            ws.set_zoom(80)
            ws.set_column("A:A", 10)
            ws.set_column("B:H", 65)
            ws.set_row(0, 37.25)
            for idx, start in enumerate(HDR_STARTS):
                for c, dname in enumerate(days):
                    ws.write(start, 1 + c, dname, format3)
                for c, val in enumerate(weeks[idx]):
                    if pd.notna(val):
                        ws.write(start + 1, 1 + c, val, format_date)
                ws.write_formula(f"A{start}", '""', format_label)
                ws.conditional_format(f"A{start+3}:H{start+3}", {"type": "no_errors", "format": format_label})
            # black bars and top two white rows per block (all sheets)
            for row in range(2, 2 + BLOCK_HEIGHT * NUM_WEEKS, BLOCK_HEIGHT):
                ws.merge_range(f"A{row}:H{row}", " ", format2)
            for i in range(NUM_WEEKS):
                ws.write(f"A{3 + i*BLOCK_HEIGHT}", "", format_date)
                ws.write(f"A{4 + i*BLOCK_HEIGHT}", "", format_date)

            text1 = (
                "Students are to alert their preceptors when they have a Clinical "
                "Reasoning Teaching Session (CRTS). Please allow the students to "
                "leave ~15 minutes prior to the start of their session so they can be prepared."
            )
            ws.merge_range("C1:F1", text1, merge_format)
            ws.write("G1", "", merge_format)
            ws.write("H1", "", merge_format)

        workbook.close()
        output.seek(0)
        return output.read()

    # ─── CSV→Excel mappings built from registry ─────────────────────────────────
    def build_mappings(NUM_WEEKS: int) -> list:
        mappings = []
        excel_column_letters = ["B", "C", "D", "E", "F", "G", "H"]
        cont_row_defs = {"AM": 6, "PM": 16}
        hd_row_defs = {"AM": {"acute_start": 6, "cont_start": 8}, "PM": {"acute_start": 16, "cont_start": 18}}

        # HOPE_DRIVE mappings
        if "HOPE_DRIVE" in SITE_CONFIGS and SITE_CONFIGS["HOPE_DRIVE"]["type"] == "hope_drive":
            for week_idx in range(1, NUM_WEEKS + 1):
                week_base = (week_idx - 1) * 24
                day_offset = (week_idx - 1) * 7
                for day_idx, col in enumerate(excel_column_letters, start=1):
                    is_weekday = day_idx <= 5
                    day_num = day_idx + day_offset
                    # AM
                    if is_weekday:
                        for prov in range(1, 2 + 1):
                            row = week_base + hd_row_defs["AM"]["acute_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_am_acute_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                        for prov in range(1, 8 + 1):
                            row = week_base + hd_row_defs["AM"]["cont_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_am_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                    else:
                        for acute_type in (1, 2):
                            row = week_base + hd_row_defs["AM"]["acute_start"] + (acute_type - 1)
                            mappings.append({"csv_column": f"hd_wknd_acute_{acute_type}_d{day_num}_1", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                        for prov in range(1, 8 + 1):
                            row = week_base + hd_row_defs["AM"]["cont_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_wknd_am_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                    # PM
                    if is_weekday:
                        for prov in range(1, 2 + 1):
                            row = week_base + hd_row_defs["PM"]["acute_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_pm_acute_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                        for prov in range(1, 8 + 1):
                            row = week_base + hd_row_defs["PM"]["cont_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_pm_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                    else:
                        for acute_type in (1, 2):
                            row = week_base + hd_row_defs["PM"]["acute_start"] + (acute_type - 1)
                            mappings.append({"csv_column": f"hd_wknd_pm_acute_{acute_type}_d{day_num}_1", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})
                        for prov in range(1, 8 + 1):
                            row = week_base + hd_row_defs["PM"]["cont_start"] + (prov - 1)
                            mappings.append({"csv_column": f"hd_wknd_pm_d{day_num}_{prov}", "excel_sheet": "HOPE_DRIVE", "excel_cell": f"{col}{row}"})

        # Continuity-only sheets from registry (auto)
        for sheet_name, cfg in SITE_CONFIGS.items():
            if cfg["type"] != "continuity":
                continue
            am_prefix = cfg["am"]["prefix"]
            pm_prefix = cfg["pm"]["prefix"]
            for week_idx in range(1, NUM_WEEKS + 1):
                week_base = (week_idx - 1) * 24
                day_offset = (week_idx - 1) * 7
                for day_idx, col in enumerate(excel_column_letters, start=1):
                    day_num = day_idx + day_offset
                    for prov in range(1, 10 + 1):
                        row = week_base + cont_row_defs["AM"] + (prov - 1)
                        mappings.append({"csv_column": f"{am_prefix}d{day_num}_{prov}", "excel_sheet": sheet_name, "excel_cell": f"{col}{row}"})
                    for prov in range(1, 10 + 1):
                        row = week_base + cont_row_defs["PM"] + (prov - 1)
                        mappings.append({"csv_column": f"{pm_prefix}d{day_num}_{prov}", "excel_sheet": sheet_name, "excel_cell": f"{col}{row}"})
        return mappings

    def update_excel_from_csv(excel_template_bytes: bytes, csv_data_bytes: bytes, mappings: list) -> bytes | None:
        from openpyxl import load_workbook
        from openpyxl.styles import Alignment
        try:
            df_csv = pd.read_csv(io.BytesIO(csv_data_bytes))
            if df_csv.empty:
                return None
            wb = load_workbook(io.BytesIO(excel_template_bytes))
            for m in mappings:
                col = m["csv_column"]
                if col not in df_csv.columns:
                    continue
                val = str(df_csv.loc[0, col]).strip()
                ws = wb[m["excel_sheet"]]
                ws[m["excel_cell"]] = val if " ~ " in val else val + " ~ "
                ws[m["excel_cell"]].alignment = Alignment(horizontal="center", vertical="center")
            out = io.BytesIO()
            wb.save(out)
            out.seek(0)
            return out.getvalue()
        except Exception:
            return None

    st.subheader("Generate OPD.xlsx (configurable, 5-week)")

    if st.button("Generate 4-Sheet OPD (5-week)"):
        excel_template_bytes = generate_opd_workbook(out_df)
        mappings = build_mappings(NUM_WEEKS)
        updated_excel_bytes = update_excel_from_csv(excel_template_bytes, csv_full, mappings)
        if not updated_excel_bytes:
            st.error("Failed to update OPD.xlsx with data.")
            st.stop()
        st.success("✅ OPD.xlsx updated successfully!")
        st.download_button(
            label="⬇️ Download OPD.xlsx",
            data=updated_excel_bytes,
            file_name="OPD_4_sheets_5_weeks.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

# Drop-in replacement / extender for your "Create Individual Schedules" mode
# Adds an optional post-step that converts 4-week tabs to a 5-week layout by
# duplicating the last weekly block's *formatting* (not values) and appending it
# as Week 5. Assumes the weekly block is 24 rows tall starting at row 6 and the
# header (day names + dates) starts at row 2. Adjust the parameters if your
# template differs.

elif mode == "Create Individual Schedules":
    import re
    from collections import defaultdict
    from io import BytesIO
    from zipfile import ZipFile, ZIP_DEFLATED

    import streamlit as st
    from openpyxl import load_workbook, Workbook
    from openpyxl.utils import column_index_from_string
    from openpyxl.cell.cell import MergedCell

    st.subheader("Individual Schedule Creator")

    uploaded = st.file_uploader(
        "Upload the master Excel (.xlsx) with one tab per person",
        type=["xlsx"],
    )

    # --- Parameters for the weekly grid (edit here if your template differs) ---
    WEEK_BLOCK_HEIGHT = 24   # rows per week (AM 10 + PM 10, plus banding)
    WEEK1_START_ROW   = 6    # first content row of Week 1 block
    HEADER1_START_ROW = 2    # row with day names for Week 1 (dates are HEADER+1)
    TARGET_WEEKS      = 5

    extend_tabs = st.checkbox(
        "Convert 4-week tabs to 5 weeks (append Week 5 layout)",
        value=True,
        help=(
            "If your tabs are 4 weeks tall, this will add a 5th week by copying the "
            "last week's *formatting* and leaving cells blank."
        ),
    )

    def copy_sheet_to_new_wb(src_ws):
        """Return BytesIO of a new .xlsx containing src_ws with formatting."""
        wb_new = Workbook()
        ws_new = wb_new.active
        ws_new.title = src_ws.title[:31]

        # Column widths & visibility
        for col_letter, dim in src_ws.column_dimensions.items():
            if dim.width is not None:
                ws_new.column_dimensions[col_letter].width = dim.width
            ws_new.column_dimensions[col_letter].hidden = dim.hidden

        # Row heights & visibility
        for idx, dim in src_ws.row_dimensions.items():
            if dim.height is not None:
                ws_new.row_dimensions[idx].height = dim.height
            ws_new.row_dimensions[idx].hidden = dim.hidden

        # Basic sheet settings (best-effort)
        try:
            ws_new.sheet_format.defaultColWidth = src_ws.sheet_format.defaultColWidth
            ws_new.sheet_format.defaultRowHeight = src_ws.sheet_format.defaultRowHeight
        except Exception:
            pass
        ws_new.freeze_panes = src_ws.freeze_panes
        try:
            ws_new.page_setup.orientation = src_ws.page_setup.orientation
            ws_new.page_setup.fitToWidth = src_ws.page_setup.fitToWidth
            ws_new.page_setup.fitToHeight = src_ws.page_setup.fitToHeight
            ws_new.page_margins = src_ws.page_margins
            ws_new.print_options.horizontalCentered = src_ws.print_options.horizontalCentered
            ws_new.print_options.verticalCentered = src_ws.print_options.verticalCentered
            ws_new.print_area = src_ws.print_area
        except Exception:
            pass

        # Common viewport tweaks
        try:
            # Zoom, widths
            ws_new.sheet_view.zoomScale = 70
            ws_new.column_dimensions["A"].width = 20
            ws_new.column_dimensions["B"].width = 30
            for col in ["C", "D", "E", "F", "G"]:
                ws_new.column_dimensions[col].width = 40
            ws_new.column_dimensions["H"].width = 155
        except Exception:
            pass

        # Copy cells (values + styles)
        for row in src_ws.iter_rows():
            for cell in row:
                # Skip non-master cells from merged ranges
                if isinstance(cell, MergedCell):
                    continue
                col_idx = getattr(cell, "col_idx", None)
                if col_idx is None:
                    col = cell.column
                    col_idx = col if isinstance(col, int) else column_index_from_string(col)
                tgt = ws_new.cell(row=cell.row, column=col_idx, value=cell.value)
                if getattr(cell, "has_style", False):
                    try:
                        from copy import copy
                        if cell.font:
                            tgt.font = copy(cell.font)
                        if cell.fill:
                            tgt.fill = copy(cell.fill)
                        if cell.border:
                            tgt.border = copy(cell.border)
                        if cell.alignment:
                            tgt.alignment = copy(cell.alignment)
                        if cell.protection:
                            tgt.protection = copy(cell.protection)
                        tgt.number_format = cell.number_format
                    except Exception:
                        pass

        # Copy merged ranges
        for merged in list(src_ws.merged_cells.ranges):
            try:
                ws_new.merge_cells(str(merged))
            except Exception:
                pass

        # Copy data validations
        try:
            if src_ws.data_validations and src_ws.data_validations.dataValidation:
                from openpyxl.worksheet.datavalidation import DataValidation
                for dv in src_ws.data_validations.dataValidation:
                    dv_new = DataValidation(
                        type=dv.type,
                        formula1=dv.formula1,
                        formula2=dv.formula2,
                        allow_blank=dv.allow_blank,
                        operator=dv.operator,
                        showDropDown=dv.showDropDown,
                        showErrorMessage=dv.showErrorMessage,
                        errorTitle=dv.errorTitle,
                        error=dv.error,
                        promptTitle=dv.promptTitle,
                        prompt=dv.prompt,
                    )
                    for sqref in getattr(dv, "sqref", []):
                        dv_new.add(sqref)
                    ws_new.add_data_validation(dv_new)
        except Exception:
            pass

        # Filters
        try:
            ws_new.auto_filter.ref = getattr(src_ws.auto_filter, "ref", None)
        except Exception:
            pass

        buf = BytesIO()
        wb_new.save(buf)
        buf.seek(0)
        return buf

    def _copy_row_styles(src_ws, dst_ws, src_r, dst_r, max_col):
        from copy import copy
        for c in range(1, max_col + 1):
            s = src_ws.cell(row=src_r, column=c)
            d = dst_ws.cell(row=dst_r, column=c)
            # copy styles
            if getattr(s, "has_style", False):
                if s.font:
                    d.font = copy(s.font)
                if s.fill:
                    d.fill = copy(s.fill)
                if s.border:
                    d.border = copy(s.border)
                if s.alignment:
                    d.alignment = copy(s.alignment)
                if s.protection:
                    d.protection = copy(s.protection)
                d.number_format = s.number_format
            # clear value by default
            d.value = None

    def extend_sheet_to_5_weeks(xlsx_bytes: bytes,
                                block_height=WEEK_BLOCK_HEIGHT,
                                week1_start=WEEK1_START_ROW,
                                header1_start=HEADER1_START_ROW,
                                target_weeks=TARGET_WEEKS) -> bytes:
        """Append a Week 5 band by cloning the last week's formatting.
        Leaves values blank, preserves styles/merges as best as openpyxl allows.
        Safe to call on already-5+ week sheets (no-op).
        """
        bio = BytesIO(xlsx_bytes)
        wb = load_workbook(bio)
        ws = wb.active

        # Estimate existing weeks from height
        used_last_row = ws.max_row or 0
        if used_last_row < week1_start:
            out = BytesIO()
            wb.save(out)
            out.seek(0)
            return out.getvalue()

        # Heuristic: assume there are 4 weeks if the fifth block start would be beyond max_row
        # You can tighten this by scanning for your black bars or headers.
        week4_start = week1_start + 3 * block_height
        week5_start = week1_start + 4 * block_height
        hdr4_start  = header1_start + 3 * block_height
        hdr5_start  = header1_start + 4 * block_height

        # If there is already a 5th block (content or header rows extend that far), skip.
        if used_last_row >= week5_start + 1:
            out = BytesIO()
            wb.save(out)
            out.seek(0)
            return out.getvalue()

        max_col = ws.max_column or 8

        # Make space for: 2 header rows + one content block
        ws.insert_rows(hdr5_start, amount=2)
        ws.insert_rows(week5_start, amount=block_height)

        # Copy header (day names) row styles from week 4's header row
        _copy_row_styles(ws, ws, hdr4_start, hdr5_start, max_col)
        # Copy date row styles from week 4's date row; leave values empty
        _copy_row_styles(ws, ws, hdr4_start + 1, hdr5_start + 1, max_col)

        # Copy AM/PM band (24 rows) styles from week 4 to week 5; keep col A labels
        for i in range(block_height):
            _copy_row_styles(ws, ws, week4_start + i, week5_start + i, max_col)
            # copy column A label text (e.g., AM/PM bands) if present
            src_label = ws.cell(row=week4_start + i, column=1).value
            ws.cell(row=week5_start + i, column=1).value = src_label

        # Replicate merged ranges that fall entirely within the copied regions
        def _shift_range(rng, row_delta):
            # rng like 'A6:H6' or 'C80:D80'
            from openpyxl.utils.cell import range_boundaries, get_column_letter
            min_col, min_row, max_coln, max_row = range_boundaries(str(rng))
            return f"{get_column_letter(min_col)}{min_row + row_delta}:{get_column_letter(max_coln)}{max_row + row_delta}"

        merged_to_add = []
        for rng in list(ws.merged_cells.ranges):
            # Header row merge copy
            if hdr4_start <= rng.min_row <= rng.max_row <= hdr4_start + 1:
                merged_to_add.append(_shift_range(rng, hdr5_start - hdr4_start))
            # Content block merge copy
            if week4_start <= rng.min_row <= rng.max_row <= week4_start + block_height - 1:
                merged_to_add.append(_shift_range(rng, week5_start - week4_start))
        for addr in merged_to_add:
            try:
                ws.merge_cells(addr)
            except Exception:
                pass

        out = BytesIO()
        wb.save(out)
        out.seek(0)
        return out.getvalue()

    if uploaded is not None:
        # Keep formulas/formatting -> data_only=False
        wb = load_workbook(uploaded, data_only=False)
        st.write(f"Found **{len(wb.sheetnames)}** tabs.")

        if st.button("Split tabs and build ZIP"):
            zip_buf = BytesIO()
            with ZipFile(zip_buf, mode="w", compression=ZIP_DEFLATED) as zf:
                used_names = defaultdict(int)

                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]

                    # Skip truly empty sheets (no cells with value)
                    has_any_value = any(cell.value is not None for row in ws.iter_rows() for cell in row)
                    if not has_any_value:
                        continue

                    # Safe file name
                    base = re.sub(r"[^A-Za-z0-9._-]+", "_", sheet_name).strip("_") or "sheet"
                    used_names[base] += 1
                    safe_name = base if used_names[base] == 1 else f"{base}_{used_names[base]}"

                    # Copy this sheet into its own new workbook (preserving formatting)
                    out_buf = copy_sheet_to_new_wb(ws)

                    # Optional: append Week 5 band
                    if extend_tabs:
                        try:
                            out_buf = BytesIO(extend_sheet_to_5_weeks(out_buf.getvalue()))
                        except Exception:
                            # On any error, fall back to the original copy
                            pass

                    # Add to ZIP
                    zf.writestr(f"{safe_name}.xlsx", out_buf.getvalue())

            zip_buf.seek(0)
            st.download_button(
                label="Download individual schedules (ZIP)",
                data=zip_buf,
                file_name="individual_schedules_formatted.zip",
                mime="application/zip",
            )

# Drop‑in replacement for your "Create Student Schedule" mode
# Converts the whole flow from fixed 4 weeks → configurable N weeks (default 5).
# It keeps your structure, colors, and AM/PM parsing logic — only the math is dynamic.

elif mode == "Create Student Schedule":
    import io
    import re
    import pandas as pd
    import xlsxwriter
    from collections import defaultdict
    from datetime import timedelta
    import streamlit as st
    from openpyxl import load_workbook

    st.subheader("Create Student Schedule")

    # ========= Configure here =========
    NUM_WEEKS = 5  # ← change to 4/5/6... if ever needed
    BLOCK_HEIGHT = 8  # visual rows per week segment in your template pattern
    # ==================================

    def save_to_session(filename, fileobj, namespace="uploaded_files"):
        st.session_state.setdefault(namespace, {})[filename] = fileobj

    # ───────── Helper to load & stash uploads ─────────
    def load_workbook_df(label, types, key):
        upload = st.file_uploader(label, type=types, key=key)
        if not upload:
            st.info(f"Please upload {label}.")
            return None
        st.session_state[f"{key}_file"] = upload
        try:
            if upload.name.lower().endswith(".csv"):
                return pd.read_csv(upload)
            else:
                return pd.read_excel(upload)
        except Exception as e:
            st.error(f"Error loading {upload.name}: {e}")
            return None

    def _seq(start, step, n):
        """Helper: n values starting at start, step apart (1‑based row numbers)."""
        return [start + i * step for i in range(n)]

    def create_ms_schedule_template(students, dates):
        buf = io.BytesIO()
        wb = xlsxwriter.Workbook(buf, {"in_memory": True})

        # — Formats — (unchanged palette)
        f1 = wb.add_format({"font_size": 14, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "text_wrap": True, "bg_color": "#FEFFCC", "border": 1})
        f2 = wb.add_format({"font_size": 10, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "yellow", "bg_color": "black", "border": 1, "text_wrap": True})
        f3 = wb.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "bg_color": "#FFC7CE", "border": 1})
        f4 = wb.add_format({"num_format": "mm/dd/yyyy", "font_size": 12, "bold": 1, "align": "center",
                            "valign": "vcenter", "font_color": "black", "bg_color": "#F4F6F7", "border": 1})
        f5 = wb.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "bg_color": "#F4F6F7", "border": 1})
        f6 = wb.add_format({"bg_color": "black", "border": 1})
        f7 = wb.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "bg_color": "#90EE90", "border": 1})
        f8 = wb.add_format({"font_size": 12, "bold": 1, "align": "center", "valign": "vcenter",
                            "font_color": "black", "bg_color": "#89CFF0", "border": 1})

        days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        # Derive all row anchors for N weeks (1‑based for A1 merges; numeric writes may be slightly offset like your original)
        start_rows = _seq(2, BLOCK_HEIGHT, NUM_WEEKS)           # day headers row per week (2,10,18,26,34)
        week_label_rows = _seq(4, BLOCK_HEIGHT, NUM_WEEKS)      # "Week i" labels (4,12,20,28,36)
        am_rows = _seq(6, BLOCK_HEIGHT, NUM_WEEKS)              # AM label row (6,14,22,30,38)
        pm_rows = _seq(7, BLOCK_HEIGHT, NUM_WEEKS)              # PM label row (7,15,23,31,39)
        green_fillers = _seq(8, BLOCK_HEIGHT, NUM_WEEKS)        # green row anchors (8,16,24,32,40)
        separators = _seq(10, BLOCK_HEIGHT, NUM_WEEKS)          # black bar rows (10,18,26,34,42)

        # Optional: extend your due text list to 5, with a safe fallback
        due_texts = [
            "Quiz 1 Due",
            "Quiz 2, Pediatric Documentation #1, 1 Clinical Encounter Log Due",
            "Quiz 3 Due",
            "Quiz 4, Pediatric Documentation #2, Social Drivers of Health Assessment Form, Developmental Assessment of Pediatric Patient Form, All Clinical Encounter Logs are Due!",
            "Final Evaluations & Wrap‑Up Due",  # Week 5 (new)
        ]

        for name in students:
            title = name[:31].replace("/", "-").replace("\\", "-")
            ws = wb.add_worksheet(title)
            ws.set_zoom(70)

            # Header
            ws.merge_range("A1:A2", "Student Name:", f1)
            ws.merge_range("B1:B2", title, f1)
            note = (
                "*Note* Asynchronous time is for coursework only. During this time period, "
                "we expect students to do coursework, be available for any additional educational "
                "activities, and any extra clinical time that may be available. If the student is not "
                "available during this time period and has not made an absence request, the student "
                "will be cited for unprofessionalism and will risk failing the course."
            )
            ws.merge_range("C1:H2", note, f2)

            # Column widths & row height
            ws.set_column("A:A", 20)
            ws.set_column("B:B", 30)
            ws.set_column("C:G", 40)
            ws.set_column("H:H", 155)
            ws.set_row(0, 37.25)

            # Day names + dates
            date_idx = 0
            for row in start_rows:
                # day names (row, B..H). Using numeric write like your original
                for col_offset, day in enumerate(days, start=1):
                    ws.write(row, col_offset, day, f3)
                # dates directly beneath
                for col_offset in range(7):
                    if date_idx < len(dates):
                        ws.write(row + 1, col_offset + 1, dates[date_idx], f4)
                        date_idx += 1

            # Week labels
            for i, row in enumerate(week_label_rows):
                ws.write(f"A{row}", f"Week {i+1}", f3)

            # AM / PM labels
            for r in am_rows:
                ws.write(f"A{r}", "AM", f3)
            for r in pm_rows:
                ws.write(f"A{r}", "PM", f3)

            # Fill AM/PM blocks with Asynchronous Time (cols B..H)
            for i in range(NUM_WEEKS):
                am_r = 5 + i * BLOCK_HEIGHT  # keep your original indexing style
                pm_r = 6 + i * BLOCK_HEIGHT
                for col in range(1, 8):
                    ws.write(am_r, col, "Asynchronous Time", f5)
                    ws.write(pm_r, col, "Asynchronous Time", f5)

            # Separators (black bars)
            for sep in separators:
                ws.merge_range(f"A{sep}:H{sep}", "", f6)

            # Green filler rows
            for filler in green_fillers:
                for col in range(8):
                    ws.write(filler, col, " ", f7)

            # Assignment‑due rows
            for i, base in enumerate(green_fillers):
                ws.write(f"A{base}", "ASSIGNMENT DUE:", f8)
                for col in range(1, 8):
                    if col == 5:
                        ws.write(base - 1, col, "Ask for Feedback!", f8)
                    elif col == 7:
                        ws.write(base - 1, col, due_texts[i] if i < len(due_texts) else "", f8)
                    else:
                        ws.write(base - 1, col, " ", f8)

        wb.close()
        buf.seek(0)
        return buf

    def _cluster_blocks(marker_rows):
        marker_rows = sorted(marker_rows)
        if not marker_rows:
            return []
        blocks, curr = [], [marker_rows[0]]
        for r in marker_rows[1:]:
            if r == curr[-1] + 1:
                curr.append(r)
            else:
                blocks.append(curr)
                curr = [r]
        blocks.append(curr)
        return blocks

    def assign_preceptors_all_weeks_am(opd_file, ms_file):
        """Copy AM 'Preceptor ~ Student' assignments from OPD (B..H) into each student's tab.
        Maps week blocks → target rows [6,14,22,30,38] for 5 weeks.
        """
        opd_wb = load_workbook(opd_file, data_only=True)
        ms_wb = load_workbook(ms_file)

        target_ms_rows = _seq(6, BLOCK_HEIGHT, NUM_WEEKS)  # [6,14,22,30,38]

        for site in opd_wb.sheetnames:
            ws_opd = opd_wb[site]
            am_rows = [c.row for c in ws_opd['A'] if isinstance(c.value, str) and re.match(r"^\s*AM\b", c.value, re.IGNORECASE)]
            blocks = _cluster_blocks(am_rows)
            for week_idx, block in enumerate(blocks[:NUM_WEEKS]):
                ms_row = target_ms_rows[week_idx]
                for col in range(2, 9):  # B..H
                    for r in block:
                        val = ws_opd.cell(row=r, column=col).value
                        if not val or "~" not in str(val):
                            continue
                        pre, student = [s.strip() for s in str(val).split("~", 1)]
                        if not student or student not in ms_wb.sheetnames:
                            continue
                        ws_ms = ms_wb[student]
                        ws_ms.cell(row=ms_row, column=col).value = f"{pre} - [{site}]"

        out = io.BytesIO()
        ms_wb.save(out)
        out.seek(0)
        return out

    def assign_preceptors_all_weeks_pm(opd_file, ms_file):
        """Copy PM assignments; week rows [7,15,23,31,39] for 5 weeks."""
        opd_wb = load_workbook(opd_file, data_only=True)
        ms_wb = load_workbook(ms_file)

        target_ms_rows = _seq(7, BLOCK_HEIGHT, NUM_WEEKS)  # [7,15,23,31,39]

        for site in opd_wb.sheetnames:
            ws_opd = opd_wb[site]
            pm_rows = [c.row for c in ws_opd['A'] if isinstance(c.value, str) and re.match(r"^\s*PM\b", c.value, re.IGNORECASE)]
            blocks = _cluster_blocks(pm_rows)
            for week_idx, block in enumerate(blocks[:NUM_WEEKS]):
                ms_row = target_ms_rows[week_idx]
                for col in range(2, 9):  # B..H
                    for r in block:
                        val = ws_opd.cell(row=r, column=col).value
                        if not val or "~" not in str(val):
                            continue
                        pre, student = [s.strip() for s in str(val).split("~", 1)]
                        if not student or student not in ms_wb.sheetnames:
                            continue
                        ws_ms = ms_wb[student]
                        ws_ms.cell(row=ms_row, column=col).value = f"{pre} - [{site}]"

        out = io.BytesIO()
        ms_wb.save(out)
        out.seek(0)
        return out

    def detect_shift_conflicts(opd_file):
        """Same as before, but scans NUM_WEEKS blocks instead of 4."""
        wb = load_workbook(opd_file, data_only=True)
        days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        am_marker = re.compile(r"^\s*AM\b", re.IGNORECASE)
        pm_marker = re.compile(r"^\s*PM\b", re.IGNORECASE)
        conflicts = []

        def find_blocks(ws, marker_re):
            rows = [c.row for c in ws['A'] if isinstance(c.value, str) and marker_re.match(c.value)]
            return _cluster_blocks(rows)[:NUM_WEEKS]

        tpl = wb[wb.sheetnames[0]]
        am_blocks = find_blocks(tpl, am_marker)
        pm_blocks = find_blocks(tpl, pm_marker)

        for shift, blocks in (("AM", am_blocks), ("PM", pm_blocks)):
            for week_idx, block_rows in enumerate(blocks, start=1):
                for day_idx, day_name in enumerate(days):
                    col = 2 + day_idx
                    locs = defaultdict(list)
                    for sheet in wb.sheetnames:
                        ws = wb[sheet]
                        for r in block_rows:
                            raw = ws.cell(row=r, column=col).value
                            text = str(raw or "")
                            if "~" not in text:
                                continue
                            pre, student = [s.strip() for s in text.split("~", 1)]
                            if not student:
                                continue
                            coord = ws.cell(row=r, column=col).coordinate
                            locs[student].append((sheet, coord))
                    for student, occ in locs.items():
                        if len(occ) > 1:
                            conflicts.append({
                                "student": student,
                                "week": week_idx,
                                "day": day_name,
                                "shift": shift,
                                "occurrences": occ,
                            })
        return conflicts

    # ───────── Load OPD & Rotation Schedule ─────────
    df_opd = load_workbook_df("Upload OPD.xlsx file", ["xlsx"], key="opd_main")
    df_rot = load_workbook_df("Upload Rotation Schedule (.xlsx or .csv)", ["xlsx", "csv"], key="rot_main")

    # ───────── Check for duplicates ─────────
    if df_opd is not None:
        conflicts = detect_shift_conflicts(st.session_state["opd_main_file"])
        if conflicts:
            for c in conflicts:
                occ_str = "; ".join(f"{sheet}@{coord}" for sheet, coord in c["occurrences"])
                st.warning(
                    f"⚠️ Week {c['week']} {c['day']} {c['shift']}: "
                    f"{c['student']} double‑booked ({occ_str})"
                )
        else:
            st.success("No AM/PM shift conflicts detected.")

    # ───────── Build, Assign & Download ─────────
    if df_opd is not None and df_rot is not None:
        df_rot["start_date"] = pd.to_datetime(df_rot["start_date"])
        monday = df_rot["start_date"].min()
        monday = monday - pd.Timedelta(days=monday.weekday())  # normalize to Monday
        dates = pd.date_range(start=monday, periods=NUM_WEEKS * 7, freq="D").tolist()

        students = df_rot["legal_name"].dropna().unique().tolist()

        if st.button("Create & Download Fully‑Populated MS_Schedule"):
            blank_buf = create_ms_schedule_template(students, dates)
            am_buf = assign_preceptors_all_weeks_am(
                opd_file=st.session_state["opd_main_file"], ms_file=blank_buf
            )
            full_buf = assign_preceptors_all_weeks_pm(
                opd_file=st.session_state["opd_main_file"], ms_file=am_buf
            )
            st.download_button(
                "Download MS_Schedule.xlsx",
                data=full_buf.getvalue(),
                file_name="MS_Schedule.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
    else:
        st.info("Please upload both OPD.xlsx and the rotation schedule above to proceed.")
